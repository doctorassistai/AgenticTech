"""
services/docx_generator.py

Two Word-document exports for a case:

  build_case_docx(case_data)
      Generic, non-templated .docx covering every case field + the parsed
      conclusion — good for copy-paste into whatever format is needed.

  build_formatted_docx(case_data)
      Converts the SAME rendered HTML the PDF uses (via
      pdf_generator.render_case_html) into an editable .docx, so the
      doctor gets a Word document that mirrors the actual insurer/TPA
      report layout (headings, tables, images) instead of a flat field
      dump. This is the "same format as the PDF, but editable" option.

Deliberately excludes two things that exist in the DB but would only
duplicate what's already shown elsewhere:
  - raw_llama_markdown        → raw OCR source text; already fully absorbed
                                 into the structured fields below.
  - case_documents[].extracted_data / extracted_flat
                                → per-document extraction results that were
                                  already merged into the top-level case
                                  fields during ingestion. Only the document
                                  *names* are listed, as an enclosures list.

Requires: pip install python-docx beautifulsoup4 htmldocx premailer
"""

import os
import re
import base64
import logging
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup, NavigableString
from premailer import transform as inline_css

logger = logging.getLogger(__name__)

BASE_DIR      = os.path.dirname(os.path.dirname(__file__))
GENERATED_DIR = os.path.join(BASE_DIR, "generated")
os.makedirs(GENERATED_DIR, exist_ok=True)


def _filename_base(case_data: dict) -> str:
    """Same rule as pdf_generator.get_pdf_filename_base: insurerRef first,
    falling back to caseId, sanitized for filesystem use."""
    raw = str(case_data.get("insurerRef") or case_data.get("caseId") or "unknown").strip()
    safe = re.sub(r"[^A-Za-z0-9_\-]+", "_", raw).strip("_")
    return safe or "unknown"


# ─────────────────────────────────────────────────────────────────────────────
# GENERIC STYLE / TABLE HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h


def _humanize_key(key: str) -> str:
    """turns 'accident_date_time' or 'roomRentLimit' into 'Accident Date Time' / 'Room Rent Limit'"""
    s = re.sub(r"(?<!^)(?=[A-Z])", " ", key)      # camelCase -> spaced
    s = s.replace("_", " ")
    return s.strip().title()


def _stringify(value):
    if isinstance(value, (list, tuple)):
        return "; ".join(str(v) for v in value if v not in (None, ""))
    return str(value)


def _add_field_row(table, label, value):
    row = table.add_row()
    row.cells[0].text = str(label or "")
    row.cells[1].text = _stringify(value) if value not in (None, "", [], {}) else "—"
    if row.cells[0].paragraphs[0].runs:
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    if row.cells[1].paragraphs[0].runs:
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)


def _add_details_table(doc, rows, skip_if_empty=True):
    """rows: list of (label, value) tuples. Skips rows with no value."""
    visible = [(l, v) for l, v in rows if v not in (None, "", {}, [])]
    if not visible:
        if not skip_if_empty:
            doc.add_paragraph("— No data captured —")
        return False
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for label, value in visible:
        _add_field_row(table, label, value)
    doc.add_paragraph("")
    return True


def _add_optional_dict_section(doc, title, data: dict, exclude_keys=None, label_overrides=None):
    """
    Renders any non-empty dict as a details table under its own heading,
    skipping the whole section if there's nothing in it. Used for the
    schema-optional blocks (accidentDetails, deathDetails, obstetricDetails,
    riskDetails, locationDetails, etc.) so we never print an empty section.
    """
    if not data:
        return
    exclude_keys = exclude_keys or set()
    label_overrides = label_overrides or {}
    rows = [
        (label_overrides.get(k, _humanize_key(k)), v)
        for k, v in data.items()
        if k not in exclude_keys and v not in (None, "", {}, [])
    ]
    if not rows:
        return
    _set_heading(doc, title)
    _add_details_table(doc, rows)


def _add_paragraph_block(doc, label, text, heading_level=3):
    if not text:
        return
    doc.add_paragraph(label, style=f"Heading {heading_level}")
    doc.add_paragraph(text)


# ─────────────────────────────────────────────────────────────────────────────
# HTML CONCLUSION → DOCX  (rich editor output: <p>, <b>/<strong>, <i>/<em>,
# <ul>/<ol>/<li>, <br>, and <img src="data:image/...;base64,...">)
# ─────────────────────────────────────────────────────────────────────────────

def _add_image_from_data_uri(doc, src):
    try:
        _, b64data = src.split(",", 1)
        img_bytes = base64.b64decode(b64data)
        doc.add_picture(BytesIO(img_bytes), width=Inches(5.5))
    except Exception as exc:
        logger.warning("Could not embed conclusion image: %s", exc)


def _render_inline(paragraph, node):
    for child in node.children:
        name = getattr(child, "name", None)
        if name is None:
            if str(child).strip():
                paragraph.add_run(str(child))
        elif name in ("b", "strong"):
            paragraph.add_run(child.get_text()).bold = True
        elif name in ("i", "em"):
            paragraph.add_run(child.get_text()).italic = True
        elif name == "br":
            paragraph.add_run().add_break()
        elif name == "img":
            continue  # images pulled out at block level
        else:
            paragraph.add_run(child.get_text())


def _render_html_node(doc, node):
    if isinstance(node, NavigableString):
        text = str(node).strip()
        if text:
            doc.add_paragraph(text)
        return

    tag = node.name
    if tag in ("p", "div"):
        img = node.find("img")
        if img and img.get("src", "").startswith("data:image"):
            _add_image_from_data_uri(doc, img["src"])
        if node.get_text(strip=True):
            p = doc.add_paragraph()
            _render_inline(p, node)
    elif tag == "img":
        if node.get("src", "").startswith("data:image"):
            _add_image_from_data_uri(doc, node["src"])
    elif tag in ("ul", "ol"):
        style = "List Bullet" if tag == "ul" else "List Number"
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style=style)
            _render_inline(p, li)
    elif tag == "br":
        pass
    else:
        for child in node.children:
            _render_html_node(doc, child)


def _detect_verdict(text: str) -> str:
    """Mirrors the frontend's detectVerdict — last occurrence of the two words wins."""
    if not text:
        return ""
    lower = text.lower()
    li_s = lower.rfind("suspected")
    li_g = lower.rfind("genuine")
    if li_s == -1 and li_g == -1:
        return ""
    return "SUSPECTED" if li_s > li_g else "GENUINE"


def _add_conclusion_section(doc, conclusion_raw):
    _set_heading(doc, "Investigation Conclusion", level=1)

    if not conclusion_raw:
        doc.add_paragraph("No conclusion provided.")
        return

    is_html = bool(re.search(r"<[a-zA-Z][^>]*>", conclusion_raw))

    if is_html:
        soup = BeautifulSoup(conclusion_raw, "html.parser")
        root = soup.body or soup
        for child in root.children:
            _render_html_node(doc, child)
        return

    # Plain text (TRIGGER:/SECTION legacy format, or the new unified
    # SECTION 1/2/3 format used in the sample case) — surface structural
    # markers as headings, drop the ==== rule lines.
    for line in conclusion_raw.splitlines():
        stripped = line.strip()
        if not stripped or set(stripped) <= {"="}:
            continue
        if (
            re.match(r"^SECTION\s+\d", stripped, re.IGNORECASE)
            or stripped.startswith("TRIGGER:")
            or stripped.startswith("OVERALL CASE VERDICT")
            or re.match(r"^[A-D]\.\s+[A-Z]", stripped)          # "A. DISCREPANCIES" etc.
        ):
            _set_heading(doc, stripped, level=2)
        else:
            doc.add_paragraph(stripped)


# ─────────────────────────────────────────────────────────────────────────────
# CURATED "ADDITIONAL INVESTIGATION FACTS" FROM pre_extracted_facts
# ─────────────────────────────────────────────────────────────────────────────

_FACT_KEYS_ALREADY_SHOWN = {
    "admission_date", "admission_time", "discharge_date", "discharge_time",
    "patient_name", "patient_age_gender", "hospital_name", "hospital_address",
    "treating_doctor", "doctor_reg_number", "surgeon_name", "surgery_date_time",
    "provisional_diagnosis", "final_diagnosis", "chief_complaints", "past_history",
    "bill_amount", "gross_bill_amount", "discount_amount", "payment_mode",
    "room_type", "accident_date_time", "accident_place", "accident_narration",
    "investigating_agency", "field_officer_name", "data_collected_from_name",
    "data_collected_from_designation", "data_collected_from_phone",
    "vitals_on_admission", "vitals_at_discharge",
}

_FACT_GROUPS = [
    ("Identifiers", [
        "ip_number", "uhid_number", "guardian_name", "bill_number",
    ]),
    ("Accident / Incident Verification", [
        "mlc_number", "mlc_registered", "mlc_collected", "fir_number",
        "police_station", "witness_name", "witness_statement", "brought_by",
        "first_aid_hospital", "first_aid_details", "vehicle_type",
        "helmet_worn", "seatbelt_worn", "alcohol_smell_noted",
        "alcohol_test_done", "alcohol_test_result", "intoxication_mentioned",
    ]),
    ("Hospital / Records Verification", [
        "hospital_registration_number", "hospital_registration_validity",
        "hospital_registration_valid", "hospital_reg_valid_till",
        "hospital_reg_issuing_authority", "hospital_bed_strength",
        "hospital_nabh", "hospital_watchlist_status", "hospital_empanelment_status",
        "first_registration_date", "inhouse_lab", "inhouse_lab_present",
        "lab_register_collected", "lab_register_attached", "lab_photos_attached",
        "lab_vicinity_to_hospital", "pharmacy_register_collected",
        "ip_register_collected", "ip_register_attached", "ot_register_attached",
        "reg_certificate_attached", "tariff_attached", "operation_record_attached",
        "physical_visit_confirmed", "postoperative_period",
        "anaesthesia_type", "anaesthetist_name",
    ]),
    ("Financial Verification", [
        "patient_stated_bill_amount", "net_amount_received", "mode_of_payment",
        "room_tariff_per_day", "icu_charges", "room_charges", "medicine_charges",
        "lab_charges", "bill_breakdown", "bill_breakdown_items",
        "contradictions_found",
    ]),
    ("Policy / PED History", [
        "policy_inception_date", "policy_start_date", "policy_end_date",
        "ped_declared_at_proposal", "ped_mentioned_in_records",
        "pre_admission_opd_visits", "pre_admission_prescriptions",
        "opd_history_before_admission", "short_duration_policy",
        "policy_type", "previous_claims", "claim_frequency",
    ]),
    ("Employment (if applicable)", [
        "employer_name", "employee_id", "employment_verification_done",
    ]),
    ("Death Claim (if applicable)", [
        "death_date", "death_time", "death_place", "cause_of_death",
        "postmortem_done", "postmortem_report", "death_certificate_available",
        "beneficiary_name", "beneficiary_relationship",
        "suicidal_history", "psychiatric_history", "alcohol_history",
    ]),
    ("Investigator's Notes", [
        "investigating_agency_type", "verification_mode",
        "field_officer_hospital_opinion", "field_officer_member_opinion",
        "discrepancies_verbatim", "final_verdict_verbatim",
    ]),
]


def _add_additional_facts_section(doc, facts: dict):
    if not facts:
        return

    any_rendered = False
    for group_title, keys in _FACT_GROUPS:
        rows = [
            (_humanize_key(k), facts.get(k))
            for k in keys
            if k not in _FACT_KEYS_ALREADY_SHOWN and facts.get(k) not in (None, "", [], {})
        ]
        if not rows:
            continue
        if not any_rendered:
            _set_heading(doc, "Additional Investigation Facts")
            doc.add_paragraph(
                "Verification detail captured during document extraction that "
                "isn't already covered in the sections above.",
            ).italic = True
            any_rendered = True
        doc.add_paragraph(group_title, style="Heading 3")
        _add_details_table(doc, rows)


# ─────────────────────────────────────────────────────────────────────────────
# TRIGGERS — dedupe the 4 overlapping fields down to one clean list
# ─────────────────────────────────────────────────────────────────────────────

def _add_triggers_section(doc, case_data: dict):
    conclusion_labels = case_data.get("conclusionTriggerLabels") or []
    claim_triggers     = case_data.get("claimTriggers") or []
    suggested_triggers = case_data.get("suggestedTriggers") or []

    if not (conclusion_labels or claim_triggers or suggested_triggers):
        return

    _set_heading(doc, "Investigation Triggers")

    if conclusion_labels:
        primary_label, primary_list = "Triggers Assessed in Conclusion", conclusion_labels
    elif claim_triggers:
        primary_label, primary_list = "Triggers Actioned", [_humanize_key(t) for t in claim_triggers]
    else:
        primary_label, primary_list = "Suggested Triggers", [_humanize_key(t) for t in suggested_triggers]

    doc.add_paragraph(primary_label, style="Heading 3")
    for t in primary_list:
        doc.add_paragraph(t, style="List Bullet")

    if suggested_triggers and claim_triggers:
        extra = [t for t in suggested_triggers if t not in claim_triggers]
        if extra:
            doc.add_paragraph("Other Suggested Triggers (not actioned)", style="Heading 3")
            for t in extra:
                doc.add_paragraph(_humanize_key(t), style="List Bullet")

    doc.add_paragraph("")


# ─────────────────────────────────────────────────────────────────────────────
# ENCLOSURES
# ─────────────────────────────────────────────────────────────────────────────

def _add_enclosures_section(doc, case_data: dict):
    docs = ((case_data.get("case_documents") or {}).get("documents")) or []
    enclosures_text = case_data.get("enclosures")

    if not docs and not enclosures_text:
        return

    _set_heading(doc, "Enclosures / Source Documents")

    if docs:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Document", "Label"
        for d in docs:
            row = table.add_row()
            row.cells[0].text = d.get("file_name", "—")
            row.cells[1].text = d.get("display_label", "—")
        doc.add_paragraph("")

    if enclosures_text:
        doc.add_paragraph("Additional Documents Collected", style="Heading 3")
        doc.add_paragraph(enclosures_text)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN BUILDER — generic, non-templated export
# ─────────────────────────────────────────────────────────────────────────────

def build_case_docx(case_data: dict) -> str:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    hospital      = case_data.get("hospitalDetails", {})          or {}
    critical      = case_data.get("criticalDetails", {})          or {}
    additional    = case_data.get("additionalMedicalDetails", {}) or {}
    billing       = case_data.get("billingDetails", {})           or {}
    investigation = case_data.get("investigationDetails", {})     or {}
    medical_staff = case_data.get("medicalStaff", {})              or {}
    accident      = case_data.get("accidentDetails", {})          or {}
    policy        = case_data.get("policyDetails", {})            or {}
    cashless      = case_data.get("cashlessDetails", {})          or {}
    checklist     = case_data.get("checklist", {})                or {}
    interview     = case_data.get("interviewDetails", {})         or {}
    risk          = case_data.get("riskDetails", {})              or {}
    location      = case_data.get("locationDetails", {})          or {}

    # ── Title block ──────────────────────────────────────────────────────
    title = doc.add_heading("Investigation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref = case_data.get("insurerRef") or case_data.get("caseId", "")
    run = sub.add_run(f"{case_data.get('insurer', '') or ''}  ·  Claim ID: {ref}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    verdict = _detect_verdict(case_data.get("conclusion") or "")
    if verdict:
        vsub = doc.add_paragraph()
        vsub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vrun = vsub.add_run(f"VERDICT: {verdict}")
        vrun.bold = True
        vrun.font.size = Pt(11)
        vrun.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B) if verdict == "SUSPECTED" else RGBColor(0x27, 0xAE, 0x60)

    doc.add_paragraph("")

    if case_data.get("description"):
        _set_heading(doc, "Case Summary")
        doc.add_paragraph(case_data["description"])

    _set_heading(doc, "Claim Details")
    _add_details_table(doc, [
        ("Insurer", case_data.get("insurer")),
        ("TPA", case_data.get("tpaName") or cashless.get("tpaName")),
        ("Claim ID / Insurer Ref", case_data.get("insurerRef")),
        ("Case Status", case_data.get("status")),
        ("Claim Mode", case_data.get("claimMode")),
        ("Claim Subtype", case_data.get("claimSubtype")),
        ("Claim Priority", case_data.get("claimPriority")),
        ("Claim Source", case_data.get("claimSource")),
        ("SLA Category", case_data.get("slaCategory")),
        ("Claimed Amount", case_data.get("claimedAmount")),
        ("Sum Insured", case_data.get("sumInsured")),
        ("Date of Incident", case_data.get("dateOfIncident")),
        ("Date of Intimation", case_data.get("dateOfIntimation")),
        ("Target Date", case_data.get("targetDate")),
        ("Insurer Contact", case_data.get("insurerContact") or case_data.get("insurerContactInfo")),
    ])

    if case_data.get("emailInstructions"):
        _set_heading(doc, "Insurer Investigation Instructions")
        doc.add_paragraph(case_data["emailInstructions"])

    _add_triggers_section(doc, case_data)

    _set_heading(doc, "Claimant / Insured Details")
    _add_details_table(doc, [
        ("Name of Insured", case_data.get("claimantName")),
        ("Age", case_data.get("claimantAge")),
        ("Relationship to Policyholder", case_data.get("relationship")),
        ("Address", case_data.get("claimantAddress")),
        ("Contact", case_data.get("claimantMobile")),
        ("Alternate Contact", case_data.get("altContact")),
        ("ID Proof Type", case_data.get("idProofType")),
        ("ID Proof Number", case_data.get("idProofNumber")),
        ("City", case_data.get("city")),
        ("District", case_data.get("district")),
        ("PIN Code", case_data.get("pinCode")),
        ("Type of Policy", case_data.get("policyType")),
    ])

    _set_heading(doc, "Policy Details")
    _add_details_table(doc, [
        ("Policy Number", case_data.get("policyNumber")),
        ("Coverage Type", policy.get("coverageType")),
        ("Policy Start Date", policy.get("startDate")),
        ("Policy End Date", policy.get("endDate")),
        ("First Commencement Date", policy.get("inceptionDate")),
        ("Room Rent Limit", policy.get("roomRentLimit")),
        ("Pre-Existing Disease", policy.get("preExistingDisease")),
        ("Sum Insured", policy.get("sumInsured")),
        ("Original Sum Insured", policy.get("originalSumInsured")),
        ("Sum Insured Enhancement", policy.get("sumInsuredEnhancement")),
        ("No. of Policy Years", policy.get("policyYears")),
        ("Agent / Broker", policy.get("agentBroker")),
    ])

    _set_heading(doc, "Hospital Details")
    _add_details_table(doc, [
        ("Hospital Name", hospital.get("name")),
        ("Hospital Address", hospital.get("address")),
        ("Hospital Type", hospital.get("type")),
        ("Department", hospital.get("department")),
        ("Registration Number", hospital.get("registrationNumber")),
        ("PPN / Non-PPN", hospital.get("ppnStatus")),
        ("Contact Person", hospital.get("contactPerson")),
        ("Hospital Contact Number", hospital.get("hospitalContactNumber")),
        ("Hospital Email", hospital.get("hospitalEmail")),
        ("Date of Admission", hospital.get("admissionDate")),
        ("Date of Discharge", hospital.get("dischargeDate")),
        ("Treating Doctor", hospital.get("doctorName")),
        ("Doctor Registration No.", hospital.get("doctorRegNumber")),
    ])

    _set_heading(doc, "Diagnosis & Clinical Details")
    _add_details_table(doc, [
        ("Cashless / Non-Cashless", case_data.get("claimMode")),
        ("Diagnosis", critical.get("diagnosis")),
        ("Procedure", critical.get("procedure")),
        ("Implants", critical.get("implants")),
        ("Surgery Date", critical.get("surgeryDate")),
        ("First Consultation Date", additional.get("firstConsultationDate")),
        ("ICU Admission Details", cashless.get("icuDetails")),
        ("Estimated Treatment Cost", cashless.get("estimatedCost")),
        ("Amount Authorized", cashless.get("amountAuthorized")),
        ("Pre-Auth Details", cashless.get("preAuthDetails")),
        ("Referral Doctor", additional.get("referralDoctor")),
        ("Date of First Onset", case_data.get("dateOfIncident")),
    ])

    if additional.get("chiefComplaints"):
        doc.add_paragraph("Chief Complaints", style="Heading 3")
        for c in additional["chiefComplaints"]:
            doc.add_paragraph(c, style="List Bullet")

    for label, key in [
        ("Clinical Summary", "clinicalSummary"),
        ("Diagnosis Summary", "diagnosisSummary"),
        ("General Examination", "generalExamination"),
        ("Local Examination", "localExamination"),
        ("Vitals", "vitals"),
        ("Past History", "pastHistory"),
        ("Initial Treatment Details", "initialTreatment"),
        ("Pre-Hospitalisation Details", "preHospitalisationDetails"),
        ("Post-Hospitalisation Details", "postHospitalisationDetails"),
        ("Investigations Suggesting Diagnosis", "investigationsSuggestingDiagnosis"),
        ("Investigator's Hospital-Side Opinion", "investigatorHospitalOpinion"),
        ("Investigator's Member-Side Opinion", "investigatorMemberOpinion"),
    ]:
        _add_paragraph_block(doc, label, additional.get(key))

    if billing:
        _set_heading(doc, "Billing Details")
        _add_details_table(doc, [
            ("Gross Bill Amount", billing.get("grossAmount")),
            ("Discount Amount", billing.get("discountAmount")),
            ("Final Bill Amount", billing.get("finalBillAmount")),
            ("Payment Mode", billing.get("paymentMode")),
            ("Room Type", billing.get("roomType")),
        ])

    _add_optional_dict_section(doc, "Accident Details", accident)

    for title, block in [
        ("Death Claim Details", case_data.get("deathDetails")),
        ("Obstetric Details", case_data.get("obstetricDetails")),
        ("Railway Claim Details", case_data.get("railwayDetails")),
        ("Reimbursement Details", case_data.get("reimbursementDetails")),
        ("Medical Staff", medical_staff),
        ("Risk Assessment", risk),
        ("Location Details", location),
    ]:
        _add_optional_dict_section(doc, title, block)

    checklist_labels = [
        ("idProofInsured", "1. ID proof of insured patient"),
        ("hospitalExistence", "2. Hospital existence & registration"),
        ("admissionVerified", "3. Admission of patient on stated dates"),
        ("treatmentParticulars", "4. Treatment particulars"),
        ("copyOfICP", "5. Copy of ICP"),
        ("labVicinity", "6. Lab in vicinity / far off"),
        ("labRegistersVerified", "7. Lab registers verified"),
        ("billsReceipts", "8. Bills & receipts verified"),
        ("medicinePurchases", "9. Medicine shop purchases"),
        ("signatureMatching", "10. Signature matching"),
        ("otReceiptBooks", "11. OT / receipt book copies"),
        ("anyOther", "12. Any other"),
    ]
    if any(checklist.get(k) for k, _ in checklist_labels):
        _set_heading(doc, "Checklist")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Item", "Status"
        for k, label in checklist_labels:
            _add_field_row(table, label, checklist.get(k, "NA"))
        doc.add_paragraph("")

    _set_heading(doc, "Interview & Data Collection")
    _add_details_table(doc, [
        ("Data Collected From", investigation.get("dataCollectedFrom")),
        ("Investigator Name", investigation.get("investigatorName")),
        ("Investigator Designation", investigation.get("investigatorDesignation")),
    ])
    if interview.get("neighbours"):
        doc.add_paragraph("Neighbours", style="Heading 3")
        doc.add_paragraph(interview.get("neighbours"))
    for k, v in interview.items():
        if k == "neighbours":
            continue
        _add_paragraph_block(doc, _humanize_key(k), v)

    _add_additional_facts_section(doc, case_data.get("pre_extracted_facts") or {})

    _add_enclosures_section(doc, case_data)

    _set_heading(doc, "Report Meta")
    _add_details_table(doc, [
        ("Report Date", case_data.get("reportDate")),
        ("Case Created", case_data.get("createdAt")),
        ("Case Last Updated", case_data.get("updatedAt")),
        ("Conclusion Generated At", case_data.get("conclusionGeneratedAt")),
        ("Assignment Notes", case_data.get("assignmentNotes")),
        ("Tags", case_data.get("tags")),
    ])

    doc.add_page_break()

    _add_conclusion_section(doc, case_data.get("conclusion") or "")

    filename_base = _filename_base(case_data)
    output_path = os.path.join(GENERATED_DIR, f"{filename_base}_edited.docx")
    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# HTML PREP FOR htmldocx — inline CSS, strip print-only positioning,
# convert flex layouts to tables (htmldocx has none of these natively)
# ─────────────────────────────────────────────────────────────────────────────

_FLEX_RE  = re.compile(r"display\s*:\s*flex", re.IGNORECASE)
_FIXED_RE = re.compile(r"position\s*:\s*fixed", re.IGNORECASE)
_ABS_RE   = re.compile(r"position\s*:\s*absolute", re.IGNORECASE)


def _flex_to_table(soup: BeautifulSoup) -> None:
    """
    htmldocx has no flexbox support — a flex container's children just get
    stacked as separate block paragraphs instead of sitting side by side.
    After premailer has inlined every style attribute, find any element
    whose inline style contains display:flex and rewrite it as a borderless
    one-row HTML table, one <td> per direct child, so the children end up
    side by side in Word the way they appear in the PDF (e.g. logo next to
    address text, or a row of signature/date fields).
    """
    flex_elements = soup.find_all(style=_FLEX_RE)

    for el in flex_elements:
        children = [c for c in el.find_all(recursive=False)]
        if not children:
            continue

        table = soup.new_tag("table")
        table["style"] = "width:100%; border-collapse:collapse; border:none;"
        tr = soup.new_tag("tr")
        col_pct = round(100 / len(children), 2)

        for child in children:
            td = soup.new_tag("td")
            td["style"] = "border:none; vertical-align:middle; padding:2px 8px;"
            child.extract()
            td.append(child)
            tr.append(td)

        table.append(tr)
        el.replace_with(table)

_SHORT_HEX_RE = re.compile(r'#([0-9a-fA-F]{3})\b')

def _expand_short_hex_colors(soup: BeautifulSoup) -> None:
    """
    htmldocx's color parser (add_styles_to_run / add_styles_to_run for
    background-color) assumes every '#...' color is 6 hex digits and does
    color[i:i+2] for i in (0, 2, 4) with no length check. A 3-digit
    shorthand hex (e.g. '#fff', '#333') is only 3 chars long, so the last
    slice color[4:6] comes back as '' and int('', 16) raises exactly the
    ValueError seen in the crash. Expand any shorthand hex found in inline
    style="" attributes to full 6-digit form before handing HTML to
    htmldocx, since htmldocx itself won't be patched here.
    """
    for el in soup.find_all(style=True):
        style = el["style"]
        if "#" not in style:
            continue

        def _expand(match):
            h = match.group(1)
            return "#" + "".join(c * 2 for c in h)

        el["style"] = _SHORT_HEX_RE.sub(_expand, style)
from PIL import Image

_PX_RE = re.compile(r'(?:^|;)\s*(width|height)\s*:\s*([\d.]+)px', re.IGNORECASE)
_CSS_DPI    = 96   # standard css-px -> inch assumption (matches WeasyPrint/browsers)
_RENDER_DPI = 150  # oversample so the image still looks crisp in Word


def _resize_data_uri_image(src, target_w_px=None, target_h_px=None):
    try:
        _, b64data = src.split(",", 1)
        img = Image.open(BytesIO(base64.b64decode(b64data)))
        if img.mode in ("P", "LA"):
            img = img.convert("RGBA")
        orig_w, orig_h = img.size

        if target_w_px and not target_h_px:
            target_h_px = target_w_px * orig_h / orig_w
        elif target_h_px and not target_w_px:
            target_w_px = target_h_px * orig_w / orig_h
        elif not target_w_px and not target_h_px:
            return src

        new_w = max(1, round((target_w_px / _CSS_DPI) * _RENDER_DPI))
        new_h = max(1, round((target_h_px / _CSS_DPI) * _RENDER_DPI))

        out = BytesIO()
        img.resize((new_w, new_h), Image.LANCZOS).save(
            out, format="PNG", dpi=(_RENDER_DPI, _RENDER_DPI)
        )
        return "data:image/png;base64," + base64.b64encode(out.getvalue()).decode("ascii")
    except Exception as exc:
        logger.warning("Could not resize embedded image for docx: %s", exc)
        return src

import tempfile
import shutil

def _materialize_data_uri_images(soup: BeautifulSoup, tmp_dir: str) -> None:
    """
    htmldocx's handle_img forwards <img src="..."> straight into
    python-docx's add_picture(), which does open(src, 'rb'). It has no
    data-URI support at all, so a base64 src blows up as a bogus,
    absurdly-long "file path" (OSError: File name too long).

    Decode every data:image src to real bytes, write it to a temp PNG on
    disk, and rewrite src to that path so add_picture()'s open() call
    actually succeeds.
    """
    for i, img in enumerate(soup.find_all("img")):
        src = img.get("src", "")
        if not src.startswith("data:image"):
            continue
        try:
            _, b64data = src.split(",", 1)
            img_bytes = base64.b64decode(b64data)
            file_path = os.path.join(tmp_dir, f"embedded_{i}.png")
            with open(file_path, "wb") as f:
                f.write(img_bytes)
            img["src"] = file_path
        except Exception as exc:
            logger.warning("Could not materialize embedded image %d for docx: %s", i, exc)
            img.decompose()  # drop rather than crash the whole export

def _load_image_bytes(src: str) -> bytes:
    if src.startswith("data:image"):
        _, b64data = src.split(",", 1)
        return base64.b64decode(b64data)
    # local filesystem path, e.g. "templates/assets/optimus_header.jpg" —
    # resolve relative to BASE_DIR the same way WeasyPrint does via base_url,
    # instead of relying on the process's cwd like htmldocx's raw add_picture would.
    path = src if os.path.isabs(src) else os.path.join(BASE_DIR, src)
    with open(path, "rb") as f:
        return f.read()


def _resize_image_src(src: str, target_w_px=None, target_h_px=None) -> str:
    try:
        img_bytes = _load_image_bytes(src)
        img = Image.open(BytesIO(img_bytes))
        if img.mode in ("P", "LA"):
            img = img.convert("RGBA")
        orig_w, orig_h = img.size

        if target_w_px and not target_h_px:
            target_h_px = target_w_px * orig_h / orig_w
        elif target_h_px and not target_w_px:
            target_w_px = target_h_px * orig_w / orig_h
        elif not target_w_px and not target_h_px:
            return src

        new_w = max(1, round((target_w_px / _CSS_DPI) * _RENDER_DPI))
        new_h = max(1, round((target_h_px / _CSS_DPI) * _RENDER_DPI))

        out = BytesIO()
        img.resize((new_w, new_h), Image.LANCZOS).save(
            out, format="PNG", dpi=(_RENDER_DPI, _RENDER_DPI)
        )
        return "data:image/png;base64," + base64.b64encode(out.getvalue()).decode("ascii")
    except Exception as exc:
        logger.warning("Could not resize image %r for docx: %s", src, exc)
        return src


def _fix_image_sizes_for_docx(soup: BeautifulSoup) -> None:
    for img in soup.find_all("img"):
        src = img.get("src", "")
        if not src:
            continue
        style = img.get("style", "")
        dims = {k.lower(): float(v) for k, v in _PX_RE.findall(style)}
        w, h = dims.get("width"), dims.get("height")
        if w is None and h is None:
            continue
        img["src"] = _resize_image_src(src, w, h)

def _prepare_html_for_docx(html_content: str, tmp_dir: str) -> str:
    inlined = inline_css(
        html_content,
        remove_classes=False,
        keep_style_tags=False,
        strip_important=False,
    )

    soup = BeautifulSoup(inlined, "html.parser")

    for el in soup.find_all(style=_FIXED_RE):
        el.decompose()
    for el in soup.find_all(style=_ABS_RE):
        el.decompose()

    _flex_to_table(soup)
    _expand_short_hex_colors(soup)
    _fix_image_sizes_for_docx(soup)      # still resizes, now just adjusts px→data-URI in place
    _materialize_data_uri_images(soup, tmp_dir)   # NEW — must run last, after resizing
    return str(soup)



# ─────────────────────────────────────────────────────────────────────────────
# FORMATTED BUILDER — mirrors the actual PDF template layout, editable
# ─────────────────────────────────────────────────────────────────────────────

def build_formatted_docx(case_data: dict) -> str:
    from services.pdf_generator import render_case_html
    from htmldocx import HtmlToDocx

    html_content = render_case_html(case_data)

    tmp_dir = tempfile.mkdtemp(prefix="docx_img_")
    try:
        html_content = _prepare_html_for_docx(html_content, tmp_dir)

        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)

        parser = HtmlToDocx()
        parser.add_html_to_document(html_content, doc)

        filename_base = _filename_base(case_data)
        output_path = os.path.join(GENERATED_DIR, f"{filename_base}_formatted.docx")
        doc.save(output_path)
        return output_path
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)