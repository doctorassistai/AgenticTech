"""
ingestion.py
============
Document ingestion layer for the Agentic Graph RAG clinical pipeline.

Dispatches PDF / DOCX / TXT / URL inputs to AgenticGraphRAGPipeline
(rag_pipeline.py). All models come from models.py.

Supported inputs
────────────────
  • PDF files   (.pdf)               — pypdf / PyMuPDF
  • DOCX files  (.docx / .doc)       — python-docx
  • Plain text  (.txt / .md / .csv)
  • URLs / links                     — Firecrawl (primary) + httpx fallback

Install
───────
  pip install pypdf python-docx httpx beautifulsoup4 firecrawl-py python-magic
"""

from __future__ import annotations

import io
import os
import re
import asyncio
from pathlib import Path
from typing import List, Optional, Tuple, Any

from loguru import logger

from Agentic.evidence_models import (
    ClinicalKnowledgeGraph,
    DocumentSource,
    GuidelineSource,
    SourceType,
)
from Agentic.evidence_rag_pipeline import AgenticGraphRAGPipeline


# ─────────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────────

FIRECRAWL_API_KEY: str = os.getenv("FIRECRAWL_API_KEY", "")
GROQ_API_KEY:      str = os.getenv("GROQ_API_KEY", "")

# Noise patterns safe to strip before the LLM sees the text
_NOISE_PATTERNS = [
    r"NCCN Guidelines for Patients",
    r"All rights reserved",
    r"Follow us",
    r"Page \d+",
    r"Table of Contents",
    r"Downloaded from",
    r"Printed by",
    r"www\.[^\s]+",
    r"http[s]?://\S+",
]

# Heuristic source detection
_SOURCE_HINTS: List[Tuple[str, GuidelineSource]] = [
    ("nccn",   GuidelineSource.NCCN),
    ("acog",   GuidelineSource.ACOG),
    ("esmo",   GuidelineSource.ESMO),
    ("nejm",   GuidelineSource.NEJM),
    ("lancet", GuidelineSource.LANCET),
    ("asco",   GuidelineSource.ASCO),
]


# ─────────────────────────────────────────────────────────────────
# PIPELINE SINGLETON
# ─────────────────────────────────────────────────────────────────

_pipeline: Optional[AgenticGraphRAGPipeline] = None


def get_pipeline(
    prior_graph: Optional[ClinicalKnowledgeGraph] = None,
) -> AgenticGraphRAGPipeline:
    """Return (or lazily create) the shared pipeline instance."""
    global _pipeline
    if _pipeline is None or prior_graph is not None:
        _pipeline = AgenticGraphRAGPipeline(
            api_key=GROQ_API_KEY,
            prior_graph=prior_graph,
        )
        logger.info("AgenticGraphRAGPipeline initialised.")
    return _pipeline


def reset_pipeline(
    prior_graph: Optional[ClinicalKnowledgeGraph] = None,
) -> AgenticGraphRAGPipeline:
    """Force-create a fresh pipeline instance."""
    global _pipeline
    _pipeline = AgenticGraphRAGPipeline(
        api_key=GROQ_API_KEY,
        prior_graph=prior_graph,
    )
    logger.info("AgenticGraphRAGPipeline reset.")
    return _pipeline


# ─────────────────────────────────────────────────────────────────
# TEXT CLEANING
# ─────────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Minimal noise-strip of raw extracted text."""
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\r",   "\n", text)
    for pattern in _NOISE_PATTERNS:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)
    text = re.sub(r"[ \t]+",  " ",    text)
    text = re.sub(r"\n{3,}",  "\n\n", text)
    text = re.sub(r"[_\-]{4,}", "",   text)
    return text.strip()


# ─────────────────────────────────────────────────────────────────
# GUIDELINE SOURCE DETECTION
# ─────────────────────────────────────────────────────────────────

def _detect_guideline_source(name: str, text_sample: str = "") -> GuidelineSource:
    combined = (name + " " + text_sample[:500]).lower()
    for keyword, source in _SOURCE_HINTS:
        if keyword in combined:
            return source
    return GuidelineSource.OTHER


# ─────────────────────────────────────────────────────────────────
# PDF EXTRACTOR
# ─────────────────────────────────────────────────────────────────

def extract_pdf(
    file_bytes:       bytes,
    source_name:      str,
    guideline_source: Optional[GuidelineSource] = None,
    version:          Optional[str] = None,
    run_pipeline:     bool = True,
) -> ClinicalKnowledgeGraph:
    """
    Extract text from a PDF and run the 8-stage pipeline.

    Args:
        file_bytes:       Raw PDF bytes.
        source_name:      Human-readable document name / filename.
        guideline_source: Override auto-detected GuidelineSource.
        version:          Guideline version string, e.g. "2024.1".
        run_pipeline:     Set False to skip the pipeline (stub graph returned).

    Returns:
        ClinicalKnowledgeGraph
    """
    # Try PyMuPDF first, fall back to pypdf
    pages_text: List[str] = []
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            t = _clean_text(page.get_text())
            if t.strip():
                pages_text.append(t)
    except ImportError:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page_num, page in enumerate(reader.pages):
                try:
                    t = _clean_text(page.extract_text() or "")
                    if t.strip():
                        pages_text.append(t)
                except Exception as exc:
                    logger.warning(f"PDF page {page_num} extraction failed: {exc}")
        except ImportError:
            raise ImportError("Install PyMuPDF or pypdf: pip install PyMuPDF")

    full_text = "\n\n".join(pages_text)
    if not full_text.strip():
        raise ValueError(f"No text extracted from PDF: {source_name}")

    logger.info(f"PDF extracted | {source_name} | {len(pages_text)} pages | {len(full_text):,} chars")

    gs = guideline_source or _detect_guideline_source(source_name, full_text)
    source = DocumentSource(
        source_type=SourceType.PDF,
        guideline_source=gs,
        name=source_name,
        version=version,
    )

    if not run_pipeline:
        return _stub_graph(source_name, gs)

    return get_pipeline().run_from_pdf(file_bytes, source)


# ─────────────────────────────────────────────────────────────────
# DOCX EXTRACTOR
# ─────────────────────────────────────────────────────────────────

def extract_docx(
    file_bytes:       bytes,
    source_name:      str,
    guideline_source: Optional[GuidelineSource] = None,
    version:          Optional[str] = None,
) -> ClinicalKnowledgeGraph:
    """Extract text from a DOCX file and run the pipeline."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("pip install python-docx")

    doc    = Document(io.BytesIO(file_bytes))
    chunks: List[str] = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            chunks.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                chunks.append(row_text)

    full_text = _clean_text("\n\n".join(chunks))
    if not full_text.strip():
        raise ValueError(f"No text extracted from DOCX: {source_name}")

    logger.info(f"DOCX extracted | {source_name} | {len(chunks)} blocks | {len(full_text):,} chars")

    gs = guideline_source or _detect_guideline_source(source_name, full_text)
    source = DocumentSource(
        source_type=SourceType.DOCUMENT,
        guideline_source=gs,
        name=source_name,
        version=version,
    )
    return get_pipeline().run_from_text(full_text, source)


# ─────────────────────────────────────────────────────────────────
# PLAIN TEXT EXTRACTOR
# ─────────────────────────────────────────────────────────────────

def extract_text_file(
    file_bytes:       bytes,
    source_name:      str,
    guideline_source: Optional[GuidelineSource] = None,
    version:          Optional[str] = None,
) -> ClinicalKnowledgeGraph:
    """Extract plain text (.txt / .md / .csv) and run the pipeline."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")

    text = _clean_text(text)
    if not text.strip():
        raise ValueError(f"Empty text file: {source_name}")

    gs = guideline_source or _detect_guideline_source(source_name, text)
    source = DocumentSource(
        source_type=SourceType.TEXT,
        guideline_source=gs,
        name=source_name,
        version=version,
    )
    return get_pipeline().run_from_text(text, source)


# ─────────────────────────────────────────────────────────────────
# URL EXTRACTOR
# ─────────────────────────────────────────────────────────────────

async def extract_url(
    url:              str,
    guideline_source: Optional[GuidelineSource] = None,
    version:          Optional[str] = None,
) -> ClinicalKnowledgeGraph:
    """
    Fetch a URL and run the pipeline.
    Primary: Firecrawl (FIRECRAWL_API_KEY env var).
    Fallback: httpx + BeautifulSoup.
    """
    if FIRECRAWL_API_KEY:
        try:
            return await _extract_url_firecrawl(url, guideline_source, version)
        except Exception as exc:
            logger.warning(f"Firecrawl failed for {url}: {exc} — falling back to httpx")
    return await _extract_url_httpx(url, guideline_source, version)


async def _extract_url_firecrawl(
    url: str,
    guideline_source: Optional[GuidelineSource],
    version: Optional[str],
) -> ClinicalKnowledgeGraph:
    try:
        from firecrawl import FirecrawlApp
    except ImportError:
        raise ImportError("pip install firecrawl-py")

    app  = FirecrawlApp(api_key=FIRECRAWL_API_KEY)
    data = app.scrape_url(url, params={"formats": ["markdown"]})
    if not data or not data.get("markdown"):
        raise ValueError(f"Firecrawl returned no content for {url}")

    text = _clean_text(data["markdown"])
    logger.info(f"Firecrawl extracted | {url} | {len(text):,} chars")

    gs = guideline_source or _detect_guideline_source(url, text)
    source = DocumentSource(
        source_type=SourceType.LINK,
        guideline_source=gs,
        name=url,
        version=version,
    )
    return await get_pipeline().run_from_url(url, source)


async def _extract_url_httpx(
    url: str,
    guideline_source: Optional[GuidelineSource],
    version: Optional[str],
) -> ClinicalKnowledgeGraph:

    try:
        import httpx
        from bs4 import BeautifulSoup

    except ImportError:
        raise ImportError(
            "pip install httpx beautifulsoup4"
        )

    # =========================================================
    # REAL BROWSER HEADERS
    # =========================================================

    headers = {

        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        ),

        "Accept": (
            "text/html,application/xhtml+xml,"
            "application/xml;q=0.9,image/webp,*/*;q=0.8"
        ),

        "Accept-Language":
            "en-US,en;q=0.9",

        "Accept-Encoding":
            "gzip, deflate, br",

        "Referer":
            "https://www.google.com/",

        "DNT":
            "1",

        "Connection":
            "keep-alive",

        "Upgrade-Insecure-Requests":
            "1",
    }

    # =========================================================
    # FETCH URL
    # =========================================================

    async with httpx.AsyncClient(

        follow_redirects=True,

        timeout=60,

        headers=headers,

        verify=False,

    ) as client:

        response = await client.get(url)

        print("URL STATUS =", response.status_code)
        print("FINAL URL =", str(response.url))

        response.raise_for_status()

    # =========================================================
    # PARSE HTML
    # =========================================================

    soup = BeautifulSoup(
        response.text,
        "html.parser"
    )

    # remove noisy tags
    for tag in soup([
        "script",
        "style",
        "nav",
        "footer",
        "header",
        "aside",
        "form",
    ]):
        tag.decompose()

    text = _clean_text(
        soup.get_text(separator="\n")
    )

    if not text.strip():

        raise ValueError(
            f"No text extracted from URL: {url}"
        )

    logger.info(
        f"httpx extracted | {url} | {len(text):,} chars"
    )

    # =========================================================
    # BUILD SOURCE
    # =========================================================

    gs = guideline_source or _detect_guideline_source(
        url,
        text
    )

    source = DocumentSource(

        source_type=SourceType.LINK,

        guideline_source=gs,

        name=url,

        version=version,
    )

    # =========================================================
    # RUN PIPELINE
    # =========================================================

    return await get_pipeline().run_from_text(
        text,
        source
    )


# ─────────────────────────────────────────────────────────────────
# FILE DISPATCHER
# ─────────────────────────────────────────────────────────────────

def extract_file(
    file_bytes:       bytes,
    filename:         str,
    guideline_source: Optional[GuidelineSource] = None,
    version:          Optional[str] = None,
) -> ClinicalKnowledgeGraph:
    """
    Route to the correct extractor by file extension.
    Called by the FastAPI upload endpoint.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(file_bytes, filename,
                           guideline_source=guideline_source, version=version)
    if ext in (".docx", ".doc"):
        return extract_docx(file_bytes, filename,
                            guideline_source=guideline_source, version=version)
    if ext in (".txt", ".md", ".csv"):
        return extract_text_file(file_bytes, filename,
                                 guideline_source=guideline_source, version=version)

    # MIME sniff for unknown extensions
    try:
        import magic
        mime = magic.from_buffer(file_bytes[:2048], mime=True)
        if "pdf" in mime:
            return extract_pdf(file_bytes, filename,
                               guideline_source=guideline_source, version=version)
        if "wordprocessingml" in mime or "msword" in mime:
            return extract_docx(file_bytes, filename,
                                guideline_source=guideline_source, version=version)
        if "text" in mime:
            return extract_text_file(file_bytes, filename,
                                     guideline_source=guideline_source, version=version)
    except ImportError:
        pass

    raise ValueError(
        f"Unsupported file type: '{ext}'. "
        "Supported: .pdf  .docx  .doc  .txt  .md  .csv"
    )


# ─────────────────────────────────────────────────────────────────
# BATCH INGESTION
# ─────────────────────────────────────────────────────────────────

async def ingest_urls(
    urls:             List[str],
    guideline_source: Optional[GuidelineSource] = None,
    version:          Optional[str] = None,
) -> List[ClinicalKnowledgeGraph]:
    """
    Extract and process a list of URLs concurrently.
    Failed URLs are logged and skipped.
    """
    tasks   = [extract_url(url, guideline_source, version) for url in urls]
    results = await asyncio.gather(*tasks, return_exceptions=True)
    graphs: List[ClinicalKnowledgeGraph] = []
    for url, result in zip(urls, results):
        if isinstance(result, Exception):
            logger.error(f"URL ingestion failed for {url}: {result}")
        else:
            graphs.append(result)
    return graphs


def ingest_files(
    file_pairs:       List[Tuple[bytes, str]],
    guideline_source: Optional[GuidelineSource] = None,
    version:          Optional[str] = None,
) -> ClinicalKnowledgeGraph:
    """
    Ingest multiple files into a SINGLE unified ClinicalKnowledgeGraph.
    Uses AgenticGraphRAGPipeline.run_batch so all documents share one graph.
    """
    pipeline = get_pipeline()
    batch: List[Tuple[Any, DocumentSource]] = []

    for file_bytes, filename in file_pairs:
        ext = Path(filename).suffix.lower()
        try:
            if ext == ".pdf":
                try:
                    import fitz
                    doc = fitz.open(stream=file_bytes, filetype="pdf")
                    raw_text = _clean_text("\n\n".join(p.get_text() for p in doc))
                except ImportError:
                    import pypdf
                    reader   = pypdf.PdfReader(io.BytesIO(file_bytes))
                    raw_text = _clean_text(
                        "\n\n".join((p.extract_text() or "") for p in reader.pages)
                    )
                src_type = SourceType.PDF

            elif ext in (".docx", ".doc"):
                from docx import Document as DocxDoc
                d        = DocxDoc(io.BytesIO(file_bytes))
                raw_text = _clean_text(
                    "\n\n".join(p.text.strip() for p in d.paragraphs if p.text.strip())
                )
                src_type = SourceType.DOCUMENT

            elif ext in (".txt", ".md", ".csv"):
                raw_text = _clean_text(file_bytes.decode("utf-8", errors="replace"))
                src_type = SourceType.TEXT

            else:
                logger.warning(f"Skipping unsupported file: {filename}")
                continue

        except Exception as exc:
            logger.error(f"Extraction failed for {filename}: {exc}")
            continue

        if not raw_text.strip():
            logger.warning(f"Empty content — skipping {filename}")
            continue

        gs = guideline_source or _detect_guideline_source(filename, raw_text)
        batch.append((raw_text, DocumentSource(
            source_type=src_type,
            guideline_source=gs,
            name=filename,
            version=version,
        )))

    if not batch:
        raise ValueError("No valid files could be extracted.")

    logger.info(f"Running batch pipeline on {len(batch)} documents …")
    return pipeline.run_batch(batch)


# ─────────────────────────────────────────────────────────────────
# INTERNAL HELPERS
# ─────────────────────────────────────────────────────────────────

# if __name__ == "__main__":
#     import sys
#     from loguru import logger as log

#     log.remove()
#     log.add(
#         sys.stderr,
#         level="INFO",
#         format="<green>{time:HH:mm:ss}</green> | "
#                "<level>{level:<8}</level> | "
#                "<cyan>{function}</cyan> — {message}"
#     )

#     sample = """
#     NCCN Clinical Practice Guidelines — Non-Small Cell Lung Cancer v2024.1

#     RECOMMENDATION (Category 1):
#     For metastatic NSCLC with PD-L1 TPS >= 50% without actionable genomic
#     alterations, pembrolizumab monotherapy is recommended as preferred first-line
#     therapy (KEYNOTE-024, HR 0.50, 95% CI 0.37-0.68, p<0.001).
#     """

#     graph = extract_text_file(
#         file_bytes=sample.encode("utf-8"),
#         source_name="NCCN_NSCLC_2024.txt",
#         guideline_source=GuidelineSource.NCCN,
#         version="2024.1",
#     )