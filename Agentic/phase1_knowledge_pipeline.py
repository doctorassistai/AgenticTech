"""
phase1_knowledge_pipeline_v14.py
=================================
Phase 1 – Clinical Knowledge Creation Pipeline  (v14 — Retrieval-First Agentic Architecture
                                                        + Structural Capability Extraction)
==============================================================================================

ARCHITECTURE CHANGE (v13 → v14):

    Everything from v13 is UNCHANGED (PDF parsing, sections, ChromaDB, Discovery Agent,
    Disease Registry, Query Builder, Retrieval, Coverage Validator, Disease/Subtype
    Knowledge Generation, Relationship Extraction, Knowledge Graph, Neo4j, MongoDB).

    The ONLY change is in skill aggregation (collect_skills / Section 16b-16c):

    v13 problem:
        Skill count == however many "skills" the LLM happened to enumerate inside its
        diagnosis/treatment JSON response. Knowledge richness (diseases, subtypes,
        biomarkers, stages, regimens, if/then rules, etc.) was NOT reliably reflected
        in skill count, because skill generation depended on the LLM remembering to
        list them, rather than on what was actually extracted.

    v14 fix — Structural Capability Extraction (NEW, Section 16c):
        Skills are now derived DIRECTLY from the SHAPE of the already-extracted
        diagnosis/treatment knowledge (no new LLM calls, no keyword lists, no
        per-disease/per-specialty branching). If diagnosis["biomarkers"] has content,
        a "Biomarker Interpretation" skill exists. If treatment["chemotherapy"]["regimens"]
        has content, a "Chemotherapy Regimen Selection" skill exists. Etc.

        This guarantees skill count scales with knowledge richness, generalizes to
        ANY specialty (the rules key off field shape, never off disease names), and
        costs ZERO additional LLM calls or retrieval calls.

        collect_skills() now merges TWO sources before the existing semantic dedup:
          1) LLM-declared skills (unchanged from v13 — kept, not removed)
          2) Structurally-derived skills (NEW)

pipeline_version = "v14"
"""

from __future__ import annotations

# ── Standard library ──────────────────────────────────────────────
import asyncio
import copy
import hashlib
import io
import json
import os
import re
import subprocess
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional
import time
# ── Third-party ───────────────────────────────────────────────────
import requests as _requests   # OpenRouter embedding calls
import chromadb
from groq import Groq
from loguru import logger
from motor.motor_asyncio import AsyncIOMotorClient

# from Agentic.skill_markdown import skill_body_to_markdown
try:
    from neo4j import GraphDatabase
    _NEO4J_AVAILABLE = True
except ImportError:
    _NEO4J_AVAILABLE = False


# ═════════════════════════════════════════════════════════════════
# SECTION 1 — CONFIGURATION
# ═════════════════════════════════════════════════════════════════

GROQ_API_KEY     = os.getenv("GROQ_API_KEY",  "")
MONGO_URI        = os.getenv("MONGO_URI",     "")
MONGO_DB         = "doctorassistai"

# ── OpenRouter embedding ───────────────────────────────────────────
OPENROUTER_API_KEY   = os.getenv("OPENAI_API_ROUTER_KEY", "")
OPENROUTER_EMBED_URL = "https://openrouter.ai/api/v1/embeddings"
EMBEDDING_MODEL      = "openai/text-embedding-3-large"
EMBEDDING_DIMENSION  = 3072

# ── ChromaDB ────────────────────────────────────────────────────────
CHROMA_PERSIST_PATH   = os.getenv("CHROMA_PERSIST_PATH", "./knowledge_db")
CHROMA_COLLECTION_NAME = "medical_guidelines"

# ── Neo4j ───────────────────────────────────────────────────────────
NEO4J_URI      = os.getenv("NEO4J_URI",      "")
NEO4J_USER     = os.getenv("NEO4J_USER",     "neo4j")
NEO4J_PASSWORD = os.getenv("NEO4J_PASSWORD", "")

_groq_client = Groq(api_key=GROQ_API_KEY)
MODEL        = "llama-3.3-70b-versatile"
TEMPERATURE  = 0.1
MAX_TOKENS   = 10000

# ── Section / chunking limits ─────────────────────────────────────
MAX_SECTION_CHARS         = 12_000
MIN_SECTION_CHARS         = 1_500
MAX_MERGE_RESULT_CHARS    = 20_000
TOC_MIN_AVG_SECTION_CHARS = 2_000
MAX_CONCURRENT_GROQ_CALLS = 1

# ── Retrieval (ChromaDB, no-LLM) ───────────────────────────────────
MAX_AGENT_CHARS          = 20_000
SIMILARITY_THRESHOLD  = 0.30   # L2-based cosine approx needs lower threshold
MAX_RETRIEVE_LIMIT    = 100

# ── Subtype batching ────────────────────────────────────────────────
SUBTYPE_BATCH_SIZE           = 5
SUBTYPE_MIN_SPECIFICITY      = 0.30
SUBTYPE_INHERITED_RATIO_WARN = 0.80

# ── Per-disease skill cap ───────────────────────────────────────────
MAX_PER_DISEASE_SKILLS = 20

# ── Skill dedup / quality (NEW) ──────────────────────────────────────
SKILL_DEDUP_THRESHOLD   = 0.96   # stricter than the old 0.92 — compares skill+basis text, not just name
SKILL_MIN_QUALITY_SCORE = 0      # 0 = score+sort only, don't drop anything yet. Raise once you've seen real score distribution.

_LLM_CACHE: dict[str, dict] = {}


# ═════════════════════════════════════════════════════════════════
# SECTION 2 — CLINICAL RELATIONSHIP STORE  (kept — feeds Knowledge Graph)
# ═════════════════════════════════════════════════════════════════

@dataclass
class ClinicalRelationship:
    source:      str
    relation:    str
    target:      str
    source_page: int
    confidence:  float = 1.0
    section:     str   = ""
    disease:     str   = ""   # NEW (Phase 2a) — which disease this triple belongs to, if identifiable
    subtype:     str   = ""   # NEW (Phase 2a) — which subtype, if identifiable

    def to_dict(self) -> dict:
        return {
            "source":      self.source,
            "relation":    self.relation,
            "target":      self.target,
            "source_page": self.source_page,
            "confidence":  self.confidence,
            "section":     self.section,
            "disease":     self.disease,
            "subtype":     self.subtype,
        }


class RelationshipStore:
    def __init__(self):
        self._rels: list[ClinicalRelationship] = []
        self._seen: set[str] = set()

    def add(self, rel: ClinicalRelationship) -> bool:
        key = f"{rel.source.lower()}|{rel.relation}|{rel.target.lower()}|{rel.subtype.lower()}"
        if key in self._seen:
            return False
        self._seen.add(key)
        self._rels.append(rel)
        return True

    def all_relationships(self) -> list[ClinicalRelationship]:
        return list(self._rels)

    def __len__(self) -> int:
        return len(self._rels)


_RELATIONSHIP_EXTRACT_SYSTEM = """
You are a clinical relationship extraction specialist.
Extract explicit clinical relationships from these sections as typed triples.

If the surrounding text explicitly names a specific disease SUBTYPE (not just
the parent disease) that this relationship applies to, capture it in
"subtype". If the text explicitly names the parent disease, capture it in
"disease". If either is not explicitly stated or is ambiguous, leave it as an
empty string — do NOT guess or infer from general medical knowledge.

Return ONLY valid JSON:
{
  "relationships": [
    {
      "source":      "the entity (biomarker, drug, test, procedure, condition)",
      "relation":    "evaluates|selects|indicates|contraindicated_with|monitors|predicts|treats|stages|differentiates",
      "target":      "the target entity",
      "source_page": 0,
      "confidence":  1.0,
      "disease":     "",
      "subtype":     ""
    }
  ]
}

Extract ONLY relationships EXPLICITLY stated in the text. No prior knowledge.
Return [] if none found. Do not assume any specific medical specialty.
"""

_RELATION_PHASE_EMBED_CACHE: dict[str, list[float]] = {}
_RELATION_PHASE_RESULT_CACHE: dict[str, str] = {}   # NEW — caches the final answer per relation string

_RELATION_PHASE_ANCHORS = [
    "diagnostic or evaluative relationship",
    "treatment or management relationship",
]

def _relation_phase(relation: str) -> str:
    """
    Returns 'diagnosis' or 'treatment' for any relation type via embedding
    similarity against two generic phase anchors — no fixed vocabulary
    split. Result is memoized per relation string (CRITICAL — without this,
    a nested loop of skills × relationships fires one embedding API call
    PER PAIR, which can reach tens of thousands of calls on a large
    document and effectively hang the pipeline).
    """
    key = relation.lower().strip()
    if key in _RELATION_PHASE_RESULT_CACHE:
        return _RELATION_PHASE_RESULT_CACHE[key]

    global _RELATION_PHASE_EMBED_CACHE
    missing = [a for a in _RELATION_PHASE_ANCHORS if a not in _RELATION_PHASE_EMBED_CACHE]
    if missing:
        vecs = _embed_batch(missing)
        for a, v in zip(missing, vecs):
            if v is not None:
                _RELATION_PHASE_EMBED_CACHE[a] = v

    if not all(a in _RELATION_PHASE_EMBED_CACHE for a in _RELATION_PHASE_ANCHORS):
        result = "treatment"
    else:
        rel_vec = _embed_text(relation)
        if rel_vec is None:
            result = "treatment"
        else:
            diag_vec, treat_vec = (_RELATION_PHASE_EMBED_CACHE[a] for a in _RELATION_PHASE_ANCHORS)
            diag_sim  = sum(a * b for a, b in zip(rel_vec, diag_vec))
            treat_sim = sum(a * b for a, b in zip(rel_vec, treat_vec))
            result = "diagnosis" if diag_sim >= treat_sim else "treatment"

    _RELATION_PHASE_RESULT_CACHE[key] = result
    return result

async def _extract_relationships(
    sections: list["Section"],
    semaphore: asyncio.Semaphore,
) -> list[ClinicalRelationship]:
    if not sections:
        return []

    batches = _batch_sections(sections)
    all_rels: list[ClinicalRelationship] = []

    for batch in batches:
        parts = []
        for sec in batch:
            parts.append(
                f"=== {sec.section_title} (pages {sec.page_start}–{sec.page_end}) ===\n"
                f"{sec.text[:3000]}"
            )
        user_msg = "\n\n".join(parts)

        try:
            async with semaphore:
                result = await _chat_async(_RELATIONSHIP_EXTRACT_SYSTEM, user_msg)
        except Exception as e:
            logger.warning(f"[RelationshipExtract] batch failed: {e}")
            continue

        for rel_dict in (result.get("relationships") or []):
            if not isinstance(rel_dict, dict):
                continue
            src = rel_dict.get("source", "").strip()
            rel = rel_dict.get("relation", "").strip()
            tgt = rel_dict.get("target", "").strip()
            if src and rel and tgt:
                all_rels.append(ClinicalRelationship(
                    source      = src,
                    relation    = rel,
                    target      = tgt,
                    source_page = rel_dict.get("source_page", batch[0].page_start or 1),
                    confidence  = rel_dict.get("confidence", 1.0),
                    section     = batch[0].section_title if batch else "",
                    disease     = (rel_dict.get("disease") or "").strip(),   # NEW
                    subtype     = (rel_dict.get("subtype") or "").strip(),  # NEW
                ))

    logger.info(f"[RelationshipExtract] {len(all_rels)} clinical relationships extracted")
    return all_rels


# ═════════════════════════════════════════════════════════════════
# SECTION 3 — SECTION DATACLASS
# ═════════════════════════════════════════════════════════════════

@dataclass
class Section:
    text:          str
    section_id:    str
    section_title: str
    page_start:    int
    page_end:      int
    section_index: int
    heading_level: int = 1


# ═════════════════════════════════════════════════════════════════
# SECTION 4 — LLM CALL HELPERS
# ═════════════════════════════════════════════════════════════════

def _cache_key(system: str, user: str) -> str:
    return hashlib.sha256((system + "\x00" + user).encode()).hexdigest()


def _trim_to_budget(text: str, max_chars: int = MAX_AGENT_CHARS * 2) -> str:
    HARD_CEILING = MAX_AGENT_CHARS * 4
    if len(text) <= max_chars:
        return text
    if len(text) <= HARD_CEILING:
        return text
    boundary = text.rfind("\n\n", 0, HARD_CEILING)
    if boundary > HARD_CEILING // 2:
        return text[:boundary]
    return text[:HARD_CEILING]


def _chat(system: str, user: str, use_cache: bool = True) -> dict:
    user = _trim_to_budget(user)
    if use_cache:
        key = _cache_key(system, user)
        if key in _LLM_CACHE:
            return _LLM_CACHE[key]

    def _attempt(extra_tokens: int = 0) -> dict:
        resp = _groq_client.chat.completions.create(
            model=MODEL,
            temperature=TEMPERATURE,
            max_tokens=MAX_TOKENS + extra_tokens,
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": user},
            ],
            response_format={"type": "json_object"},
        )
        raw           = resp.choices[0].message.content or "{}"
        logger.info(
            f"[LLM RAW RESPONSE]\n{raw[:5000]}"
        )
        finish_reason = resp.choices[0].finish_reason
        if finish_reason == "length":
            raise ValueError("length_limit")
        try:
            parsed = json.loads(raw)

            logger.info(
                f"[LLM PARSED KEYS] {list(parsed.keys())}"
            )

            return parsed
        except json.JSONDecodeError:
            clean = re.sub(r"```(?:json)?|```", "", raw).strip()
            try:
                return json.loads(clean)
            except json.JSONDecodeError:
                return json.loads(_repair_truncated_json(clean))

    def _attempt_with_retry(extra_tokens: int = 0) -> dict:
        for attempt in range(8):
            try:
                return _attempt(extra_tokens)
            except Exception as e:
                if "429" in str(e) or "rate_limit" in str(e).lower():
                    wait = min(60, 2 ** attempt)
                    logger.warning(f"[Groq] 429 rate limit — sleeping {wait}s (attempt {attempt+1}/8)")
                    time.sleep(wait)
                    continue
                raise
        raise RuntimeError("Groq rate limit: max retries exhausted")

    try:
        result = _attempt_with_retry()
    except ValueError as e:
        if "length_limit" in str(e):
            try:
                result = _attempt_with_retry(extra_tokens=MAX_TOKENS // 4)
            except (ValueError, json.JSONDecodeError) as e2:
                logger.error(f"LLM retry also failed: {e2}")
                result = {}
        else:
            raise

    if use_cache and result:
        _LLM_CACHE[_cache_key(system, user)] = result
    return result


def _repair_truncated_json(text: str) -> str:
    depth_curly  = text.count("{") - text.count("}")
    depth_square = text.count("[") - text.count("]")
    last_safe = max(text.rfind("},"), text.rfind("},\n"), text.rfind("]"), text.rfind('"'))
    if last_safe > len(text) // 2:
        text = text[:last_safe + 1]
        depth_curly  = text.count("{") - text.count("}")
        depth_square = text.count("[") - text.count("]")
    if depth_square > 0:
        text += "]" * depth_square
    if depth_curly > 0:
        text += "}" * depth_curly
    return text


async def _chat_async(system: str, user: str, use_cache: bool = True) -> dict:
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: _chat(system, user, use_cache))


def _attach_page_to_items(items: list, page_range: list[int]) -> list:
    if not items or not page_range:
        return items
    result = []
    for item in items:
        if isinstance(item, dict):
            existing_page = item.get("source_page")
            if not existing_page or existing_page == 0:
                item = {**item, "source_page": page_range[0]}
            if "source_pages" not in item:
                item = {**item, "source_pages": page_range}
        result.append(item)
    return result


# ═════════════════════════════════════════════════════════════════
# SECTION 5 — OPENROUTER EMBEDDING HELPERS
# ═════════════════════════════════════════════════════════════════

def _embed_batch(texts: list[str], retry_count: int = 0) -> list[Optional[list[float]]]:
    """
    Embed a batch of texts via OpenRouter with exponential backoff on 429.
    Returns list of embeddings (None for failed items).
    """
    if not OPENROUTER_API_KEY or not texts:
        return [None] * len(texts)
    try:
        response = _requests.post(
            OPENROUTER_EMBED_URL,
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type":  "application/json",
            },
            json={"model": EMBEDDING_MODEL, "input": [t[:8000] for t in texts]},
            timeout=60,
        )
        if response.status_code == 429:
            wait = 2 ** retry_count   # 1, 2, 4, 8 seconds
            logger.warning(f"[Embedding] 429 rate limit — waiting {wait}s (retry {retry_count+1})")
            time.sleep(wait)
            if retry_count < 4:
                return _embed_batch(texts, retry_count + 1)
            return [None] * len(texts)
        response.raise_for_status()
        data = response.json()["data"]
        result = [None] * len(texts)
        for item in data:
            result[item["index"]] = item["embedding"]
        return result
    except Exception as exc:
        logger.warning(f"[Embedding] batch call failed: {exc}")
        return [None] * len(texts)


def _embed_text(text: str) -> Optional[list[float]]:
    """Single-text embed — used for query embeddings (search_guideline)."""
    if not text or not text.strip():
        return None
    results = _embed_batch([text])
    return results[0]


# ═════════════════════════════════════════════════════════════════
# SECTION 6 — CHROMADB LAYER
# ═════════════════════════════════════════════════════════════════

_chroma_client = None
_chroma_collection = None

def _get_chroma_collection():
    global _chroma_client, _chroma_collection
    if _chroma_collection is None:
        _chroma_client = chromadb.PersistentClient(path=CHROMA_PERSIST_PATH)
        _chroma_collection = _chroma_client.get_or_create_collection(CHROMA_COLLECTION_NAME)
    return _chroma_collection

EMBED_BATCH_SIZE = 20


def _chroma_doc_id(doc_id: str, section_id: str, section_index: int) -> str:
    return f"{doc_id}::{section_id}::{section_index}"


def _section_content_hash(sec: Section) -> str:
    return hashlib.sha256(sec.text.encode()).hexdigest()


def store_section_embeddings(
    sections: list[Section],
    doc_id: str,
    doctor_id: str,
    mongo_uri: str = "",
) -> int:
    """
    Batch-embed sections (20 at a time) and store in ChromaDB.
    Tracks progress in MongoDB embedding_jobs collection for resume support.
    """
    last_completed = 0
    motor_sync_client = None
    emb_jobs_coll = None

    if mongo_uri:
        try:
            from pymongo import MongoClient
            motor_sync_client = MongoClient(mongo_uri)
            emb_jobs_coll = motor_sync_client[MONGO_DB]["embedding_jobs"]
            job = emb_jobs_coll.find_one({"doc_id": doc_id})
            if job and job.get("status") == "completed":
                logger.info(f"[Chroma] Embedding already completed for doc_id={doc_id} — skipping")
                if motor_sync_client is not None: motor_sync_client.close()
                return job.get("sections_stored", 0)
            if job and job.get("last_embedded_section", 0) > 0:
                last_completed = job["last_embedded_section"]
                logger.info(f"[Chroma] Resuming from section {last_completed}")
            emb_jobs_coll.update_one(
                {"doc_id": doc_id},
                {"$set": {
                    "doc_id": doc_id, "doctor_id": doctor_id,
                    "last_embedded_section": last_completed,
                    "status": "in_progress",
                    "total_sections": len(sections),
                }},
                upsert=True,
            )
        except Exception as e:
            logger.warning(f"[EmbeddingJobs] MongoDB tracking unavailable: {e}")

    stored_count = 0
    collection = _get_chroma_collection()
    try:
        existing = collection.get(where={"doc_id": doc_id}, include=["metadatas"])
        existing_hashes = {
            m.get("content_hash", "")
            for m in (existing.get("metadatas") or [])
            if m.get("content_hash")
        }
    except Exception:
        existing_hashes = set()

    sections_to_embed = [
        s for s in sections[last_completed:]
        if _section_content_hash(s) not in existing_hashes
    ]
    skipped = len(sections[last_completed:]) - len(sections_to_embed)
    if skipped:
        logger.info(f"[Chroma] {skipped} sections skipped (already embedded by hash)")

    for batch_start in range(0, len(sections_to_embed), EMBED_BATCH_SIZE):
        batch = sections_to_embed[batch_start: batch_start + EMBED_BATCH_SIZE]
        texts = [f"{sec.section_title}. {sec.text[:6000]}" for sec in batch]
        vectors = _embed_batch(texts)

        ids, documents, embeddings, metadatas = [], [], [], []
        for sec, vector in zip(batch, vectors):
            if vector is None:
                logger.warning(f"[Chroma] Embedding failed for '{sec.section_title}' — skipping")
                continue
            ids.append(_chroma_doc_id(doc_id, sec.section_id, sec.section_index))
            documents.append(sec.text)
            embeddings.append(vector)
            metadatas.append({
                "doc_id":        doc_id,
                "doctor_id":     doctor_id,
                "section_id":    sec.section_id,
                "title":         sec.section_title,
                "page_start":    sec.page_start,
                "page_end":      sec.page_end,
                "heading_level": sec.heading_level,
                "section_index": sec.section_index,
                "content_hash":  _section_content_hash(sec),
            })

        from collections import Counter
        if ids:
            duplicates = [k for k, v in Counter(ids).items() if v > 1]
            if duplicates:
                logger.error(f"[Chroma] Duplicate IDs detected: {duplicates}")
                for sec in batch:
                    logger.error(
                        f"section_id={sec.section_id} | "
                        f"section_index={sec.section_index} | title={sec.section_title}"
                    )
            collection.add(ids=ids, documents=documents, embeddings=embeddings, metadatas=metadatas)
            stored_count += len(ids)

        completed_idx = last_completed + batch_start + len(batch)
        if emb_jobs_coll is not None:
            try:
                emb_jobs_coll.update_one(
                    {"doc_id": doc_id},
                    {"$set": {"last_embedded_section": completed_idx, "sections_stored": stored_count}}
                )
            except Exception:
                pass

        logger.info(
            f"[Chroma] Batch {batch_start//EMBED_BATCH_SIZE + 1} | "
            f"stored {stored_count}/{len(sections)} sections"
        )

    if emb_jobs_coll is not None:
        try:
            emb_jobs_coll.update_one(
                {"doc_id": doc_id},
                {"$set": {"status": "completed", "sections_stored": stored_count}}
            )
        except Exception:
            pass
    if motor_sync_client is not None:
        motor_sync_client.close()

    logger.info(f"[Chroma] Embedding complete | {stored_count}/{len(sections)} stored | doc_id={doc_id}")
    return stored_count


def _chroma_hit_to_section(doc_text: str, metadata: dict) -> Section:
    return Section(
        text          = doc_text,
        section_id    = metadata.get("section_id", ""),
        section_title = metadata.get("title", "?"),
        page_start    = metadata.get("page_start", 1),
        page_end      = metadata.get("page_end", metadata.get("page_start", 1)),
        section_index = metadata.get("section_index", 0),
        heading_level = metadata.get("heading_level", 1),
    )


def search_guideline(query: str, doc_id: str, n_results: int = MAX_RETRIEVE_LIMIT) -> list[Section]:
    """
    Semantic retrieval with similarity filtering — pure ChromaDB, NO LLM.
    Retrieves broadly then filters by threshold — no fixed top-K cutoff.
    """
    query_vector = _embed_text(query)
    if query_vector is None:
        return []
    collection = _get_chroma_collection()
    try:
        results = collection.query(
            query_embeddings=[query_vector],
            n_results=n_results,
            where={"doc_id": doc_id},
            include=["documents", "metadatas", "distances"],
        )
    except Exception as exc:
        logger.warning(f"[Chroma] query failed: {exc}")
        return []

    sections: list[Section] = []
    docs      = (results.get("documents")  or [[]])[0]
    metas     = (results.get("metadatas")  or [[]])[0]
    distances = (results.get("distances")  or [[]])[0]

    for doc_text, meta, distance in zip(docs, metas, distances):
        similarity = max(0.0, 1.0 - (distance / 2.0))
        if similarity >= SIMILARITY_THRESHOLD:
            sections.append(_chroma_hit_to_section(doc_text, meta))

    logger.debug(
        f"[Chroma] query='{query[:60]}' | retrieved={len(docs)} | "
        f"above_threshold={len(sections)} | threshold={SIMILARITY_THRESHOLD}"
    )
    return sections


# ═══════════════════════════════════════════════════════════════
# SECTION 6b — DISCOVERY AGENT + DISEASE/SUBTYPE REGISTRIES
# ═══════════════════════════════════════════════════════════════

_DISCOVERY_EXTRACT_SYSTEM = """
You are a clinical knowledge discovery specialist.
Read these retrieved guideline sections and identify ALL entities covered.
Do NOT use any prior medical knowledge — only what is EXPLICITLY stated in the text.

Classify every entity into exactly ONE bucket using these structural definitions
(these are definitions of TYPE, not examples — apply them to whatever condition
this guideline is actually about, in ANY specialty):

- disease: a named clinical condition/disorder that is the subject of management
  in this guideline (the thing being diagnosed/treated).
- subtype: a clinically distinct variant/category OF a disease, defined by
  intrinsic biological, anatomical, or pathological characteristics
  (not by severity, time-course, or response status).
- stage: a label describing anatomical extent/severity/progression of a disease
  at a point in time (any staging or grading system explicitly used in the text,
  in any classification system the guideline uses).
- risk_group: a label describing a patient's PROGNOSTIC or treatment-decision
  risk category (e.g. derived from a scoring system, risk calculator, or
  explicit "risk" terminology in the text) — distinct from anatomical stage.
- biomarker: any measurable molecular, genetic, serological, or physiological
  marker used for diagnosis, prognosis, or treatment selection.

Disambiguation rule: if an entity describes EXTENT/SPREAD of disease at a point
in time -> stage. If it describes a PROGNOSTIC/RISK classification derived from
multiple factors -> risk_group. If it describes a BIOLOGICAL VARIANT of the
disease itself -> subtype. When genuinely ambiguous from the text, prefer the
bucket the source document itself names it as (use the heading/label context).

For every subtype, you MUST also identify which disease in the "diseases" list
it belongs to, based on where it is discussed in the text. If a subtype is
discussed under more than one disease, list all applicable parents.

Return ONLY valid JSON:
{
  "diseases": ["<disease name>"],
  "subtypes": [
    {"name": "<subtype name>", "parent_diseases": ["<disease name>"]}
  ],
  "stages": ["<stage label>"],
  "risk_groups": ["<risk group label>"],
  "biomarkers": ["<biomarker name>"],
  "drugs": ["<drug name>"],
  "procedures": ["<procedure name>"],
  "special_populations": ["<population>"]
}
"""

def _semantic_dedup(items: list[dict], threshold: float = 0.92, text_fn=None) -> list[dict]:
    """
    Embedding-based dedup. text_fn, if given, builds the comparison string per
    item (e.g. skill + basis) instead of just item["name"] — lets callers dedup
    on richer context without duplicating this function.
    """
    if len(items) <= 1:
        return items
    if text_fn is None:
        texts = [it.get("name", "") for it in items]
    else:
        texts = [text_fn(it) for it in items]
    vectors = _embed_batch(texts)
    kept: list[dict] = []
    kept_vectors: list = []
    for item, vec in zip(items, vectors):
        if vec is None:
            kept.append(item)
            kept_vectors.append(None)
            continue
        is_dup = False
        for kv in kept_vectors:
            if kv is None:
                continue
            sim = sum(a * b for a, b in zip(vec, kv))
            if sim >= threshold:
                is_dup = True
                break
        if not is_dup:
            kept.append(item)
            kept_vectors.append(vec)
    return kept

async def run_discovery_agent(
    doc_id: str,
    semaphore: asyncio.Semaphore,
    all_sections: list[Section],
) -> dict:
    """
    Run the discovery agent across ALL embedded sections to auto-identify
    diseases, subtypes, biomarkers, stages, and special populations.
    Reads every section — not just top-K semantic matches — so nothing is missed
    in 17–300+ page documents. This is the ONLY full-document-read LLM stage;
    everything downstream is targeted retrieval.
    """
    if not all_sections:
        logger.warning("[DiscoveryAgent] No sections available")
        return {
            "diseases": [], "subtypes": [], "biomarkers": [],
            "drugs": [], "procedures": [], "stages": [], "special_populations": [],
        }

    logger.info(f"[DiscoveryAgent] Processing ALL {len(all_sections)} sections for discovery")

    logger.info(
        f"[Discovery] Total Sections={len(all_sections)}"
    )

    for s in all_sections[:10]:
        logger.info(
            f"[Section] {s.section_title}"
        )
    batches  = _batch_sections(all_sections)
    results: list[dict] = []

    for batch_idx, batch in enumerate(batches):
        user_msg = _build_agent_user_message("discovery", batch)

        # DEBUG 1: What are we sending?
        logger.info(
            f"[Discovery Prompt Batch {batch_idx}] "
            f"chars={len(user_msg)}\n"
            f"{user_msg[:3000]}"
        )

        logger.info(
            f"[Discovery Batch {batch_idx}] "
            f"sections={len(batch)}"
        )

        try:
            async with semaphore:
                result = await _chat_async(
                    _DISCOVERY_EXTRACT_SYSTEM,
                    user_msg
                )

            # DEBUG 2: What did Groq return?
            logger.info(
                f"[Discovery Raw Batch {batch_idx}]\n"
                f"{json.dumps(result, indent=2)[:5000]}"
            )

            # DEBUG 3: What keys exist?
            if isinstance(result, dict):
                logger.info(
                    f"[Discovery Keys Batch {batch_idx}] "
                    f"{list(result.keys())}"
                )
                results.append(result)
            else:
                logger.error(
                    f"[Discovery] Non-dict response: {type(result)}"
                )

        except Exception as e:
            logger.warning(
                f"[DiscoveryAgent] batch {batch_idx} failed: {e}"
            )

    logger.info(
        f"[Discovery] Total batch results received: {len(results)}"
    )

    for idx, r in enumerate(results):
        logger.info(
            f"[Discovery Result {idx}] "
            f"keys={list(r.keys()) if isinstance(r, dict) else 'NOT_DICT'}"
        )

    merged: dict = {
        "diseases": [], "subtypes": [], "biomarkers": [],
        "drugs": [], "procedures": [], "stages": [], "special_populations": [],
    }
    for r in results:
        for key in merged:
            merged[key].extend(r.get(key) or [])

        logger.info(
            f"[FINAL DISCOVERY COUNTS] "
            f"diseases={len(merged['diseases'])} "
            f"subtypes={len(merged['subtypes'])} "
            f"drugs={len(merged['drugs'])} "
            f"biomarkers={len(merged['biomarkers'])}"
        )

        logger.info(
            "[FINAL DISCOVERY OUTPUT]\n"
            + json.dumps(merged, indent=2)[:15000]
        )
    for key in merged:
        seen = set()
        cleaned = []
        for item in merged[key]:
            if key == "subtypes":
                if isinstance(item, dict):
                    name = item.get("name", "").strip()
                    parents = item.get("parent_diseases") or []
                elif isinstance(item, str):
                    name, parents = item.strip(), []
                else:
                    continue
                if _is_noise_entity(name):
                    continue
                if name and name.lower() not in seen:
                    cleaned.append({"name": name, "parent_diseases": parents})
                    seen.add(name.lower())
                elif name and name.lower() in seen:
                    # merge parent lists across duplicate batches
                    for c in cleaned:
                        if c["name"].lower() == name.lower():
                            c["parent_diseases"] = list(set(c["parent_diseases"] + parents))
            else:
                # existing string/dict handling for diseases, stages, risk_groups, biomarkers, etc.
                if isinstance(item, dict):
                    name = item.get("name", "").strip()
                elif isinstance(item, str):
                    name = item.strip()
                else:
                    continue
                if _is_noise_entity(name):
                    continue
                if name and name.lower() not in seen:
                    cleaned.append(item if isinstance(item, dict) else {"name": name, "source_page": None})
                    seen.add(name.lower())
        merged[key] = cleaned

    # ── NEW: semantic dedup pass (embedding-based, catches near-duplicates
    # that exact-string matching above misses, e.g. "HER2+" vs "HER2 Positive") ──
    for key in ("diseases", "biomarkers"):
        merged[key] = _semantic_dedup(merged[key])

    logger.info(
        "[Discovery Merged Output]\n"
        + json.dumps(merged, indent=2)[:5000]
    )

    logger.info(
        f"[DiscoveryAgent] diseases={len(merged['diseases'])} subtypes={len(merged['subtypes'])} "
        f"biomarkers={len(merged['biomarkers'])} stages={len(merged['stages'])}"
    )

    if not merged["diseases"]:
        raise ValueError(
            f"Discovery failed. Merged output={merged}"
        )
    return merged

def _extract_entity_names(items):
    names = []

    for item in items:
        if isinstance(item, dict):
            name = item.get("name", "").strip()
            if name:
                names.append(name)

        elif isinstance(item, str):
            if item.strip():
                names.append(item.strip())

    return names

async def save_registries(
    doctor_id: str,
    doc_id: str,
    discovery: dict,
    mongo_uri: str,
) -> None:
    """Persist disease_registry and subtype_registry collections to MongoDB."""
    if not mongo_uri:
        return
    try:
        motor_client = AsyncIOMotorClient(mongo_uri)
        db           = motor_client[MONGO_DB]
        now          = datetime.now(timezone.utc).isoformat()

        disease_reg_coll  = db["disease_registry"]
        subtype_reg_coll  = db["subtype_registry"]

        diseases = _extract_entity_names(
            discovery.get("diseases", [])
        )

        subtypes = _extract_entity_names(
            discovery.get("subtypes", [])
        )
        if diseases:
            for disease in diseases:
                await disease_reg_coll.update_one(
                    {"disease": disease, "doc_id": doc_id},
                    {"$set": {
                        "disease": disease, "subtypes": subtypes,
                        "doc_id": doc_id, "doctor_id": doctor_id, "updated_at": now,
                    }},
                    upsert=True,
                )

        if subtypes:
            for subtype in subtypes:
                await subtype_reg_coll.update_one(
                    {"subtype": subtype, "doc_id": doc_id},
                    {"$set": {
                        "subtype": subtype, "related_diseases": diseases,
                        "doc_id": doc_id, "doctor_id": doctor_id, "updated_at": now,
                    }},
                    upsert=True,
                )

        motor_client.close()
        logger.info(
            f"[Registries] Saved {len(diseases)} disease records, "
            f"{len(subtypes)} subtype records | doc_id={doc_id}"
        )
    except Exception as e:
        logger.warning(f"[Registries] Failed to save: {e}")


async def save_subtype_taxonomy(
    doctor_id: str, doc_id: str,
    primary_disease: str,
    subtype_parent_map: dict[str, str],
    mongo_uri: str,
) -> None:
    if not mongo_uri:
        return
    try:
        client = AsyncIOMotorClient(mongo_uri)
        db = client[MONGO_DB]
        coll = db["phase1_subtype_taxonomy"]
        now = datetime.now(timezone.utc).isoformat()

        disease_node_id = f"{doc_id}:disease:{primary_disease}"
        await coll.update_one(
            {"node_id": disease_node_id, "doc_id": doc_id},
            {"$set": {
                "node_id": disease_node_id, "level": "disease",
                "name": primary_disease, "parent_id": None,
                "doc_id": doc_id, "doctor_id": doctor_id,
                "node_score": 1.0, "updated_at": now,
            }}, upsert=True,
        )
        for subtype_name, parent in subtype_parent_map.items():
            if not subtype_name:
                continue
            parent_name = parent or primary_disease
            sub_node_id = f"{doc_id}:subtype:{subtype_name}"
            await coll.update_one(
                {"node_id": sub_node_id, "doc_id": doc_id},
                {"$set": {
                    "node_id": sub_node_id, "level": "subtype",
                    "name": subtype_name, "parent_id": f"{doc_id}:disease:{parent_name}",
                    "doc_id": doc_id, "doctor_id": doctor_id,
                    "node_score": 1.0, "updated_at": now,
                }}, upsert=True,
            )
        client.close()
        logger.info(f"[SubtypeTaxonomy] saved {len(subtype_parent_map)} subtype nodes | doc_id={doc_id}")
    except Exception as e:
        logger.warning(f"[SubtypeTaxonomy] save failed: {e}")

async def _classify_primary_vs_differential(
    discovery: dict,
    filename: str,
    semaphore: asyncio.Semaphore,
) -> dict:
    """
    Use a single LLM call to separate primary diseases (the guideline's subject)
    from co-mentioned conditions (differentials, comorbidities, related diseases).
    """
    all_diseases = _extract_entity_names(
        discovery.get("diseases", [])
    )
    if not all_diseases:
        return discovery

    try:
        async with semaphore:
            result = await _chat_async(
                system=(
                    "You are a clinical knowledge classification expert. "
                    "Given a list of disease names extracted from a clinical guideline and the guideline filename, "
                    "classify each disease as either:\n"
                    "  - 'primary': the main condition(s) the guideline covers\n"
                    "  - 'differential': conditions mentioned only as differentials, comorbidities, or comparators\n"
                    "Return ONLY valid JSON:\n"
                    "{\n"
                    "  \"primary_diseases\": [\"disease1\"],\n"
                    "  \"differential_diseases\": [\"disease2\"]\n"
                    "}"
                ),
                user=(
                    f"Guideline filename: {filename}\n"
                    f"All extracted disease names: {all_diseases}"
                ),
            )
        primary = result.get("primary_diseases") or []
        differential = result.get("differential_diseases") or []
        if primary:
            discovery["primary_diseases"] = primary
            discovery["differential_diseases"] = differential
            discovery["diseases"] = primary
            logger.info(f"[DiseaseClassifier] primary={primary} | differential={differential}")
    except Exception as e:
        logger.warning(f"[DiseaseClassifier] failed: {e}")

    return discovery


# ═════════════════════════════════════════════════════════════════
# SECTION 7 — TEXT CLEANING
# ═════════════════════════════════════════════════════════════════

_NOISE = [
    r"All rights reserved", r"Downloaded from", r"Printed by",
    r"www\.[^\s]+", r"http[s]?://\S+", r"Page \d+", r"Table of Contents",
]


def _clean(text: str) -> str:
    text = re.sub(r"\r\n|\r", "\n", text)
    for p in _NOISE:
        text = re.sub(p, "", text, flags=re.IGNORECASE)
    text = re.sub(r"[ \t]+",  " ",    text)
    text = re.sub(r"\n{3,}",  "\n\n", text)
    text = re.sub(r"[_\-]{4,}", "",   text)
    return text.strip()


# ═════════════════════════════════════════════════════════════════
# SECTION 8 — PDF / TEXT EXTRACTION  (kept fully — required stage)
# ═════════════════════════════════════════════════════════════════

def _extract_table_from_blocks(page_dict: dict) -> str:
    blocks = page_dict.get("blocks", [])
    text_blocks = []
    for b in blocks:
        if b.get("type") == 0:
            bbox = b.get("bbox", (0, 0, 0, 0))
            text = ""
            for line in b.get("lines", []):
                for span in line.get("spans", []):
                    text += span.get("text", "")
                text += " "
            text = text.strip()
            if text:
                text_blocks.append({"bbox": bbox, "text": text, "y": bbox[1], "x": bbox[0]})

    if len(text_blocks) < 4:
        return ""

    rows: dict[int, list] = {}
    for tb in text_blocks:
        y_key = round(tb["y"] / 5) * 5
        rows.setdefault(y_key, [])
        rows[y_key].append(tb)

    table_rows = {y: sorted(cells, key=lambda c: c["x"])
                  for y, cells in rows.items() if len(cells) >= 2}

    if len(table_rows) < 2:
        return ""

    table_lines = []
    for y in sorted(table_rows.keys()):
        cells = table_rows[y]
        table_lines.append(" | ".join(c["text"] for c in cells))

    if table_lines:
        return "[TABLE]\n" + "\n".join(table_lines) + "\n[/TABLE]\n"
    return ""


def _extract_via_fitz(file_bytes: bytes, filename: str) -> tuple[str, list[dict]]:
    import fitz
    full_text:   list[str]  = []
    toc_entries: list[dict] = []

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    toc = doc.get_toc()
    for level, title, page in toc:
        toc_entries.append({"level": level, "title": title.strip(), "page": page})

    for i, page in enumerate(doc):
        text = page.get_text("text")
        try:
            page_dict  = page.get_text("dict")
            table_text = _extract_table_from_blocks(page_dict)
        except Exception:
            table_text = ""

        page_content = f"\n--- PAGE {i+1} ---\n"
        if text.strip():
            page_content += text
        if table_text:
            page_content += "\n" + table_text
        full_text.append(page_content)

    doc.close()
    combined = _clean("\n".join(full_text))
    logger.info(f"fitz: {len(full_text)} pages, {len(combined):,} chars, {len(toc_entries)} ToC entries")
    return combined, toc_entries


def _extract_via_ocr(pdf_path: str, page_count: int) -> str:
    import glob as _glob
    import pytesseract
    from PIL import Image
    full_text: list[str] = []
    for page_num in range(1, (page_count or 999) + 1):
        try:
            subprocess.run(
                ["pdftoppm", "-jpeg", "-r", "200",
                 "-f", str(page_num), "-l", str(page_num),
                 pdf_path, "/tmp/ocr_page"],
                check=True, capture_output=True, timeout=60,
            )
        except subprocess.CalledProcessError:
            break
        files = sorted(_glob.glob("/tmp/ocr_page-*.jpg"))
        if not files:
            break
        try:
            text = pytesseract.image_to_string(Image.open(files[0]))
            full_text.append(f"--- PAGE {page_num} ---\n{text}")
        finally:
            for f in files:
                try:
                    os.remove(f)
                except OSError:
                    pass
    return _clean("\n".join(full_text))


def detect_pdf_type(pdf_path: str) -> dict:
    try:
        pdfinfo_out = subprocess.run(
            ["pdfinfo", pdf_path], capture_output=True, text=True, timeout=30
        ).stdout
        page_count = 0
        for line in pdfinfo_out.splitlines():
            if line.startswith("Pages:"):
                page_count = int(line.split(":")[1].strip())
        pdffonts_out = subprocess.run(
            ["pdffonts", pdf_path], capture_output=True, text=True, timeout=30
        ).stdout
        font_lines = [l for l in pdffonts_out.strip().splitlines() if l.strip()]
        has_fonts  = len(font_lines) > 2
        sample = subprocess.run(
            ["pdftotext", "-f", "1", "-l", "1", pdf_path, "-"],
            capture_output=True, text=True, timeout=30
        ).stdout.strip()
        has_text = has_fonts and len(sample) > 50
    except (FileNotFoundError, subprocess.TimeoutExpired):
        logger.warning("poppler not available; assuming text-based PDF")
        page_count, has_text = 0, True
    return {"page_count": page_count, "has_text": has_text,
            "needs_ocr": not has_text, "pdf_path": pdf_path}


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, list[dict]]:
    ext = Path(filename).suffix.lower()
    toc_entries: list[dict] = []

    if ext == ".pdf":
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        detection = detect_pdf_type(tmp_path)
        try:
            if detection["needs_ocr"]:
                full = _extract_via_ocr(tmp_path, detection["page_count"])
            else:
                try:
                    full, toc_entries = _extract_via_fitz(file_bytes, filename)
                except Exception as e:
                    logger.warning(f"fitz failed ({e}) — falling back to OCR")
                    full = _extract_via_ocr(tmp_path, detection.get("page_count", 0))
        finally:
            try:
                os.remove(tmp_path)
            except OSError:
                pass
        if not full.strip():
            raise ValueError(f"No text extracted from: {filename}")
        return full, toc_entries

    if ext in (".docx", ".doc"):
        from docx import Document
        doc    = Document(io.BytesIO(file_bytes))
        chunks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                rt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if rt:
                    chunks.append(rt)
        full = _clean("\n\n".join(chunks))
        if not full.strip():
            raise ValueError(f"No text extracted from: {filename}")
        return full, []

    if ext in (".txt", ".md", ".csv"):
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1")
        full = _clean(text)
        if not full.strip():
            raise ValueError(f"Empty file: {filename}")
        return full, []

    raise ValueError(f"Unsupported file type: {ext}")


# ═════════════════════════════════════════════════════════════════
# SECTION 9 — HEADING-AWARE SECTION DETECTION  (kept fully)
# ═════════════════════════════════════════════════════════════════

_HEADING_PATTERNS = [
    (1, re.compile(r"^([A-Z]{2,8}-\d+[A-Z]?)\s*$", re.MULTILINE)),
    (1, re.compile(r"^([A-Z]{2,8}-\d+[A-Z]?)\s+(.+)$", re.MULTILINE)),
    (1, re.compile(r"^(\d+)\.\s+([A-Z].{3,60})$", re.MULTILINE)),
    (2, re.compile(r"^(\d+\.\d+)\s+([A-Z].{3,60})$", re.MULTILINE)),
    (3, re.compile(r"^(\d+\.\d+\.\d+)\s+([A-Z].{2,50})$", re.MULTILINE)),
    (1, re.compile(r"^([A-Z][A-Z\s]{4,50})$", re.MULTILINE)),
    (1, re.compile(r"^([IVX]{1,6})\.\s+(.+)$", re.MULTILINE)),
    (1, re.compile(r"^#{1}\s+(.+)$", re.MULTILINE)),
    (2, re.compile(r"^#{2}\s+(.+)$", re.MULTILINE)),
    (3, re.compile(r"^#{3}\s+(.+)$", re.MULTILINE)),
]


def _extract_page_number(line: str) -> Optional[int]:
    m = re.match(r"---\s*PAGE\s+(\d+)\s*---", line.strip())
    return int(m.group(1)) if m else None


def _is_heading(line: str) -> tuple[bool, int, str]:
    stripped = line.strip()
    if not stripped or len(stripped) > 120:
        return False, 0, ""
    for level, pattern in _HEADING_PATTERNS:
        m = pattern.match(stripped)
        if m:
            groups = m.groups()
            title  = groups[-1].strip() if groups else stripped
            return True, level, title
    return False, 0, ""


def _detect_sections_from_toc(full_text: str, toc_entries: list[dict]) -> list[Section]:
    if not toc_entries:
        return []
    anchors = [(e["page"], e["title"], e["level"]) for e in toc_entries]
    anchors.sort(key=lambda x: x[0])
    page_texts: dict[int, str] = {}
    current_page  = 1
    current_lines: list[str] = []
    for line in full_text.splitlines():
        pn = _extract_page_number(line)
        if pn is not None:
            page_texts[current_page] = "\n".join(current_lines)
            current_page  = pn
            current_lines = []
        else:
            current_lines.append(line)
    page_texts[current_page] = "\n".join(current_lines)
    sections: list[Section] = []
    for i, (page, title, level) in enumerate(anchors):
        end_page      = anchors[i + 1][0] - 1 if i + 1 < len(anchors) else current_page
        content_parts = [page_texts.get(p, "") for p in range(page, end_page + 1)]
        content       = (f"--- PAGE {page} ---\n" + "\n".join(content_parts)).strip()
        if content:
            sections.append(Section(
                text=content, section_id=f"toc_{i:04d}", section_title=title,
                page_start=page, page_end=end_page, section_index=i, heading_level=level,
            ))
    logger.info(f"ToC-based sections: {len(sections)}")
    return sections


def _detect_sections_from_headings(full_text: str) -> list[Section]:
    lines          = full_text.splitlines()
    sections:      list[Section] = []
    buffer:        list[str]     = []
    current_page   = 1
    sec_page_start = 1
    sec_title      = "Preamble"
    sec_id         = "preamble"
    sec_level      = 1
    sec_index      = 0
    heading_count  = 0

    def _flush():
        nonlocal sec_index
        text = "\n".join(buffer).strip()
        if text:
            sections.append(Section(
                text=text, section_id=sec_id, section_title=sec_title,
                page_start=sec_page_start, page_end=current_page,
                section_index=sec_index, heading_level=sec_level,
            ))
            sec_index += 1
        buffer.clear()

    for line in lines:
        pn = _extract_page_number(line)
        if pn is not None:
            current_page = pn
            buffer.append(line)
            continue
        is_h, level, title = _is_heading(line)
        if is_h:
            heading_count += 1
            _flush()
            sec_title      = title or line.strip()[:80]
            sec_id         = re.sub(r"\W+", "_", sec_title.lower())[:40]
            sec_level      = level
            sec_page_start = current_page
            buffer.append(line)
        else:
            buffer.append(line)
    _flush()
    logger.info(f"Heading-based sections: {heading_count} headings → {len(sections)} sections")
    return sections


def _merge_small_sections(sections: list[Section]) -> list[Section]:
    import dataclasses as _dc
    original_count = len(sections)
    merged: list[Section] = []
    buf: Optional[Section] = None
    for sec in sections:
        if buf is None:
            buf = sec
            continue
        merged_size = len(buf.text) + len(sec.text)
        if len(buf.text) < MIN_SECTION_CHARS and merged_size <= MAX_MERGE_RESULT_CHARS:
            buf = _dc.replace(buf, text=buf.text + "\n\n" + sec.text, page_end=sec.page_end)
        else:
            merged.append(buf)
            buf = sec
    if buf:
        merged.append(buf)
    if original_count >= 5 and len(merged) == 1:
        logger.warning(f"Over-merge detected: {original_count} → 1. Reverting.")
        return sections
    logger.info(f"Section merge: {original_count} → {len(merged)} sections")
    return merged


def _split_oversized_sections(sections: list[Section]) -> list[Section]:
    import dataclasses as _dc
    result: list[Section] = []
    for sec in sections:
        if len(sec.text) <= MAX_SECTION_CHARS:
            result.append(sec)
            continue
        paragraphs    = sec.text.split("\n\n")
        current_parts: list[str] = []
        part_idx      = 0
        current_len   = 0
        for para in paragraphs:
            if current_len + len(para) > MAX_SECTION_CHARS and current_parts:
                result.append(_dc.replace(
                    sec, text="\n\n".join(current_parts),
                    section_id=f"{sec.section_id}_p{part_idx}",
                    section_index=sec.section_index * 100 + part_idx,
                ))
                part_idx += 1
                current_parts = [para]
                current_len   = len(para)
            else:
                current_parts.append(para)
                current_len += len(para)
        if current_parts:
            result.append(_dc.replace(
                sec, text="\n\n".join(current_parts),
                section_id=f"{sec.section_id}_p{part_idx}",
                section_index=sec.section_index * 100 + part_idx,
            ))
    if len(result) != len(sections):
        logger.info(f"Split oversized: {len(sections)} → {len(result)} sections")
    return result


def _log_section_sizes(sections: list[Section], label: str) -> None:
    logger.info(f"── Section sizes after {label} ({len(sections)} sections) ──")
    total = sum(len(s.text) for s in sections)
    for s in sections:
        pct = round(len(s.text) / max(total, 1) * 100, 1)
        logger.info(
            f"  [{s.section_index:03d}] {s.section_title[:55]:<55} "
            f"{len(s.text):>8,} chars  ({pct:5.1f}%)  pp {s.page_start}-{s.page_end}"
        )
    logger.info(f"  Total: {total:,} chars across {len(sections)} sections")


def _toc_quality_ok(sections: list[Section]) -> bool:
    if not sections:
        return False
    avg = sum(len(s.text) for s in sections) / len(sections)
    if avg < TOC_MIN_AVG_SECTION_CHARS:
        logger.warning(f"ToC quality FAILED: avg {avg:,.0f} chars < threshold.")
        return False
    total   = sum(len(s.text) for s in sections)
    max_sec = max(len(s.text) for s in sections)
    if total > 0 and max_sec / total > 0.80:
        logger.warning(f"ToC quality FAILED: largest section is {max_sec/total*100:.0f}% of total.")
        return False
    return True


def _fixed_char_sections(full_text: str, chunk_size: int = MAX_SECTION_CHARS) -> list[Section]:
    sections: list[Section] = []
    page_re  = re.compile(r"---\s*PAGE\s+(\d+)\s*---")
    start    = 0
    idx      = 0
    while start < len(full_text):
        end = min(start + chunk_size, len(full_text))
        if end < len(full_text):
            b = full_text.rfind("\n\n", start, end)
            if b != -1 and b > start + chunk_size // 2:
                end = b
        chunk_text = full_text[start:end].strip()
        if chunk_text:
            pages   = [int(m.group(1)) for m in page_re.finditer(chunk_text)]
            p_start = pages[0]  if pages else 1
            p_end   = pages[-1] if pages else 1
            sections.append(Section(
                text=chunk_text, section_id=f"chunk_{idx:04d}",
                section_title=f"Section {idx + 1}",
                page_start=p_start, page_end=p_end, section_index=idx,
            ))
            idx += 1
        if end >= len(full_text):
            break
        start = end
    logger.info(f"Fixed-char fallback: {len(sections)} chunks")
    return sections


def build_sections(full_text: str, toc_entries: list[dict]) -> list[Section]:
    toc_sections: list[Section] = []
    if toc_entries:
        toc_sections = _detect_sections_from_toc(full_text, toc_entries)
        _log_section_sizes(toc_sections, "ToC split")
    use_toc = bool(toc_sections) and len(toc_sections) >= 3 and _toc_quality_ok(toc_sections)
    if use_toc:
        sections = toc_sections
        logger.info(f"Using ToC-based sections: {len(sections)}")
    else:
        logger.info("Using heading-pattern section detection")
        sections = _detect_sections_from_headings(full_text)
        _log_section_sizes(sections, "heading detection")
    if not sections:
        logger.warning("No sections detected — splitting full text into fixed char chunks")
        sections = _fixed_char_sections(full_text)
    sections = _merge_small_sections(sections)
    _log_section_sizes(sections, "merge")
    sections = _split_oversized_sections(sections)
    _log_section_sizes(sections, "final")
    logger.info(f"✅ Final sections ready: {len(sections)}")
    return sections


# ═════════════════════════════════════════════════════════════════
# SECTION 10 — MERGE HELPERS
# ═════════════════════════════════════════════════════════════════

def _unique_list(items: list) -> list:
    seen, out = set(), []
    for item in items:
        try:
            key = json.dumps(item, sort_keys=True) if isinstance(item, (dict, list)) else str(item)
        except (TypeError, ValueError):
            key = str(item)
        if key not in seen:
            seen.add(key)
            out.append(item)
    return out


def _safe_page_int(p) -> int | None:
    if isinstance(p, int):   return p
    if isinstance(p, float): return int(p)
    if isinstance(p, str):
        try: return int(p)
        except: return None
    if isinstance(p, dict):
        for key in ("page", "source_page", "start", "page_start"):
            v = p.get(key)
            if isinstance(v, (int, float)): return int(v)
        return None
    if isinstance(p, list):
        for x in p:
            v = _safe_page_int(x)
            if v is not None: return v
        return None
    return None


def _clean_pages(pages: list) -> list[int]:
    clean = []
    for p in pages:
        if isinstance(p, int):
            clean.append(p)
        elif isinstance(p, float):
            clean.append(int(p))
        elif isinstance(p, list):
            for x in p:
                v = _safe_page_int(x)
                if v is not None:
                    clean.append(v)
        else:
            v = _safe_page_int(p)
            if v is not None:
                clean.append(v)
    return sorted(set(clean))


def _sanitise_raw_llm_body(body: dict) -> dict:
    if not isinstance(body, dict):
        return body
    if "source_pages" in body:
        body["source_pages"] = _clean_pages(body.get("source_pages") or [])
    return body

def _is_noise_entity(name: str) -> bool:
    if not name:
        return True

    name = name.strip()

    noise_patterns = [
        r"^source_page$",
        r"^source_pages$",
        r"^:\s*\d+$",
        r"^\d+$",
        r"^null$",
        r"^none$",
        r"^page$",
        r"^source$",
    ]

    return any(
        re.match(p, name, re.IGNORECASE)
        for p in noise_patterns
    )

def _merge_understanding(results: list[dict]) -> dict:
    merged: dict = {
        "disease_name": None, "disease_category": None, "disease_type": None,
        "diseases": [], "subtypes": [], "stages": [], "biomarkers": [],
        "investigations": [], "procedures": [], "drugs": [], "regimens": [],
        "guideline_name": None, "guideline_version": None,
        "specialty": None, "specialties": [],
    }
    list_fields   = {"diseases", "subtypes", "stages", "biomarkers", "investigations", "procedures", "drugs", "regimens", "specialties"}
    scalar_fields = {"disease_name", "disease_category", "disease_type", "guideline_name", "guideline_version", "specialty"}
    _DESCRIPTION_PHRASES = {"disorder", "disease of", "condition", "syndrome of", "benign", "malignant", "inflammatory", "autoimmune", "neoplasm", "neoplastic"}

    def _is_classification_phrase(s: str) -> bool:
        if not s: return False
        return any(kw in s.lower() for kw in _DESCRIPTION_PHRASES)

    for r in results:
        if not isinstance(r, dict): continue
        for f in list_fields:
            merged[f].extend(r.get(f) or [])
        for f in scalar_fields:
            if not merged[f] and r.get(f):
                merged[f] = r[f]

    for f in list_fields:
        merged[f] = _unique_list(merged[f])

    flat_subtypes = []
    for s in merged["subtypes"]:
        if isinstance(s, list):
            flat_subtypes.extend(str(x).strip() for x in s if x)
        elif isinstance(s, dict):
            name = s.get("name") or s.get("subtype") or s.get("disease_name") or ""
            if name.strip(): flat_subtypes.append(name.strip())
        elif isinstance(s, str) and s.strip():
            flat_subtypes.append(s.strip())
    merged["subtypes"] = _unique_list(flat_subtypes)

    if not merged.get("specialty") and merged.get("specialties"):
        merged["specialty"] = merged["specialties"][0]

    dn = merged.get("disease_name", "")
    dc = merged.get("disease_type", "")
    if dn and _is_classification_phrase(dn):
        merged["disease_category"] = merged["disease_category"] or dn
        merged["disease_name"]     = None
    if not merged.get("disease_name") and merged.get("diseases"):
        for candidate in merged["diseases"]:
            if not _is_classification_phrase(candidate):
                merged["disease_name"] = candidate
                break
    if not merged.get("disease_category"):
        if dc and _is_classification_phrase(dc):
            merged["disease_category"] = dc
        elif dn and _is_classification_phrase(dn):
            merged["disease_category"] = dn
    merged["disease_type"] = merged.get("disease_name") or merged.get("disease_category") or "Unknown"

    all_disease_names = list(dict.fromkeys(
        d.strip() for d in (
            ([merged["disease_name"]] if merged.get("disease_name") else [])
            + merged.get("diseases", [])
        )
        if d and isinstance(d, str) and d.strip()
    ))
    merged["diseases"] = all_disease_names
    return merged


def _merge_obj(old: dict, new: dict) -> dict:
    for k, v in new.items():
        if k.startswith("_"):
            continue
        if isinstance(v, list):
            existing = old.get(k)
            if not isinstance(existing, list):
                existing = [existing] if existing else []
            old[k] = _unique_list(existing + v)
        elif isinstance(v, dict):
            existing = old.get(k)
            if isinstance(existing, dict):
                old[k] = _merge_obj(existing, v)
            elif not existing:
                old[k] = v
        elif isinstance(v, str):
            if not v.strip():
                continue
            old_v = old.get(k, "")
            if isinstance(old_v, dict):
                pass
            elif isinstance(old_v, list):
                if v not in old_v:
                    old_v.append(v)
            elif not old_v:
                old[k] = v
            elif v and v not in old_v:
                old[k] = old_v + " " + v
    return old


def _empty_diagnosis() -> dict:
    return {
        "disease_overview":       {"definition": "", "epidemiology": "", "risk_factors": []},
        "clinical_presentation":  {"symptoms": [], "signs": [], "chief_complaints": []},
        "diagnostic_criteria":    "",
        "investigations":         {"laboratory_tests": [], "imaging": [], "pathology": []},
        "biomarkers":             [],
        "molecular_testing":      {"genetic_mutations": [], "genomic_tests": []},
        "staging":                [],
        "risk_stratification":    {"low_risk": "", "intermediate_risk": "", "high_risk": ""},
        "subtypes":               [],
        "diagnostic_pathway":     [],
        "special_populations":    {"elderly": "", "pregnancy": "", "renal_impairment": "", "hepatic_impairment": "", "pediatric": ""},
        "key_evidence":           [],
        "differential_diagnosis": [],
        "exclusion_criteria":     [],
        "skill_boundaries":       {"does_not_cover": [], "related_skills": []},
        "gaps":                   [],
        "source_pages":           [],
    }


def _empty_treatment() -> dict:
    return {
        "treatment_principles":  "",
        "stage_wise_treatment":  [],
        "surgery":               {"procedures": [], "indications": "", "notes": ""},
        "radiation":             {"indications": "", "protocols": [], "dose": ""},
        "chemotherapy":          {"regimens": []},
        "therapeutic_procedures": [],   # NEW (Priority 1) — catches items the LLM
                                         # mis-placed under investigations that are
                                         # actually interventions (surgery, radiotherapy,
                                         # chemoimmunotherapy, etc.), moved here by
                                         # _reclassify_investigations post-processing.
        "immunotherapy":         {"drugs": [], "indications": "", "biomarker_selection": ""},
        "targeted_therapy":      {"drugs": [], "targets": [], "indications": ""},
        "hormone_therapy":       {"drugs": [], "indications": ""},
        "follow_up":             {"schedule": "", "monitoring_tests": [], "recurrence_detection": ""},
        "supportive_care":       {"pain_management": "", "nutrition": "", "palliative_care": ""},
        "risk_stratification":   [],
        "contraindications":     [],
        "dose_modifications":    [],
        "monitoring":            {"parameters": [], "frequency": "", "red_flags": []},
        "special_populations":   {"elderly": "", "pregnancy": "", "renal_impairment": "", "hepatic_impairment": "", "pediatric": ""},
        "key_evidence":          [],
        "recommendations":       [],
        "if_then_rules":         [],
        "toxicity_monitoring":   [],
        "dose_hold_criteria":    [],
        "dose_resume_criteria":  [],
        "surveillance_schedule": [],
        "response_assessment":   {"criteria": "", "timing": "", "tools": []},
        "skill_boundaries":      {"does_not_cover": [], "related_skills": []},
        "gaps":                  [],
        "source_pages":          [],
    }


def _merge_diagnosis(results: list[dict]) -> dict:
    base = _empty_diagnosis()
    all_source_pages: list[int] = []
    for r in results:
        if not isinstance(r, dict): continue
        if r.get("_source_pages"):
            all_source_pages.extend(_clean_pages(r["_source_pages"]))
        base = _merge_obj(base, r)

    base["source_pages"] = _clean_pages(all_source_pages + base.get("source_pages", []))

    for sub in ("investigations", "molecular_testing", "clinical_presentation"):
        for k in list(base[sub].keys()):
            if isinstance(base[sub][k], list):
                base[sub][k] = _unique_list(base[sub][k])
    base["disease_overview"]["risk_factors"] = _unique_list(base["disease_overview"]["risk_factors"])
    for f in ("staging", "biomarkers", "subtypes", "diagnostic_pathway", "key_evidence", "gaps",
              "differential_diagnosis", "exclusion_criteria"):
        base[f] = _unique_list(base[f])
    for k in base["skill_boundaries"]:
        base["skill_boundaries"][k] = _unique_list(base["skill_boundaries"][k])

    base["staging"]   = _filter_empty_list_items(base["staging"],   ["stage", "criteria"])
    base["biomarkers"] = _filter_empty_list_items(base["biomarkers"], ["name"])
    base["differential_diagnosis"] = _filter_empty_list_items(base["differential_diagnosis"], ["condition"])
    base["diagnostic_pathway"] = _filter_empty_list_items(base["diagnostic_pathway"], ["step", "action"])
    base["key_evidence"] = _filter_empty_list_items(base["key_evidence"], ["trial", "finding"])
    base["exclusion_criteria"] = _filter_empty_list_items(base["exclusion_criteria"], ["condition"])
    return base


def _merge_treatment(results: list[dict]) -> dict:
    base = _empty_treatment()
    all_source_pages: list[int] = []
    for r in results:
        if not isinstance(r, dict): continue
        if r.get("_source_pages"):
            all_source_pages.extend(_clean_pages(r["_source_pages"]))
        base = _merge_obj(base, r)

    base["source_pages"] = _clean_pages(all_source_pages + base.get("source_pages", []))

    for sub in ("surgery", "radiation", "immunotherapy", "targeted_therapy", "hormone_therapy", "follow_up"):
        for k in list(base[sub].keys()):
            if isinstance(base[sub][k], list):
                base[sub][k] = _unique_list(base[sub][k])
    base["chemotherapy"]["regimens"] = _unique_list(base["chemotherapy"]["regimens"])
    for f in ("stage_wise_treatment", "contraindications", "dose_modifications",
              "risk_stratification", "key_evidence", "gaps", "if_then_rules", "recommendations"):
        base[f] = _unique_list(base[f])
    base["monitoring"]["parameters"] = _unique_list(base["monitoring"]["parameters"])
    base["monitoring"]["red_flags"]  = _unique_list(base["monitoring"]["red_flags"])
    for k in base["skill_boundaries"]:
        base["skill_boundaries"][k] = _unique_list(base["skill_boundaries"][k])

    base["contraindications"] = _filter_empty_list_items(base["contraindications"], ["drug_or_action", "condition"])
    base["recommendations"] = _filter_empty_list_items(base["recommendations"], ["statement"])
    base["if_then_rules"] = _filter_empty_list_items(base["if_then_rules"], ["condition", "action"])
    base["stage_wise_treatment"] = _filter_empty_list_items(base["stage_wise_treatment"], ["stage"])
    base["dose_modifications"] = _filter_empty_list_items(base["dose_modifications"], ["condition", "drug"])
    base["key_evidence"] = _filter_empty_list_items(base["key_evidence"], ["trial", "finding"])
    return base


def _attach_source_pages_to_all_items(body: dict, page_range: list[int]) -> dict:
    list_fields = [
        "biomarkers", "staging", "key_evidence", "differential_diagnosis",
        "exclusion_criteria", "if_then_rules", "contraindications",
        "dose_modifications", "recommendations", "diagnostic_pathway",
        "toxicity_monitoring", "dose_hold_criteria", "dose_resume_criteria",
        "surveillance_schedule",
    ]
    for field in list_fields:
        if field in body and isinstance(body[field], list):
            body[field] = _attach_page_to_items(body[field], page_range)
    return body


# ═════════════════════════════════════════════════════════════════
# SECTION 11 — EMPTY FIELD CLEANUP
# ═════════════════════════════════════════════════════════════════

def _strip_empty_fields(obj: Any, depth: int = 0) -> Any:
    if depth > 10:
        return obj
    if isinstance(obj, dict):
        cleaned = {}
        for k, v in obj.items():
            if k.startswith("_"):
                cleaned[k] = v
                continue
            cleaned_v = _strip_empty_fields(v, depth + 1)
            if cleaned_v is None:
                continue
            if isinstance(cleaned_v, str) and not cleaned_v.strip():
                continue
            if isinstance(cleaned_v, list) and len(cleaned_v) == 0:
                continue
            if isinstance(cleaned_v, dict) and len(cleaned_v) == 0:
                continue
            cleaned[k] = cleaned_v
        return cleaned if cleaned else None
    if isinstance(obj, list):
        cleaned = [_strip_empty_fields(item, depth + 1) for item in obj]
        cleaned = [item for item in cleaned if item is not None]
        cleaned = [item for item in cleaned
                   if not (isinstance(item, str) and not item.strip())
                   and not (isinstance(item, dict) and not item)]
        return cleaned
    if isinstance(obj, str):
        return obj.strip() if obj.strip() else None
    return obj


def _filter_empty_list_items(items: list, required_fields: list[str]) -> list:
    """Remove list items where ALL required fields are empty/None."""
    result = []
    for item in items:
        if not isinstance(item, dict):
            if item:
                result.append(item)
            continue
        has_content = any(
            item.get(f) and str(item.get(f)).strip()
            for f in required_fields
        )
        if has_content:
            result.append(item)
    return result


def clean_empty_fields(body: dict) -> dict:
    cleaned = _strip_empty_fields(body)
    if not isinstance(cleaned, dict):
        return body
    for key in ("source_pages", "skill_boundaries", "gaps"):
        if key not in cleaned:
            original = body.get(key)
            if original is not None:
                cleaned[key] = original
    return cleaned


# ═════════════════════════════════════════════════════════════════
# SECTION 12 — EXTRACTION PROMPTS  (kept — diagnosis/treatment only)
# ═════════════════════════════════════════════════════════════════

_DIAGNOSIS_EXTRACT_SYSTEM = """
You are a senior clinical diagnostician.
Extract ALL DIAGNOSIS knowledge. Every item MUST include source_page.
Return ONLY valid JSON:
{
  "diagnosis": {
    "disease_overview": {"definition":"","epidemiology":"","risk_factors":[]},
    "clinical_presentation": {"symptoms":[],"signs":[],"chief_complaints":[]},
    "diagnostic_criteria": "",
    "investigations": {"laboratory_tests":[],"imaging":[],"pathology":[]},
    "biomarkers": [{"name":"","significance":"","source_stated":true,"source_page":0}],
    "molecular_testing": {"genetic_mutations":[],"genomic_tests":[]},
    "staging": [{"stage":"","criteria":"","description":"","source_page":0}],
    "risk_stratification": {"low_risk":"","intermediate_risk":"","high_risk":""},
    "subtypes": [{"name":"","definition":"","key_features":[]}],
    "diagnostic_pathway": [{"step":1,"action":"","rationale":"","next_if_positive":"","next_if_negative":"","source_page":0}],
    "differential_diagnosis": [{"condition":"","distinguishing_feature":"","key_test":"","source_page":0}],
    "exclusion_criteria": [{"condition":"","test_to_exclude":"","rationale":""}],
    "special_populations": {"elderly":"","pregnancy":"","renal_impairment":"","hepatic_impairment":"","pediatric":""},
    "key_evidence": [{"trial":"","finding":"","source_section":"","source_page":0}],
    "skill_boundaries": {"does_not_cover":[],"related_skills":[]},
    "gaps": []
  },
  "llm_confidence": 1.0
}
Extract ONLY what is explicitly stated. Do not assume any specific specialty.
"""

_TREATMENT_EXTRACT_SYSTEM = """
You are a senior clinical specialist.
Extract ALL TREATMENT knowledge. Every item MUST include source_page.
Recommendations MUST include evidence_level if stated in the text.
Return ONLY valid JSON:
{
  "treatment": {
    "treatment_principles": "",
    "stage_wise_treatment": [{"stage":"","subtype":"","intent":"","primary_treatment":"","conditions":[],"options":[{"regimen_name":"","drugs":[],"modality":"","line":"","evidence_trial":"","notes":"","condition":"","source_page":0}],"surgery":{"procedure":"","indication":""},"radiation":{"protocol":"","indication":""}}],
    "surgery": {"procedures":[],"indications":"","notes":""},
    "radiation": {"indications":"","protocols":[],"dose":""},
    "chemotherapy": {"regimens":[{"name":"","drugs":[],"indication":""}]},
    "immunotherapy": {"drugs":[],"indications":"","biomarker_selection":""},
    "targeted_therapy": {"drugs":[],"targets":[],"indications":""},
    "hormone_therapy": {"drugs":[],"indications":""},
    "follow_up": {"schedule":"","monitoring_tests":[],"recurrence_detection":""},
    "supportive_care": {"pain_management":"","nutrition":"","palliative_care":""},
    "risk_stratification": [{"risk_group":"","criteria":"","implication":""}],
    "contraindications": [{"drug_or_action":"","condition":"","reason":"","source_page":0}],
    "dose_modifications": [{"condition":"","drug":"","modification":""}],
    "monitoring": {"parameters":[],"frequency":"","red_flags":[]},
    "special_populations": {"elderly":"","pregnancy":"","renal_impairment":"","hepatic_impairment":"","pediatric":""},
    "key_evidence": [{"trial":"","finding":"","source_section":"","source_page":0}],
    "recommendations": [{"statement":"","evidence_level":"","strength":"Strong|Moderate|Weak|Expert Opinion","conditional":false,"source_page":0}],
    "if_then_rules": [{"condition":"","action":"","population":"","biomarker_condition":"","source_page":0}],
    "skill_boundaries": {"does_not_cover":[],"related_skills":[]},
    "gaps": []
  },
  "llm_confidence": 1.0
}
"""


# ═════════════════════════════════════════════════════════════════
# SECTION 13 — QUERY BUILDER  (NEW — NO LLM, deterministic templates)
# ═════════════════════════════════════════════════════════════════

RETRIEVAL_QUERY_TEMPLATES: dict[str, str] = {
    "diagnosis": (
        "{disease}\n\n"
        "diagnosis criteria\nclinical presentation\nsymptoms signs\n"
        "investigations imaging laboratory tests\nbiomarkers molecular testing\nstaging risk stratification"
    ),
    "treatment": (
        "{disease}\n\n"
        "treatment therapy management\nsurgery procedure\nchemotherapy regimen\n"
        "radiation\ntargeted therapy immunotherapy hormone therapy\nfirst line second line maintenance"
    ),
    "evidence": (
        "{disease}\n\n"
        "clinical trial evidence\nrecommendation grade level\nlandmark study randomized"
    ),
    "monitoring": (
        "{disease}\n\n"
        "monitoring follow up\nsurveillance schedule\ntoxicity grading\ndose hold resume criteria\nred flags"
    ),
    "rare_population": (
        "{disease}\n\n"
        "pregnancy lactation\nelderly geriatric\nrenal impairment\nhepatic impairment\npediatric adolescent"
    ),
}


def build_agent_query(disease: str, query_type: str, biomarkers: list[str] = None, stages: list[str] = None) -> str:
    template = RETRIEVAL_QUERY_TEMPLATES.get(query_type)
    if not template:
        raise ValueError(f"Unknown query_type: {query_type}")
    base = template.format(disease=disease)
    extras = []
    if biomarkers:
        extras.append(" ".join(biomarkers[:8]))
    if stages and query_type in ("diagnosis", "treatment"):
        extras.append(" ".join(stages[:6]))
    if extras:
        base = base + "\n\n" + "\n".join(extras)
    return base


def build_context_aware_query(
    entity_name: str,
    parent_disease: Optional[str] = None,
    related_biomarkers: list[str] = None,
    related_stages: list[str] = None,
    related_risk_groups: list[str] = None,
) -> str:
    """
    Generic, specialty-agnostic context query — works for ANY entity type
    (subtype, biomarker-driven disease state, stage-specific lookup, etc.)
    by composing whatever relational context is actually known about it.
    No hardcoded entity-type logic — just concatenates whatever's passed in.
    """
    parts = [entity_name]
    if parent_disease:
        parts.insert(0, parent_disease)
    if related_biomarkers:
        parts.append(" ".join(related_biomarkers[:5]))
    if related_stages:
        parts.append(" ".join(related_stages[:3]))
    if related_risk_groups:
        parts.append(" ".join(related_risk_groups[:3]))
    return " ".join(p for p in parts if p)

def _dedup_sections(secs: list[Section]) -> list[Section]:
    seen, out = set(), []
    for s in secs:
        if s.section_id not in seen:
            seen.add(s.section_id)
            out.append(s)
    return out



# ═══════════════════════════════════════════════════════════════
# SECTION 13b — DISEASE REGISTRY / SKILL QUEUE BUILDER  (NEW)
# ═══════════════════════════════════════════════════════════════

@dataclass
class SkillTarget:
    name:          str
    target_type:   str   # "primary_disease" | "differential_disease" | "subtype"
    parent_disease: Optional[str] = None   # set for subtypes
    biomarkers:    list[str] = None
    stages:        list[str] = None


def build_skill_queue(discovery: dict) -> list[SkillTarget]:
    primary = _extract_entity_names(discovery.get("primary_diseases") or discovery.get("diseases") or [])
    differential = discovery.get("differential_diseases") or []
    biomarkers = discovery.get("biomarkers") or []
    stages = discovery.get("stages") or []

    queue: list[SkillTarget] = []
    for d in primary:
        queue.append(SkillTarget(d, "primary_disease", biomarkers=biomarkers, stages=stages))
    for d in differential:
        queue.append(SkillTarget(d, "differential_disease", biomarkers=biomarkers, stages=stages))

    for s in discovery.get("subtypes") or []:
        if isinstance(s, dict):
            name = s.get("name", "").strip()
            parents = s.get("parent_diseases") or []
        else:
            name, parents = str(s).strip(), []
        if not name:
            continue
        # use the actual discovered parent; fall back to primary[0] ONLY if discovery gave nothing
        parent = parents[0] if parents else (primary[0] if primary else None)
        queue.append(SkillTarget(name, "subtype", parent_disease=parent,
                                  biomarkers=biomarkers, stages=stages))
    return queue

# ═════════════════════════════════════════════════════════════════
# SECTION 14 — RETRIEVAL (NO LLM) + COVERAGE VALIDATOR  (NEW)
# ═════════════════════════════════════════════════════════════════

async def retrieve_context_for_disease(
    disease: str, doc_id: str,
    biomarkers: list[str] = None, stages: list[str] = None,
) -> dict[str, list[Section]]:
    loop = asyncio.get_event_loop()
    query_types = list(RETRIEVAL_QUERY_TEMPLATES.keys())

    async def _retrieve(qtype: str) -> tuple[str, list[Section]]:
        query = build_agent_query(disease, qtype, biomarkers, stages)
        sections = await loop.run_in_executor(None, lambda q=query: search_guideline(q, doc_id))
        return qtype, sections

    results = await asyncio.gather(*[_retrieve(qt) for qt in query_types])
    context = {qt: secs for qt, secs in results}
    logger.info(f"[Retrieval] '{disease}' | " + " | ".join(f"{k}={len(v)}" for k, v in context.items()))
    return context

def validate_disease_coverage(retrieved_by_disease: dict[str, dict[str, list[Section]]]) -> list[str]:
    """Diseases with ZERO retrieved sections across ALL query types — a retrieval gap."""
    gaps = [
        disease for disease, ctx in retrieved_by_disease.items()
        if sum(len(v) for v in ctx.values()) == 0
    ]
    if gaps:
        logger.warning(f"[CoverageValidator] zero-retrieval diseases: {gaps}")
    else:
        logger.info("[CoverageValidator] all diseases have retrieved content")
    return gaps


async def revalidate_and_retry(
    gaps: list[str],
    doc_id: str,
    retrieved_by_disease: dict[str, dict[str, list[Section]]],
) -> dict[str, dict[str, list[Section]]]:
    """
    For diseases with zero retrieval, fall back to a single broad query
    (just the disease name, no template) — still no LLM involved.
    """
    loop = asyncio.get_event_loop()
    for disease in gaps:
        sections = await loop.run_in_executor(None, lambda d=disease: search_guideline(d, doc_id))
        if sections:
            retrieved_by_disease[disease]["diagnosis"] = sections
            logger.info(f"[CoverageValidator] recovered {len(sections)} sections for '{disease}' via broad re-query")
        else:
            logger.warning(f"[CoverageValidator] '{disease}' has NO retrievable content even after broad re-query")
    return retrieved_by_disease


def compute_disease_page_coverage(
    retrieved_by_disease: dict[str, dict[str, list[Section]]],
    total_pages: int,
) -> dict:
    """Lightweight, retrieval-only coverage report (no LLM sweep needed)."""
    covered_pages: set[int] = set()
    for ctx in retrieved_by_disease.values():
        for secs in ctx.values():
            for sec in secs:
                if sec.page_start and sec.page_end:
                    covered_pages.update(range(sec.page_start, sec.page_end + 1))
    coverage_pct = round(len(covered_pages) / max(total_pages, 1) * 100, 1)
    logger.info(f"[Coverage] {len(covered_pages)}/{total_pages} pages covered ({coverage_pct}%)")
    return {
        "covered_pages": sorted(covered_pages),
        "coverage_pct": coverage_pct,
        "total_pages": total_pages,
    }


# ═════════════════════════════════════════════════════════════════
# SECTION 15 — DISEASE SKILL GENERATION  (NEW — replaces the 8-agent system)
# ═════════════════════════════════════════════════════════════════

async def generate_disease_knowledge(
    disease: str,
    doc_id: str,
    semaphore: asyncio.Semaphore,
    context: Optional[dict[str, list[Section]]] = None,
) -> tuple[dict, dict]:
    """
    ONE diagnosis LLM call + ONE treatment LLM call per disease, built from
    pre-retrieved (no-LLM) ChromaDB context. Diagnosis context = diagnosis +
    evidence + rare_population sections. Treatment context = treatment +
    evidence + monitoring + rare_population sections (folds in what used to be
    the separate Evidence/Monitoring/RarePopulation agents).
    """
    context = context or await retrieve_context_for_disease(disease, doc_id)

    diag_sections  = _dedup_sections(
        context.get("diagnosis", []) + context.get("evidence", []) + context.get("rare_population", [])
    )
    treat_sections = _dedup_sections(
        context.get("treatment", []) + context.get("evidence", []) +
        context.get("monitoring", []) + context.get("rare_population", [])
    )

    diag_result, treat_result = {}, {}

    if diag_sections:
        diag_msg = _build_agent_user_message("diagnosis", diag_sections)
        async with semaphore:
            diag_result = await _chat_async(
                _DIAGNOSIS_EXTRACT_SYSTEM + f"\n\nFocus ONLY on disease: {disease}", diag_msg
            )

    if treat_sections:
        treat_msg = _build_agent_user_message("treatment", treat_sections)
        async with semaphore:
            treat_result = await _chat_async(
                _TREATMENT_EXTRACT_SYSTEM + f"\n\nFocus ONLY on disease: {disease}", treat_msg
            )

    diag_body  = diag_result.get("diagnosis", {})
    treat_body = treat_result.get("treatment", {})

    diag_pages  = sorted({p for s in diag_sections  for p in range(s.page_start, s.page_end + 1)})
    treat_pages = sorted({p for s in treat_sections for p in range(s.page_start, s.page_end + 1)})

    if diag_body:
        diag_body["_source_pages"] = diag_pages
        diag_body = _attach_source_pages_to_all_items(diag_body, diag_pages)
        diag_body["_source_sections"] = [
            {"section_title": s.section_title, "page_start": s.page_start} for s in diag_sections
        ]
    if treat_body:
        treat_body["_source_pages"] = treat_pages
        treat_body = _attach_source_pages_to_all_items(treat_body, treat_pages)
        treat_body["_source_sections"] = [
            {"section_title": s.section_title, "page_start": s.page_start} for s in treat_sections
        ]

    diagnosis_knowledge = _merge_diagnosis([diag_body]) if diag_body else _empty_diagnosis()
    treatment_knowledge = _merge_treatment([treat_body]) if treat_body else _empty_treatment()

    logger.info(f"[DiseaseSkill] '{disease}': diag_sec={len(diag_sections)} treat_sec={len(treat_sections)}")
    return diagnosis_knowledge, treatment_knowledge


# ═══════════════════════════════════════════════════════════════
# SECTION 15b — BATCHED DISEASE SKILL GENERATION  (NEW)
# ═══════════════════════════════════════════════════════════════

DISEASE_BATCH_SIZE = 5  # diseases per LLM call — tune to context budget

_DISEASE_BATCH_DIAGNOSIS_SYSTEM = """
You are a senior clinical diagnostician.
You will be given MULTIPLE diseases, each with its own retrieved guideline sections.
For EACH disease, extract diagnosis knowledge using ONLY that disease's sections.
Never mix content between diseases. If a disease shows [NO CONTENT RETRIEVED], omit it.

Extract EVERY biomarker, investigation, staging entry, and evidence item
explicitly named in this disease's sections — do not select a representative
subset or collapse a list into a single example. If the text names 10
biomarkers, return 10.

In addition, for each disease identify the distinct clinical skills/competencies a
clinician must apply, based ONLY on what is explicitly covered in that disease's
retrieved sections. A skill is a discrete, actionable capability (e.g. interpreting
a specific named test, applying a specific named classification or criteria,
selecting among specific explicitly-named treatment or diagnostic options) — not a
restatement of the disease name and not a vague category. Generate as many distinct
skills as the retrieved content actually supports; do not pad, and do not invent
skills that are not grounded in the text. If the retrieved content does not support
any clinically meaningful skill, return an empty list for that disease.

Return ONLY valid JSON:
{
  "diseases": {
    "<disease name>": {
      "disease_overview": {"definition":"","epidemiology":"","risk_factors":[]},
      "clinical_presentation": {"symptoms":[],"signs":[],"chief_complaints":[]},
      "diagnostic_criteria": "",
      "investigations": {"laboratory_tests":[],"imaging":[],"pathology":[]},
      "biomarkers": [{"name":"","significance":"","source_page":0}],
      "staging": [{"stage":"","criteria":"","source_page":0}],
      "subtypes": [{"name":"","definition":""}],
      "diagnostic_pathway": [{"step":1,"action":"","source_page":0}],
      "differential_diagnosis": [{"condition":"","distinguishing_feature":"","source_page":0}],
      "key_evidence": [{"trial":"","finding":"","source_page":0}],
      "skills": [{"skill":"","basis":""}]
    }
  }
}
"""

_DISEASE_BATCH_TREATMENT_SYSTEM = """
You are a senior clinical specialist.
You will be given MULTIPLE diseases, each with its own retrieved guideline sections.
For EACH disease, extract treatment knowledge using ONLY that disease's sections.
Never mix content between diseases. If a disease shows [NO CONTENT RETRIEVED], omit it.

Extract EVERY biomarker, investigation, staging entry, and evidence item
explicitly named in this disease's sections — do not select a representative
subset or collapse a list into a single example. If the text names 10
biomarkers, return 10.

In addition, for each disease identify the distinct treatment-side clinical
skills/competencies a clinician must apply, based ONLY on what is explicitly
covered in that disease's retrieved sections (e.g. selecting among specific
named regimens, applying specific named dose-modification or monitoring rules,
recognizing specific named contraindications). Do not restate the disease name
as a skill and do not invent skills not grounded in the text. If nothing
meaningful is supported, return an empty list.

Return ONLY valid JSON:
{
  "diseases": {
    "<disease name>": {
      "treatment_principles": "",
      "stage_wise_treatment": [{"stage":"","intent":"","primary_treatment":"","options":[{"regimen_name":"","drugs":[],"source_page":0}]}],
      "chemotherapy": {"regimens": [{"name":"","drugs":[],"indication":""}]},
      "targeted_therapy": {"drugs": [], "targets": [], "indications": ""},
      "contraindications": [{"drug_or_action":"","condition":"","source_page":0}],
      "recommendations": [{"statement":"","evidence_level":"","source_page":0}],
      "if_then_rules": [{"condition":"","action":"","source_page":0}],
      "key_evidence": [{"trial":"","finding":"","source_page":0}],
      "skills": [{"skill":"","basis":""}]
    }
  }
}
"""


def _format_disease_block(disease: str, sections: list[Section], max_chars: int = 6000) -> str:
    if not sections:
        return f"=== DISEASE: {disease} ===\n[NO CONTENT RETRIEVED]"
    text = "\n\n".join(
        f"[page {s.page_start}-{s.page_end}] {s.section_title}\n{s.text[:3000]}" for s in sections
    )
    return f"=== DISEASE: {disease} ===\n{_trim_to_budget(text, max_chars)}"

def _backfill_disease_entities(
    body: dict,
    dotted_path: str,
    candidate_names: list[str],
    sections: list[Section],
    item_shape_fn,
) -> int:
    """
    Deterministic, no-LLM backfill. For each candidate name (from the
    Discovery Agent's document-wide lists) not already present at
    dotted_path in `body`, check whether it appears verbatim
    (case-insensitive) in the text of sections actually retrieved for this
    disease. If so, add it back with a source_page inferred from that
    section. Zero LLM calls.
    """
    existing = _resolve_dotted(body, dotted_path)
    if not isinstance(existing, list):
        existing = []
    existing_names = {_entity_display_name(e).lower() for e in existing}
    added = 0
    for name in candidate_names:
        if not name or name.lower() in existing_names:
            continue
        for sec in sections:
            if name.lower() in sec.text.lower():
                existing.append(item_shape_fn(name, sec.page_start))
                existing_names.add(name.lower())
                added += 1
                break
    node = body
    parts = dotted_path.split(".")
    for p in parts[:-1]:
        node = node.setdefault(p, {})
    node[parts[-1]] = existing
    return added

def _investigation_categories() -> list[str]:
    """Derived from _empty_diagnosis()['investigations'] keys — not a literal list."""
    return list(_empty_diagnosis()["investigations"].keys())


def _drug_categories() -> list[str]:
    """
    Any top-level treatment bucket that structurally holds a drug/regimen
    list in the schema qualifies as a drug category — found by introspection,
    not by naming specific therapy types. Works whether the schema has
    oncology fields, cardiology fields, or anything else, since it only
    checks shape ('does this bucket have a drugs/regimens list?'), not name.
    """
    t = _empty_treatment()
    return [
        key for key, val in t.items()
        if isinstance(val, dict) and any(k in val for k in ("drugs", "regimens"))
    ]
_CATEGORY_EMBED_CACHE: dict[str, list[float]] = {}

def _category_vectors(categories: list[str]) -> dict[str, list[float]]:
    missing = [c for c in categories if c not in _CATEGORY_EMBED_CACHE]
    if missing:
        vectors = _embed_batch(missing)
        for c, v in zip(missing, vectors):
            if v is not None:
                _CATEGORY_EMBED_CACHE[c] = v
    return {c: _CATEGORY_EMBED_CACHE[c] for c in categories if c in _CATEGORY_EMBED_CACHE}

def _classify_by_embedding(names: list[str], categories: list[str], default: str) -> dict[str, str]:
    """
    Batch-classifies candidate names against category label embeddings via
    cosine similarity — one embedding call for all names, not one per name.
    Falls back to `default` only if embedding is unavailable.
    """
    cat_vectors = _category_vectors(categories)
    if not cat_vectors or not names:
        return {n: default for n in names}
    name_vectors = _embed_batch(names)
    out: dict[str, str] = {}
    for name, vec in zip(names, name_vectors):
        if vec is None:
            out[name] = default
            continue
        best_cat, best_sim = default, float("-inf")
        for cat, cvec in cat_vectors.items():
            sim = sum(a * b for a, b in zip(vec, cvec))
            if sim > best_sim:
                best_sim, best_cat = sim, cat
        out[name] = best_cat
    return out

def _reclassify_investigations(diag_body: dict, treat_body: dict) -> int:
    """
    Priority 1 fix — the diagnosis-extraction LLM sometimes places therapeutic
    items (Surgery, Radiotherapy, Chemoimmunotherapy) inside
    investigations.<category> because they textually co-occur with diagnostic
    workup in the source guideline. This reclassifies every investigation
    string against the diagnosis schema's own investigation categories PLUS
    a generic "therapeutic procedure" anchor — via embedding similarity, no
    keyword table — and moves anything that scores as a therapeutic
    procedure into treatment.therapeutic_procedures instead of silently
    poisoning the diagnosis skill's investigation lists.
    """
    inv = diag_body.get("investigations")
    if not isinstance(inv, dict):
        return 0

    all_items: list[str] = []
    item_category: dict[str, str] = {}
    for cat, items in inv.items():
        if not isinstance(items, list):
            continue
        for it in items:
            name = it if isinstance(it, str) else (it.get("name") if isinstance(it, dict) else None)
            if name:
                all_items.append(name)
                item_category[name] = cat
    if not all_items:
        return 0

    # Generic anchor, not domain vocabulary — a structural label describing
    # what kind of thing an entity IS, same pattern as the drug/investigation
    # category classifier already used for backfill.
    categories = _investigation_categories() + ["therapeutic procedure"]
    classification = _classify_by_embedding(all_items, categories, default=item_category[all_items[0]])

    moved = 0
    for name in all_items:
        assigned = classification.get(name)
        if assigned == "therapeutic procedure":
            orig_cat = item_category[name]
            if isinstance(inv.get(orig_cat), list) and name in inv[orig_cat]:
                inv[orig_cat] = [x for x in inv[orig_cat]
                                  if (x if isinstance(x, str) else x.get("name")) != name]
                treat_body.setdefault("therapeutic_procedures", [])
                if name not in treat_body["therapeutic_procedures"]:
                    treat_body["therapeutic_procedures"].append(name)
                moved += 1
    return moved

def _backfill_categorized(
    body: dict,
    candidate_names: list[str],
    sections: list[Section],
    classification: dict[str, str],
    default_category: str,
    path_for_category_fn,
    item_shape_fn,
) -> dict[str, int]:
    if not candidate_names:
        return {}
    added_by_category: dict[str, int] = {}
    for name in candidate_names:
        if not name:
            continue
        category = classification.get(name, default_category)
        dotted_path = path_for_category_fn(category)
        n = _backfill_disease_entities(body, dotted_path, [name], sections, item_shape_fn)
        if n:
            added_by_category[category] = added_by_category.get(category, 0) + n
    return added_by_category

async def generate_disease_knowledge_batched(
    diseases: list[str],
    retrieved_by_disease: dict[str, dict[str, list[Section]]],
    semaphore: asyncio.Semaphore,
    discovery_biomarkers: Optional[list[str]] = None,      # NEW
    discovery_investigations: Optional[list[str]] = None,  # NEW
    discovery_drugs: Optional[list[str]] = None,            # NEW
) -> dict[str, tuple[dict, dict]]:
    biomarker_pool     = discovery_biomarkers or []
    investigation_pool = discovery_investigations or []
    drug_pool           = discovery_drugs or []
    if not diseases:
        return {}

    # Classify candidate pools ONCE — they're the same discovery-level lists
    # for every disease, so classifying inside the per-disease loop below
    # would re-embed identical names N times for N diseases.
    inv_classification = (
        _classify_by_embedding(investigation_pool, _investigation_categories(), "laboratory_tests")
        if investigation_pool else {}
    )
    drg_classification = (
        _classify_by_embedding(drug_pool, _drug_categories(), "chemotherapy")
        if drug_pool else {}
    )

    diag_sections_map: dict[str, list[Section]] = {}
    treat_sections_map: dict[str, list[Section]] = {}
    for d in diseases:
        ctx = retrieved_by_disease.get(d, {})
        diag_sections_map[d] = _dedup_sections(
            ctx.get("diagnosis", []) + ctx.get("evidence", []) + ctx.get("rare_population", [])
        )
        treat_sections_map[d] = _dedup_sections(
            ctx.get("treatment", []) + ctx.get("evidence", []) +
            ctx.get("monitoring", []) + ctx.get("rare_population", [])
        )

    batches = [diseases[i:i + DISEASE_BATCH_SIZE] for i in range(0, len(diseases), DISEASE_BATCH_SIZE)]
    logger.info(f"[DiseaseBatch] {len(diseases)} diseases → {len(batches)} batch(es) of ≤{DISEASE_BATCH_SIZE}")

    diag_results_by_disease: dict[str, dict] = {}
    treat_results_by_disease: dict[str, dict] = {}

    for batch_idx, batch in enumerate(batches):
        diag_payload  = "\n\n".join(_format_disease_block(d, diag_sections_map.get(d, [])) for d in batch)
        treat_payload = "\n\n".join(_format_disease_block(d, treat_sections_map.get(d, [])) for d in batch)

        if any(diag_sections_map.get(d) for d in batch):
            try:
                async with semaphore:
                    diag_result = await _chat_async(_DISEASE_BATCH_DIAGNOSIS_SYSTEM, diag_payload)
                diag_results_by_disease.update(diag_result.get("diseases", {}))
            except Exception as exc:
                logger.warning(f"[DiseaseBatch {batch_idx}] diagnosis call failed: {exc}")

        if any(treat_sections_map.get(d) for d in batch):
            try:
                async with semaphore:
                    treat_result = await _chat_async(_DISEASE_BATCH_TREATMENT_SYSTEM, treat_payload)
                treat_results_by_disease.update(treat_result.get("diseases", {}))
            except Exception as exc:
                logger.warning(f"[DiseaseBatch {batch_idx}] treatment call failed: {exc}")

        logger.info(f"[DiseaseBatch {batch_idx+1}/{len(batches)}] diseases={batch}")

    out: dict[str, tuple[dict, dict]] = {}
    for d in diseases:
        diag_secs  = diag_sections_map.get(d, [])
        treat_secs = treat_sections_map.get(d, [])
        diag_pages  = sorted({p for s in diag_secs  for p in range(s.page_start, s.page_end + 1)})
        treat_pages = sorted({p for s in treat_secs for p in range(s.page_start, s.page_end + 1)})

        diag_body  = diag_results_by_disease.get(d) or {}
        treat_body = treat_results_by_disease.get(d) or {}

        # Pull skills out BEFORE merging — _merge_diagnosis/_merge_treatment
        # don't know about this key and would otherwise drop it silently.
        diag_skills_raw  = diag_body.pop("skills", []) if isinstance(diag_body, dict) else []
        treat_skills_raw = treat_body.pop("skills", []) if isinstance(treat_body, dict) else []

        if diag_body:
            diag_body["_source_pages"] = diag_pages
            diag_body = _attach_source_pages_to_all_items(diag_body, diag_pages)
        if treat_body:
            treat_body["_source_pages"] = treat_pages
            treat_body = _attach_source_pages_to_all_items(treat_body, treat_pages)

        merged_diag  = _merge_diagnosis([diag_body]) if diag_body else _empty_diagnosis()
        merged_treat = _merge_treatment([treat_body]) if treat_body else _empty_treatment()

        merged_diag["skills"]  = [s for s in diag_skills_raw if isinstance(s, dict) and s.get("skill")]
        merged_treat["skills"] = [s for s in treat_skills_raw if isinstance(s, dict) and s.get("skill")]

        # NEW (Priority 1) — reclassify any investigation items that are
        # actually therapeutic procedures, before backfill runs.
        n_reclass = _reclassify_investigations(merged_diag, merged_treat)
        if n_reclass:
            logger.info(f"[Reclassify] {d}: moved {n_reclass} items from investigations → therapeutic_procedures")

        # NEW — deterministic backfill (Issues #1/#4/#5): recover entities
        # the LLM compressed away, verified against this disease's own
        # retrieved section text so nothing unrelated gets injected.
        n_bm = _backfill_disease_entities(
            merged_diag, "biomarkers", biomarker_pool, diag_secs,
            lambda name, page: {"name": name, "significance": "", "source_page": page, "backfilled": True},
        )
        inv_added = _backfill_categorized(
            merged_diag, investigation_pool, diag_secs,
            classification=inv_classification,
            default_category="laboratory_tests",
            path_for_category_fn=lambda cat: f"investigations.{cat.replace(' ', '_')}",
            item_shape_fn=lambda name, page: name,
        )
        drg_added = _backfill_categorized(
            merged_treat, drug_pool, treat_secs,
            classification=drg_classification,
            default_category="chemotherapy",
            path_for_category_fn=lambda cat: f"{cat.replace(' ', '_')}.backfilled_drugs",
            item_shape_fn=lambda name, page: {"name": name, "source_page": page, "backfilled": True},
        )
        n_inv = sum(inv_added.values())
        n_drg = sum(drg_added.values())
        if n_bm or n_inv or n_drg:
            logger.info(
                f"[Backfill] {d}: +{n_bm} biomarkers, "
                f"+{n_inv} investigations {inv_added}, "
                f"+{n_drg} drugs {drg_added}"
            )

        out[d] = (merged_diag, merged_treat)
    return out


# ═════════════════════════════════════════════════════════════════
# SECTION 15c — SUBTYPE RELEVANCE FILTERING (NEW)
# ═════════════════════════════════════════════════════════════════
#
# Fixes the "copy.deepcopy(diagnosis_knowledge)" pollution bug (Problem: a
# subtype like iMCD-TAFRO inheriting HHV8-specific biomarkers wholesale).
#
# Three corrections vs. the naive version:
#   1) Scoring is TIERED — relationship match > source-page match > token
#      match — instead of token-matching alone (which drops things like
#      "Anasarca" for "iMCD-TAFRO" even though it's a defining feature).
#   2) Fields are discovered DYNAMICALLY by walking the dict for any list
#      of entities, at any depth — instead of a hardcoded field list that
#      breaks for new specialties.
#   3) No new LLM calls anywhere in this section.

_NON_ENTITY_KEYS = {"source_pages", "source_page", "gaps", "specificity_ratio"}


def _is_entity_list(key: str, value) -> bool:
    """True if this looks like a list of clinical entities, not metadata."""
    if key.startswith("_") or key in _NON_ENTITY_KEYS:
        return False
    if not isinstance(value, list) or not value:
        return False
    if all(isinstance(x, (int, float)) for x in value):   # e.g. page number lists
        return False
    return True


def walk_entity_fields(body: dict, path: tuple = ()) -> list[tuple[tuple, list]]:
    """
    Recursively find every list-of-entities field in a diagnosis/treatment
    body at any nesting depth — no hardcoded field names. Works for any
    specialty: biomarkers, risk_groups, complications, imaging_findings,
    histopathology, response_criteria, chemotherapy.regimens, etc.
    Returns [(path_tuple, list_value), ...].
    """
    found: list[tuple[tuple, list]] = []
    if not isinstance(body, dict):
        return found
    for key, value in body.items():
        if key.startswith("_"):
            continue
        full_path = path + (key,)
        if _is_entity_list(key, value):
            found.append((full_path, value))
        elif isinstance(value, dict):
            found.extend(walk_entity_fields(value, full_path))
    return found


def get_nested(body: dict, path: tuple):
    node = body
    for p in path:
        if not isinstance(node, dict):
            return None
        node = node.get(p)
    return node


def set_nested(body: dict, path: tuple, value) -> None:
    node = body
    for p in path[:-1]:
        node = node.setdefault(p, {})
    node[path[-1]] = value


def _entity_display_name(entity) -> str:
    if isinstance(entity, str):
        return entity
    if isinstance(entity, dict):
        return _normalize_keyword(entity) or json.dumps(entity)
    return str(entity)


def _entity_source_page(entity) -> Optional[int]:
    if isinstance(entity, dict):
        return _safe_page_int(entity.get("source_page"))
    return None


def relevance_score(
    entity,
    subtype_name: str,
    parent_disease: str = "",
    subtype_pages: Optional[set[int]] = None,
    relationship_store: Optional["RelationshipStore"] = None,
    sibling_subtypes: Optional[list[str]] = None,   # NEW
) -> int:
    """
    Tiered, rule-based relevance scoring — NO LLM.
    """
    score = 0
    name = _entity_display_name(entity)
    name_lower = name.lower()
    subtype_lower = subtype_name.lower()
    entity_text = json.dumps(entity).lower() if isinstance(entity, dict) else name_lower

    # HARD EXCLUSION (NEW): entity explicitly names a sibling subtype/stage
    # and does NOT name this subtype -> reject outright. This is what was
    # letting e.g. MCD-only stage/regimen text leak into UCD skills, because
    # the old floor below gave it a free point just for mentioning the
    # shared parent disease name.
    if sibling_subtypes:
        for sib in sibling_subtypes:
            sib_l = sib.lower().strip()
            if sib_l and sib_l != subtype_lower and sib_l in entity_text and subtype_lower not in entity_text:
                return 0

    # Tier 1 — relationship graph match
    if relationship_store:
        for rel in relationship_store.all_relationships():
            s, t = rel.source.lower(), rel.target.lower()
            entity_hit = name_lower in s or s in name_lower or name_lower in t or t in name_lower
            subtype_hit = subtype_lower in s or s in subtype_lower or subtype_lower in t or t in subtype_lower
            if entity_hit and subtype_hit:
                score += 20
                break

    # Tier 2 — source-section/page match
    if subtype_pages:
        page = _entity_source_page(entity)
        if page is not None and page in subtype_pages:
            score += 15

    # Tier 3 — token overlap
    if subtype_lower and subtype_lower in name_lower:
        score += 10
    else:
        subtype_tokens = {t for t in re.findall(r"\b\w+\b", subtype_lower) if len(t) > 2}
        score += sum(1 for t in subtype_tokens if t in entity_text) * 4

    # REMOVED the old floor ("mentions parent disease -> +1"). That floor
    # fired on almost every disease-level entity and was the actual root
    # cause of subtype contamination — it is not replaced with anything;
    # entities with zero real signal now correctly score 0.

    return score


def filter_entities_by_relevance(
    entities: list,
    subtype_name: str,
    parent_disease: str = "",
    subtype_pages: Optional[set[int]] = None,
    relationship_store: Optional["RelationshipStore"] = None,
    min_score: int = 4,                              # was 1 — now requires real signal
    sibling_subtypes: Optional[list[str]] = None,     # NEW
) -> list:
    return [
        e for e in entities
        if relevance_score(e, subtype_name, parent_disease, subtype_pages,
                            relationship_store, sibling_subtypes) >= min_score
    ]

SUBTYPE_CANDIDATE_MIN_SCORE = 4  # stricter than backfill's min_score=1 — only show
                                  # plausible candidates, not everything with a floor score

def _format_candidate_entities(candidates: dict[str, list[tuple[tuple, list]]]) -> str:
    lines = []
    for knowledge_type, field_items in candidates.items():
        for path, entities in field_items:
            names = [_entity_display_name(e) for e in entities]
            if names:
                lines.append(f"  [{knowledge_type}.{'.'.join(path)}]: {', '.join(names)}")
    if not lines:
        return ""
    return (
        "CANDIDATE ENTITIES FROM DISEASE-LEVEL KNOWLEDGE (not guaranteed relevant to "
        "THIS subtype — verify each against this subtype's own retrieved sections below "
        "before including; omit anything not actually supported here):\n"
        + "\n".join(lines)
    )

def expand_via_relationships(
    confirmed_entities: list,
    candidate_pool: list,
    relationship_store: Optional["RelationshipStore"],
) -> list:
    """
    One-hop expansion only (no chains) — for each entity already confirmed
    relevant to the subtype, pull in anything from the SAME pool that is
    explicitly linked to it in the relationship graph, even with zero
    name/token/section overlap with the subtype itself.
    """
    if not relationship_store or not confirmed_entities:
        return []
    confirmed_names = {_entity_display_name(e).lower() for e in confirmed_entities}
    pool_by_name = {_entity_display_name(e).lower(): e for e in candidate_pool}
    expanded = []
    for rel in relationship_store.all_relationships():
        s, t = rel.source.lower(), rel.target.lower()
        if s in confirmed_names and t in pool_by_name and t not in confirmed_names:
            expanded.append(pool_by_name[t])
        elif t in confirmed_names and s in pool_by_name and s not in confirmed_names:
            expanded.append(pool_by_name[s])
    return expanded


def build_subtype_candidates(
    subtype: str,
    parent_disease: str,
    subtype_pages: Optional[set[int]],
    disease_diag: Optional[dict],
    disease_treat: Optional[dict],
    relationship_store: Optional["RelationshipStore"],
    sibling_subtypes: Optional[list[str]] = None,   # NEW
) -> dict:
    candidates = {"diagnosis": [], "treatment": []}
    for knowledge_type, body in (("diagnosis", disease_diag), ("treatment", disease_treat)):
        if not body:
            continue
        for path, pool in walk_entity_fields(body):
            filtered = filter_entities_by_relevance(
                pool, subtype, parent_disease, subtype_pages, relationship_store,
                min_score=SUBTYPE_CANDIDATE_MIN_SCORE,
                sibling_subtypes=sibling_subtypes,
            )
            if filtered:
                candidates[knowledge_type].append((path, filtered[:15]))  # cap prompt size
    return candidates


def build_subtype_knowledge_package(
    subtype_name: str,
    parent_disease: str,
    subtype_diag: dict,
    subtype_treat: dict,
    disease_diag: dict,
    disease_treat: dict,
    subtype_pages: Optional[set[int]] = None,
    relationship_store: Optional["RelationshipStore"] = None,
    sibling_subtypes: Optional[list[str]] = None,   # NEW
) -> tuple[dict, dict]:
    """
    Backfill the subtype's own (possibly sparse) LLM-extracted diagnosis/
    ...
    """
    out_diag  = copy.deepcopy(subtype_diag)  if subtype_diag  else _empty_diagnosis()
    out_treat = copy.deepcopy(subtype_treat) if subtype_treat else _empty_treatment()

    for body_out, body_disease in ((out_diag, disease_diag), (out_treat, disease_treat)):
        if not body_disease:
            continue
        for path, pool in walk_entity_fields(body_disease):
            filtered = filter_entities_by_relevance(
                pool, subtype_name, parent_disease, subtype_pages, relationship_store,
                sibling_subtypes=sibling_subtypes,
            )
            expanded = expand_via_relationships(filtered, pool, relationship_store)
            combined_keys = {
                json.dumps(x, sort_keys=True) if isinstance(x, (dict, list)) else str(x)
                for x in filtered
            }
            for item in expanded:
                key = json.dumps(item, sort_keys=True) if isinstance(item, (dict, list)) else str(item)
                if key not in combined_keys:
                    filtered.append(item)
                    combined_keys.add(key)
            if not filtered:
                continue
            existing = get_nested(body_out, path)
            if not isinstance(existing, list):
                # NEW — defensive type coercion. The subtype-level LLM
                # extraction is a separate call from the disease-level one
                # and isn't strictly schema-locked, so it can occasionally
                # return a non-list (e.g. a stray string) at a path where
                # the disease-level body has a list. `or []` alone doesn't
                # catch this because a non-empty string is truthy in Python.
                if existing:
                    logger.warning(
                        f"[SubtypeKnowledgePackage] expected list at "
                        f"{'.'.join(path)} but got {type(existing).__name__} "
                        f"— discarding and starting a fresh list"
                    )
                existing = []
            existing_keys = {
                json.dumps(x, sort_keys=True) if isinstance(x, (dict, list)) else str(x)
                for x in existing
            }
            for item in filtered:
                # NEW (Phase 3b) — tag inherited items so downstream
                # consumers (and doctors reviewing the skill) can see this
                # entity came from the parent disease, not this subtype's
                # own sections. Only dict-shaped items get tagged in place;
                # bare strings are wrapped so the tag has somewhere to live.
                if isinstance(item, dict):
                    tagged_item = {**item, "inherited_from": parent_disease}
                else:
                    tagged_item = {"name": item, "inherited_from": parent_disease}
                key = json.dumps(tagged_item, sort_keys=True)
                if key not in existing_keys:
                    existing.append(tagged_item)
                    existing_keys.add(key)
            set_nested(body_out, path, existing)

    return out_diag, out_treat

def compute_entity_coverage(disease_body: dict, subtype_body: dict) -> dict:
    """
    Gap 5/12: quantifies, per field, how many disease-level entities of that
    shape ended up (post relevance-filter + backfill) in the subtype body.
    This is informational, not a pass/fail gate — a low % can be legitimate
    (most disease entities genuinely don't belong to a narrow subtype).
    """
    input_fields  = walk_entity_fields(disease_body)
    output_lookup = {path: val for path, val in walk_entity_fields(subtype_body)}

    field_report: dict[str, dict] = {}
    total_in, total_out = 0, 0
    for path, pool in input_fields:
        out_list = output_lookup.get(path, [])
        total_in  += len(pool)
        total_out += len(out_list)
        field_report[".".join(path)] = {"input_count": len(pool), "output_count": len(out_list)}

    return {
        "fields": field_report,
        "total_input_entities":  total_in,
        "total_output_entities": total_out,
        "coverage_pct": round((total_out / total_in) * 100, 1) if total_in else None,
    }

# ═════════════════════════════════════════════════════════════════
# SECTION 16 — SUBTYPE SKILL GENERATION  (NEW — batched, no sleep(8))
# ═════════════════════════════════════════════════════════════════

_SUBTYPE_BATCH_SYSTEM = """
You are a clinical knowledge extraction specialist.
You will be given MULTIPLE subtypes, each with its own retrieved guideline sections.
For EACH subtype, extract diagnosis and treatment knowledge using ONLY what is
explicitly stated in THAT subtype's sections. Never mix content between subtypes.
If a subtype shows [NO CONTENT RETRIEVED], omit it from the output entirely.

Each subtype's block includes an EXCLUDE line naming its sibling subtypes.
Retrieval is semantic, so text about a sibling subtype (e.g. "Multicentric
Castleman Disease" content) may appear inside a different subtype's block
(e.g. "Unicentric Castleman Disease") purely because the wording is similar.
Actively check every biomarker, stage, regimen, and drug against the EXCLUDE
list before including it — if a piece of text explicitly names an excluded
sibling and does not also name the current subtype, do not attribute it to
the current subtype.

Only include an entity (biomarker, drug, stage, regimen, criterion) if it is
explicitly present in THIS subtype's retrieved sections. Do not omit an entity
merely because it might also apply at the disease level — subtype specificity
matters more than avoiding repetition. If you are unsure whether something
belongs to this subtype vs. a sibling subtype, exclude it rather than guess;
a separate downstream step will backfill genuinely shared disease-level
entities using relevance filtering, so under-extraction here is safe and
over-extraction (hallucinated or borrowed entities) is not.

Where a "CANDIDATE ENTITIES" list is provided for a subtype, treat it as a
checklist to verify against — not a mandate. Include a candidate only if this
subtype's own retrieved sections actually support it; the downstream pipeline
independently backfills confirmed-relevant disease-level entities, so you do
not need to force-include unverifiable candidates just to be safe.

For each subtype, identify all clinically relevant skills, competencies,
decision-making tasks, diagnostic tasks, treatment selection tasks,
monitoring tasks, interpretation tasks, and management tasks that are
supported by the retrieved subtype content.

Skills may overlap with disease-level skills if the subtype content
provides subtype-specific criteria, biomarkers, treatments, monitoring,
or management guidance.

Do not suppress skills simply because a similar skill may also exist at
the disease level.

Return all clinically relevant skills supported by the subtype content.

If highly subtype-specific skills are not evident, return clinically
relevant diagnosis, treatment, monitoring, interpretation, assessment,
and management skills supported by the subtype content.

Generate 5-10 clinically meaningful skills whenever sufficient subtype information exists.

Do not limit output to rare or highly specialized skills.

CRITICAL REQUIREMENTS

You will receive N subtype names.

You MUST return exactly N subtype objects.

For every subtype provided in the input:

1. Create a separate output object.
2. Preserve the subtype name exactly.
3. Do not rename subtype names.
4. Do not merge subtype names.
5. Do not collapse subtype variants into a parent subtype.
6. Do not omit any subtype.

Even if two subtype names are clinically related, return separate entries for both.

If a subtype has limited information, return an empty diagnosis/treatment structure for that subtype instead of omitting it.

Return ONLY valid JSON:
{
  "subtypes": {
    "<subtype name>": {
      "diagnosis": {
        "diagnostic_criteria": "",
        "biomarkers": [{"name":"","significance":"","source_page":0}],
        "staging": [{"stage":"","criteria":"","source_page":0}],
        "clinical_presentation": {"symptoms": [], "signs": []},
        "subtypes": [],
        "key_evidence": [{"trial":"","finding":"","source_page":0}],
        "skills": [{"skill":"","basis":""}]
      },
      "treatment": {
        "stage_wise_treatment": [{"stage":"","intent":"","primary_treatment":"","options":[{"regimen_name":"","drugs":[],"source_page":0}]}],
        "chemotherapy": {"regimens": [{"name":"","drugs":[],"indication":""}]},
        "targeted_therapy": {"drugs": [], "targets": [], "indications": ""},
        "if_then_rules": [{"condition":"","action":"","source_page":0}],
        "recommendations": [{"statement":"","evidence_level":"","source_page":0}],
        "skills": [{"skill":"","basis":""}]
      }
    }
  }
}
"""


async def derive_subtype_knowledge_batched(
    subtypes: list[str],
    doc_id: str,
    semaphore: asyncio.Semaphore,
    subtype_parent_map: Optional[dict[str, str]] = None,
    biomarkers: Optional[list[str]] = None,
    stages: Optional[list[str]] = None,
    risk_groups: Optional[list[str]] = None,
    disease_diag: Optional[dict] = None,                        # NEW
    disease_treat: Optional[dict] = None,                       # NEW
    relationship_store: Optional["RelationshipStore"] = None,   # NEW
) -> dict[str, dict]:
    """
    Replaces v12's derive_subtype_knowledge() (1 LLM call + sleep(8) PER subtype).
    Step 1: retrieve sections for ALL subtypes in parallel (no LLM, no rate limit risk).
    Step 2: batch subtypes into groups of SUBTYPE_BATCH_SIZE (~20) and issue ONE
            LLM call per batch — so 174 subtypes ≈ 9 calls instead of 174.
    """
    subtype_knowledge: dict[str, dict] = {}
    if not subtypes:
        return subtype_knowledge

    loop = asyncio.get_event_loop()

    biomarker_names = _extract_entity_names(biomarkers or [])
    stage_names = _extract_entity_names(stages or [])
    risk_group_names = _extract_entity_names(risk_groups or [])

    async def _retrieve_subtype(subtype: str, parent_disease: Optional[str]) -> tuple[str, list[Section]]:
        query = build_context_aware_query(
            entity_name=subtype,
            parent_disease=parent_disease,
            related_biomarkers=biomarker_names,
            related_stages=stage_names,
            related_risk_groups=risk_group_names,
        )
        sections = await loop.run_in_executor(None, lambda q=query: search_guideline(q, doc_id))
        return subtype, sections

    logger.info(
        f"[Subtype Retrieval] count={len(subtypes)}"
    )

    logger.info(
        f"[Subtype Retrieval Sample] {subtypes[:10]}"
    )

    retrieval_results = await asyncio.gather(
        *[_retrieve_subtype(s, subtype_parent_map.get(s)) for s in subtypes]
    )
    subtype_sections_map = {s: secs for s, secs in retrieval_results}

    # Retry weak/zero-retrieval subtypes with a parent-disease-scoped query
    # before giving up on them — generic, no hardcoded thresholds beyond
    # "did retrieval return anything at all".
    subtype_parent_map = subtype_parent_map or {}
    weak = [s for s, secs in subtype_sections_map.items() if len(secs) == 0]
    if weak:
        async def _retry_subtype(subtype: str) -> tuple[str, list[Section]]:
            parent = subtype_parent_map.get(subtype)
            retry_query = build_context_aware_query(
                entity_name=subtype,
                parent_disease=parent,
                related_biomarkers=biomarker_names,
                related_stages=stage_names,
                related_risk_groups=risk_group_names,
            )
            secs = await loop.run_in_executor(None, lambda q=retry_query: search_guideline(q, doc_id))
            return subtype, secs

        retry_results = await asyncio.gather(*[_retry_subtype(s) for s in weak])
        recovered = 0
        for subtype, secs in retry_results:
            if secs:
                subtype_sections_map[subtype] = secs
                recovered += 1
        logger.info(f"[SubtypeRetry] {recovered}/{len(weak)} weak subtypes recovered via parent-scoped retry")

    subtype_candidates_map: dict[str, dict] = {}
    if disease_diag or disease_treat:
        for subtype in subtypes:
            secs = subtype_sections_map.get(subtype, [])
            pages = {p for s in secs for p in range(s.page_start, s.page_end + 1)}
            parent = (subtype_parent_map or {}).get(subtype, "")
            sibling_subtypes = [s for s in subtypes if s != subtype]
            subtype_candidates_map[subtype] = build_subtype_candidates(
                subtype, parent, pages or None, disease_diag, disease_treat, relationship_store,
                sibling_subtypes=sibling_subtypes,
            )

    batches = [subtypes[i:i + SUBTYPE_BATCH_SIZE] for i in range(0, len(subtypes), SUBTYPE_BATCH_SIZE)]
    logger.info(f"[SubtypeBatch] {len(subtypes)} subtypes → {len(batches)} batch(es) of ≤{SUBTYPE_BATCH_SIZE}")

    for batch_idx, batch in enumerate(batches):
        parts, any_content = [], False
        logger.info(f"[SubtypeBatch {batch_idx}] input={len(batch)}")
        expected_keys = "\n".join(batch)
        parts.append(
            f"Input subtype count: {len(batch)}\n\n"
            f"Subtype names:\n" + "\n".join(f"- {s}" for s in batch) + "\n\n"
            f"Return EXACTLY {len(batch)} subtype objects.\n\n"
            f"Requirements:\n"
            f"1. Every subtype listed above must appear exactly once.\n"
            f"2. Use the exact subtype names as keys.\n"
            f"3. Do not merge subtypes.\n"
            f"4. Do not omit subtypes.\n"
            f"5. Do not create new subtypes.\n"
            f"6. The number of returned subtype objects must equal {len(batch)}.\n\n"
            f"The keys under \"subtypes\" must be exactly:\n{expected_keys}\n"
            f"and no others."
        )
        for subtype in batch:
            secs = subtype_sections_map.get(subtype, [])
            if not secs:
                parts.append(f"=== SUBTYPE: {subtype} ===\n[NO CONTENT RETRIEVED]")
                continue
            any_content = True
            sec_text = "\n\n".join(
                f"[page {s.page_start}-{s.page_end}] {s.section_title}\n{s.text[:3000]}" for s in secs
            )
            cand_text = _format_candidate_entities(subtype_candidates_map.get(subtype, {}))

            # NEW (Phase 1a) — explicit exclusion list, computed from the FULL
            # subtype universe (not just this batch), so cross-batch siblings
            # are excluded too. This is what lets the LLM actively filter out
            # sibling-subtype text that ChromaDB's semantic retrieval pulled
            # in because it scored similarly (e.g. all Castleman subtypes
            # share heavy vocabulary overlap).
            all_siblings = [s for s in subtypes if s != subtype]
            exclude_line = (
                f"\nEXCLUDE information specific to these OTHER subtypes, even if it "
                f"appears in the retrieved text above: {', '.join(all_siblings[:25])}"
                if all_siblings else ""
            )

            block = f"=== SUBTYPE: {subtype} ==={exclude_line}\n{sec_text}"
            if cand_text:
                block += f"\n\n{cand_text}"
            parts.append(block)

        if not any_content:
            for subtype in batch:
                subtype_knowledge[subtype] = {
                    "diagnosis": None, "treatment": None, "specificity_score": 0.0,
                    "is_generic": True, "used_targeted_extraction": False, "rejected": True,
                    "rejection_reason": f"No content retrieved for subtype '{subtype}'.",
                    "skills": [],
                }
            continue

        try:
            async with semaphore:
                result = await _chat_async(_SUBTYPE_BATCH_SYSTEM, "\n\n".join(parts))
        except Exception as exc:
            logger.warning(f"[SubtypeBatch {batch_idx}] failed: {exc}")
            result = {}

        batch_subtypes = result.get("subtypes", {})

        logger.info(
            f"[SubtypeBatch {batch_idx}] requested={len(batch)} returned={len(batch_subtypes)}"
        )

        logger.info(
            f"[SubtypeBatch {batch_idx}] missing={set(batch)-set(batch_subtypes.keys())}"
        )
        logger.info(
            f"[SubtypeBatch {batch_idx}] returned_subtypes={list(batch_subtypes.keys())}"
        )

        logger.info(
            f"[SubtypeBatch {batch_idx}] requested_subtypes={batch}"
        )
        logger.info(f"[SubtypeBatch {batch_idx}] llm_output_keys={len(batch_subtypes)}")
        for subtype in batch:
            secs  = subtype_sections_map.get(subtype, [])
            pages = sorted({p for s in secs for p in range(s.page_start, s.page_end + 1)})
            data  = batch_subtypes.get(subtype, {})
            diag_body  = data.get("diagnosis") or {}
            treat_body = data.get("treatment") or {}

            if not diag_body and not treat_body:
                subtype_knowledge[subtype] = {
                    "diagnosis": None, "treatment": None, "specificity_score": 0.0,
                    "is_generic": True, "used_targeted_extraction": False, "rejected": True,
                    "rejection_reason": f"No extractable content for subtype '{subtype}'.",
                    "skills": [],
                }
                continue

            # Pull skills out BEFORE merging/sanitising — these helpers don't
            # know about the "skills" key and would silently drop it.
            diag_skills_raw  = diag_body.pop("skills", []) if isinstance(diag_body, dict) else []
            treat_skills_raw = treat_body.pop("skills", []) if isinstance(treat_body, dict) else []

            if diag_body:
                diag_body["_source_pages"] = pages
            if treat_body:
                treat_body["_source_pages"] = pages

            merged_diag  = _merge_diagnosis([_sanitise_raw_llm_body(diag_body)]) if diag_body else _empty_diagnosis()
            merged_treat = _merge_treatment([_sanitise_raw_llm_body(treat_body)]) if treat_body else _empty_treatment()

            n_reclass = _reclassify_investigations(merged_diag, merged_treat)
            if n_reclass:
                logger.info(f"[Reclassify] {subtype}: moved {n_reclass} items → therapeutic_procedures")

            subtype_skills = [s for s in (diag_skills_raw + treat_skills_raw) if isinstance(s, dict) and s.get("skill")]

            # ── NEW: backfill relevant disease-level entities the subtype
            #    extraction missed, filtered through the tiered scorer.
            parent = (subtype_parent_map or {}).get(subtype, "")
            subtype_page_set = set(pages) if pages else None
            if disease_diag or disease_treat:
                merged_diag, merged_treat = build_subtype_knowledge_package(
                    subtype_name=subtype,
                    parent_disease=parent,
                    subtype_diag=merged_diag,
                    subtype_treat=merged_treat,
                    disease_diag=disease_diag or {},
                    disease_treat=disease_treat or {},
                    subtype_pages=subtype_page_set,
                    relationship_store=relationship_store,
                    sibling_subtypes=[s for s in subtypes if s != subtype],
                )

            entity_coverage = None
            if disease_diag or disease_treat:
                entity_coverage = {
                    "diagnosis": compute_entity_coverage(disease_diag or {}, merged_diag),
                    "treatment": compute_entity_coverage(disease_treat or {}, merged_treat),
                }
                logger.info(
                    f"[SubtypeEntityCoverage] {subtype}: "
                    f"diag={entity_coverage['diagnosis']['coverage_pct']}% "
                    f"({entity_coverage['diagnosis']['total_output_entities']}/{entity_coverage['diagnosis']['total_input_entities']}) "
                    f"treat={entity_coverage['treatment']['coverage_pct']}% "
                    f"({entity_coverage['treatment']['total_output_entities']}/{entity_coverage['treatment']['total_input_entities']})"
                )

            logger.info(
                f"[SubtypeCoverage] {subtype}: "
                f"biomarkers={len(merged_diag.get('biomarkers', []))} "
                f"staging={len(merged_diag.get('staging', []))} "
                f"regimens={len((merged_treat.get('chemotherapy') or {}).get('regimens', []))} "
                f"if_then_rules={len(merged_treat.get('if_then_rules', []))} "
                f"recommendations={len(merged_treat.get('recommendations', []))}"
            )

            subtype_knowledge[subtype] = {
                "diagnosis": merged_diag,
                "treatment": merged_treat,
                "specificity_score": 1.0, "is_generic": False, "used_targeted_extraction": True,
                "skills": subtype_skills,
                "entity_coverage": entity_coverage,   # NEW — exposes Gap 5/12 metrics for QA/UI
            }

        logger.info(f"[SubtypeBatch {batch_idx+1}/{len(batches)}] {len(batch)} subtypes processed")

    return subtype_knowledge

# ═════════════════════════════════════════════════════════════════
# SECTION 16b — STRUCTURAL CAPABILITY EXTRACTION  (NEW in v14 — no LLM)
# ═════════════════════════════════════════════════════════════════
#
# WHY THIS EXISTS:
#   In v13, skill count depended entirely on the LLM remembering to list
#   "skills" inside its diagnosis/treatment JSON response. That meant a
#   disease with 15 biomarkers, 8 stages, and 6 regimens could still end up
#   with 0-1 skills if the LLM simply didn't enumerate them as "skills".
#
#   This section fixes that WITHOUT any new LLM calls, retrieval calls, or
#   specialty-specific logic. It looks at the SHAPE of knowledge that was
#   already extracted (diagnosis dict / treatment dict) and, for each
#   structural field that actually has content, emits a corresponding skill.
#   This is the "Capability Extraction Engine" — every rule below answers
#   "does this field have content?", never "what disease/specialty is this?".
#
#   Because the rules are keyed on FIELD SHAPE (biomarkers present? staging
#   present? chemotherapy.regimens present?) rather than on disease names or
#   keyword lists, this generalizes to any specialty (oncology, cardiology,
#   nephrology, neurology, pulmonology, endocrinology, gastroenterology,
#   pediatrics, etc.) without modification.

def _has_items(value: Any) -> bool:
    """Generic non-empty check across str / list / dict shapes."""
    if value is None:
        return False
    if isinstance(value, str):
        return bool(value.strip())
    if isinstance(value, (list, tuple, set)):
        return len(value) > 0
    if isinstance(value, dict):
        return any(_has_items(v) for v in value.values())
    return bool(value)


# Structural rules: (dotted path into the body dict, skill name template, basis label)
# Dotted path may traverse nested dicts, e.g. "chemotherapy.regimens".
_DIAGNOSIS_SKILL_RULES: list[tuple[str, str, str]] = [
    ("biomarkers",             "Biomarker Interpretation - {entity}",                  "biomarkers"),
    ("staging",                "Disease Staging - {entity}",                           "staging"),
    ("diagnostic_criteria",    "Diagnostic Assessment - {entity}",                     "diagnostic_criteria"),
    ("clinical_presentation",  "Clinical Evaluation - {entity}",                       "clinical_presentation"),
    ("investigations",         "Investigation Selection - {entity}",                   "investigations"),
    ("molecular_testing",      "Molecular Testing Interpretation - {entity}",          "molecular_testing"),
    ("risk_stratification",    "Risk Stratification - {entity}",                       "risk_stratification"),
    ("diagnostic_pathway",     "Diagnostic Pathway Navigation - {entity}",             "diagnostic_pathway"),
    ("differential_diagnosis", "Differential Diagnosis - {entity}",                    "differential_diagnosis"),
    ("exclusion_criteria",     "Diagnostic Exclusion Assessment - {entity}",           "exclusion_criteria"),
    ("special_populations",    "Special Population Diagnostic Assessment - {entity}",  "special_populations"),
    ("subtypes",               "Subtype Classification - {entity}",                    "subtypes"),
]

_TREATMENT_SKILL_RULES: list[tuple[str, str, str]] = [
    ("chemotherapy.regimens",  "Chemotherapy Regimen Selection - {entity}",            "chemotherapy"),
    ("targeted_therapy.drugs", "Targeted Therapy Selection - {entity}",                "targeted_therapy"),
    ("immunotherapy.drugs",    "Immunotherapy Selection - {entity}",                   "immunotherapy"),
    ("hormone_therapy.drugs",  "Hormone Therapy Selection - {entity}",                 "hormone_therapy"),
    ("surgery.procedures",     "Surgical Planning - {entity}",                         "surgery"),
    ("radiation.protocols",    "Radiation Treatment Planning - {entity}",              "radiation"),
    ("stage_wise_treatment",   "Stage-Wise Treatment Planning - {entity}",             "stage_wise_treatment"),
    ("if_then_rules",          "Clinical Decision Making - {entity}",                  "if_then_rules"),
    ("recommendations",        "Guideline-Based Management - {entity}",               "recommendations"),
    ("contraindications",      "Contraindication Screening - {entity}",               "contraindications"),
    ("dose_modifications",     "Dose Modification Management - {entity}",             "dose_modifications"),
    ("monitoring",             "Treatment Monitoring - {entity}",                      "monitoring"),
    ("follow_up",              "Follow-Up Planning - {entity}",                        "follow_up"),
    ("supportive_care",        "Supportive Care Management - {entity}",               "supportive_care"),
    ("special_populations",    "Special Population Treatment Adaptation - {entity}",   "special_populations"),
    ("toxicity_monitoring",    "Toxicity Monitoring - {entity}",                       "toxicity_monitoring"),
    ("response_assessment",    "Treatment Response Assessment - {entity}",             "response_assessment"),
    ("surveillance_schedule",  "Surveillance Planning - {entity}",                     "surveillance_schedule"),
]


def _resolve_dotted(body: dict, dotted_path: str):
    """Resolve a dotted path like 'chemotherapy.regimens' inside a body dict."""
    node = body
    for part in dotted_path.split("."):
        if not isinstance(node, dict):
            return None
        node = node.get(part)
    return node


def _derive_skills_from_body(
    entity_name: str,
    body: Optional[dict],
    rules: list[tuple[str, str, str]],
    entity_type: str,
    knowledge_type: str,
) -> list[dict]:
    if not body or not isinstance(body, dict):
        return []
    derived: list[dict] = []
    for dotted_path, name_template, basis in rules:
        value = _resolve_dotted(body, dotted_path)
        if _has_items(value):
            derived.append({
                "skill":          name_template.format(entity=entity_name),
                "basis":          basis,
                "source":         entity_name,
                "entity_type":    entity_type,
                "knowledge_type": knowledge_type,
                "derivation":     "structural",
            })
    return derived


def derive_skills_from_knowledge(
    entity_name: str,
    diagnosis: Optional[dict],
    treatment: Optional[dict],
    entity_type: str = "disease",
) -> list[dict]:
    """
    Structurally derive clinical skills from already-extracted diagnosis/
    treatment knowledge for a single disease or subtype — ZERO additional
    LLM calls, ZERO additional retrieval calls. A skill is created whenever
    the corresponding structural field actually has content, so skill
    coverage tracks knowledge richness instead of depending on the LLM to
    separately enumerate "skills" in its response.

    Generalizes across specialties because the rules key off field SHAPE
    (biomarkers present? staging present? regimens present?) — never off
    disease names, keyword lists, or specialty branching.
    """
    skills: list[dict] = []
    skills.extend(_derive_skills_from_body(entity_name, diagnosis, _DIAGNOSIS_SKILL_RULES, entity_type, "diagnosis"))
    skills.extend(_derive_skills_from_body(entity_name, treatment, _TREATMENT_SKILL_RULES, entity_type, "treatment"))
    return skills


# ═════════════════════════════════════════════════════════════════
# SECTION 16c — SKILL AGGREGATION  (UPDATED in v14 — merges LLM-declared
#                skills + structurally-derived skills, no extra LLM calls)
# ═════════════════════════════════════════════════════════════════

_STRUCTURAL_BASIS_SCORES: dict[str, int] = {
    "biomarkers": 2, "molecular_testing": 2, "staging": 2, "risk_stratification": 2,
    "if_then_rules": 3, "chemotherapy": 2, "targeted_therapy": 2, "immunotherapy": 2,
    "hormone_therapy": 2, "recommendations": 2, "contraindications": 2,
    "dose_modifications": 2, "toxicity_monitoring": 2, "stage_wise_treatment": 2,
    "surgery": 1, "radiation": 1, "monitoring": 1, "follow_up": 1,
    "supportive_care": 1, "diagnostic_criteria": 1, "clinical_presentation": 1,
    "investigations": 1, "diagnostic_pathway": 1, "differential_diagnosis": 1,
    "exclusion_criteria": 1, "special_populations": 1, "subtypes": 1,
    "response_assessment": 1, "surveillance_schedule": 1,
}


def _score_skill(skill: dict) -> int:
    basis = (skill.get("basis") or "")
    if skill.get("derivation") == "structural":
        score = _STRUCTURAL_BASIS_SCORES.get(basis, 1)
    else:
        text = f"{skill.get('skill','')} {basis}".lower()
        score = 0
        if any(k in text for k in ("biomarker", "marker", "mutation", "expression")):
            score += 2
        if any(k in text for k in ("drug", "regimen", "chemotherapy", "therapy", "agent")):
            score += 2
        if any(k in text for k in ("recommend", "guideline")):
            score += 2
        if "rule" in text or "criteria" in text or ("if " in text and "then" in text):
            score += 3
        if "stage" in text:
            score += 2
        score = score or 1

    if skill.get("entity_type") == "subtype":
        score += 3
    return score

def collect_skills(
    per_disease_knowledge: dict[str, tuple[dict, dict]],
    subtype_knowledge: dict[str, dict],
) -> list[dict]:
    """
    Aggregate skills from TWO sources, with ZERO extra LLM calls:

      1) LLM-declared skills — already embedded in the disease-batch and
         subtype-batch LLM responses (unchanged from v13, kept as-is).
      2) Structurally-derived skills (NEW in v14) — computed directly off
         the SHAPE of the extracted diagnosis/treatment knowledge via
         derive_skills_from_knowledge(). This is what guarantees skill
         count actually scales with knowledge richness (diseases, subtypes,
         biomarkers, stages, regimens, rules, etc.) instead of depending on
         the LLM to separately remember to list "skills".

    Both sources are merged BEFORE the existing embedding-based semantic
    dedup, so duplicate/near-duplicate skills (LLM-declared vs. structurally
    derived, or across diseases/subtypes) collapse into one entry exactly as
    before. Generic across specialties — no keyword/string heuristics.
    """
    raw_skills: list[dict] = []
    llm_skill_count = 0
    structural_skill_count = 0

    for disease, (diag, treat) in per_disease_knowledge.items():
        diag  = diag or {}
        treat = treat or {}

        # 1) LLM-declared skills — unchanged from v13
        for s in (diag.get("skills") or []) + (treat.get("skills") or []):
            name = s.get("skill") if isinstance(s, dict) else str(s)
            if name and name.strip():
                raw_skills.append({
                    "skill": name.strip(),
                    "basis": s.get("basis", "") if isinstance(s, dict) else "",
                    "source": disease,
                    "entity_type": "disease",
                    "derivation": "llm",
                })
                llm_skill_count += 1

        # 2) Structurally-derived skills — NEW in v14444
        structural = derive_skills_from_knowledge(disease, diag, treat, entity_type="disease")
        raw_skills.extend(structural)
        structural_skill_count += len(structural)

    for subtype, data in subtype_knowledge.items():
        if data.get("rejected"):
            continue
        sub_diag  = data.get("diagnosis") or {}
        sub_treat = data.get("treatment") or {}

        # 1) LLM-declared skills — unchanged from v13
        for s in data.get("skills") or []:
            name = s.get("skill") if isinstance(s, dict) else str(s)
            if name and name.strip():
                raw_skills.append({
                    "skill": name.strip(),
                    "basis": s.get("basis", "") if isinstance(s, dict) else "",
                    "source": subtype,
                    "entity_type": "subtype",
                    "derivation": "llm",
                })
                llm_skill_count += 1

        # 2) Structurally-derived skills — NEW in v14
        structural = derive_skills_from_knowledge(subtype, sub_diag, sub_treat, entity_type="subtype")
        raw_skills.extend(structural)
        structural_skill_count += len(structural)

    if not raw_skills:
        return []

    items_for_dedup = [{"name": s["skill"], **s} for s in raw_skills]
    deduped = _semantic_dedup(
        items_for_dedup,
        threshold=SKILL_DEDUP_THRESHOLD,
        text_fn=lambda it: f"{it.get('skill','')} — {it.get('basis','')}".strip(),
    )

    final = [{"skill": d["skill"], "basis": d.get("basis", ""),
              "source": d["source"], "entity_type": d["entity_type"],
              "derivation": d.get("derivation", "structural")} for d in deduped]

    for s in final:
        s["quality_score"] = _score_skill(s)
    final.sort(key=lambda s: s["quality_score"], reverse=True)

    if SKILL_MIN_QUALITY_SCORE > 0:
        before = len(final)
        final = [s for s in final if s["quality_score"] >= SKILL_MIN_QUALITY_SCORE]
        logger.info(f"[SkillQuality] dropped {before - len(final)} skills below score={SKILL_MIN_QUALITY_SCORE}")

    logger.info(
        f"[SkillAggregation] raw={len(raw_skills)} "
        f"(llm={llm_skill_count} structural={structural_skill_count}) "
        f"deduped={len(final)} avg_score={round(sum(s['quality_score'] for s in final)/max(len(final),1), 2)}"
    )
    return final

# ═════════════════════════════════════════════════════════════════
# SECTION 17 — SYNTHESIS  (kept — straight LLM consolidation, no validation gate)
# ═════════════════════════════════════════════════════════════════

_SYNTH_DIAG_SYSTEM = """
You are a senior clinical knowledge engineer.
Your job is to CONSOLIDATE — not to generate. Remove duplicates, resolve minor wording conflicts.
⚠️  SOURCE-GROUNDING CONSTRAINT: ONLY consolidate what is already present. Do NOT add facts.
    If a field is empty, leave it empty. Never fill from prior medical knowledge.
Return ONLY valid JSON using the EXACT same schema as the input.
"""

_SYNTH_TREAT_SYSTEM = """
You are a senior clinical knowledge engineer.
Your job is to CONSOLIDATE — not to generate. Remove duplicate regimen entries, merge duplicate items.
⚠️  SOURCE-GROUNDING CONSTRAINT: ONLY consolidate what is already present. Do NOT add drugs or regimens.
    if_then_rules: preserve ALL conditional logic entries exactly as given.
Return ONLY valid JSON using the EXACT same schema as the input.
"""


def _consolidate(merged: dict, system_prompt: str, label: str) -> dict:
    try:
        payload = json.dumps(merged, ensure_ascii=False)
        if len(payload) > 24_000:
            slim    = {k: v for k, v in merged.items() if k != "source_pages"}
            payload = json.dumps(slim, ensure_ascii=False)[:24_000]
        grounded_user = (
            "REMINDER: Only consolidate what is in this draft. "
            "Leave empty fields empty. Do NOT reduce the number of items in any list.\n\n"
            f"Merged draft:\n{payload}"
        )
        result = _chat(system_prompt, grounded_user)
        result["source_pages"] = _clean_pages(merged.get("source_pages") or [])
        logger.info(f"Synthesis complete: {label}")
        return result
    except Exception as exc:
        logger.warning(f"Synthesis failed for {label} ({exc}) — using raw merge")
        return merged


# ═════════════════════════════════════════════════════════════════
# SECTION 18 — KNOWLEDGE GRAPH  (kept — required for graph indexing & search)
# ═════════════════════════════════════════════════════════════════

def _make_index(prefix: str, counter: dict) -> str:
    key = prefix.upper()
    counter[key] = counter.get(key, 0) + 1
    return f"{key}{counter[key]:03d}"


def build_knowledge_graph(
    doctor_id:            str,
    doc_id:               str,
    understanding:        dict,
    diagnosis_knowledge:  dict,
    treatment_knowledge:  dict,
    subtype_knowledge:    dict[str, dict],
    relationship_store:   Optional[RelationshipStore] = None,
) -> dict:
    counter:   dict       = {}
    nodes:     list[dict] = []
    edges:     list[dict] = []
    index_map: dict       = {}

    def add_node(node_type: str, name: str, prefix: str, extra: dict = {}) -> str:
        if isinstance(name, dict):
            name = name.get("name") or name.get("id") or name.get("test") or str(name)
        if isinstance(name, list):
            name = str(name[0]) if name else "unknown"
        if not isinstance(name, str):
            name = str(name)
        if name in index_map:
            return index_map[name]
        idx = _make_index(prefix, counter)
        index_map[name] = idx
        nodes.append({
            "index": idx, "type": node_type, "name": name,
            "doctor_id": doctor_id, "doc_id": doc_id,
            "created_at": datetime.now(timezone.utc).isoformat(),
            **extra,
        })
        return idx

    def add_edge(from_idx: str, to_idx: str, rel: str, extra: dict = {}):
        edges.append({"from": from_idx, "to": to_idx, "relationship": rel,
                       "doc_id": doc_id, "doctor_id": doctor_id, **extra})

    doc_idx = add_node("doctor", doctor_id, "DR")

    disease_type = understanding.get("disease_type") or (
        (understanding.get("diseases") or ["Unknown"])[0]
    )
    dis_idx = add_node("disease", disease_type, "DIS")
    add_edge(doc_idx, dis_idx, "MANAGES")

    for _other_disease in understanding.get("diseases", []):
        if not _other_disease or _other_disease == disease_type: continue
        _od_idx = add_node("disease", _other_disease, "DIS")
        add_edge(doc_idx, _od_idx, "MANAGES")
        add_edge(dis_idx, _od_idx, "RELATED_DISEASE", extra={"relationship_type": "co_guideline"})

    for bm_item in diagnosis_knowledge.get("biomarkers", []):
        bm_name = bm_item["name"] if isinstance(bm_item, dict) else str(bm_item)
        if bm_name:
            bm_idx = add_node("biomarker", bm_name, "BM", extra={
                "significance": bm_item.get("significance", "") if isinstance(bm_item, dict) else "",
                "source_page":  bm_item.get("source_page", 0) if isinstance(bm_item, dict) else 0,
            })
            add_edge(dis_idx, bm_idx, "HAS_BIOMARKER")

    cp = diagnosis_knowledge.get("clinical_presentation", {})
    for symptom in cp.get("symptoms", []):
        if symptom:
            sym_idx = add_node("symptom", f"SYM_{symptom}", "SYM", extra={"presentation_type": "symptom"})
            add_edge(dis_idx, sym_idx, "PRESENTS_WITH")
    for sign in cp.get("signs", []):
        if sign:
            sgn_idx = add_node("symptom", f"SIGN_{sign}", "SYM", extra={"presentation_type": "sign"})
            add_edge(dis_idx, sgn_idx, "PRESENTS_WITH")

    inv = diagnosis_knowledge.get("investigations", {})
    if isinstance(inv, dict):
        for inv_type, inv_list in inv.items():
            for test in (inv_list or []):
                test_name = test if isinstance(test, str) else (test.get("test") or test.get("name") or str(test))
                if test_name:
                    inv_idx = add_node("investigation", test_name, "INV", extra={"investigation_type": inv_type})
                    add_edge(dis_idx, inv_idx, "REQUIRES_INVESTIGATION")

    mon = treatment_knowledge.get("monitoring", {})
    if isinstance(mon, dict):
        for param in (mon.get("parameters") or []):
            if param:
                mon_idx = add_node("monitoring", str(param), "MON")
                add_edge(dis_idx, mon_idx, "REQUIRES_MONITORING")
        for rf in (mon.get("red_flags") or []):
            if rf:
                comp_idx = add_node("complication", str(rf), "COMP")
                add_edge(dis_idx, comp_idx, "HAS_COMPLICATION")

    for ci in treatment_knowledge.get("contraindications", []):
        if not isinstance(ci, dict): continue
        doa = ci.get("drug_or_action", "")
        if doa:
            ci_idx = add_node("contraindication", f"CI_{doa}_{ci.get('condition', '')}", "CI", extra={
                "condition": ci.get("condition", ""), "reason": ci.get("reason", ""),
                "source_page": ci.get("source_page", 0),
            })
            add_edge(dis_idx, ci_idx, "HAS_CONTRAINDICATION")

    for rec in treatment_knowledge.get("recommendations", []):
        if not isinstance(rec, dict): continue
        stmt = rec.get("statement", "")
        if stmt:
            rec_idx = add_node("guideline_recommendation", stmt[:80], "GREC", extra={
                "evidence_level": rec.get("evidence_level", "") or rec.get("grade", ""),
                "strength": rec.get("strength", ""), "conditional": rec.get("conditional", False),
                "source_page": rec.get("source_page", 0),
            })
            add_edge(dis_idx, rec_idx, "HAS_RECOMMENDATION")

    prev_step_idx = None
    for step_num, step in enumerate(diagnosis_knowledge.get("diagnostic_pathway", []), 1):
        if not step: continue
        dp_idx = add_node("diagnostic_pathway", f"PATHWAY_STEP_{step_num}_{str(step)[:40]}", "DP",
                          extra={"step_number": step_num, "description": str(step)})
        add_edge(dis_idx, dp_idx, "HAS_PATHWAY_STEP")
        if prev_step_idx:
            add_edge(prev_step_idx, dp_idx, "FOLLOWED_BY")
        prev_step_idx = dp_idx

    for rule_num, rule in enumerate(treatment_knowledge.get("if_then_rules", []), 1):
        if not isinstance(rule, dict): continue
        condition = rule.get("condition", "")
        action    = rule.get("action", "") or rule.get("recommended_action", "")
        if condition and action:
            rule_idx = add_node("if_then_rule", f"RULE_{rule_num}_{condition[:40]}", "RULE", extra={
                "condition": condition, "action": action,
                "biomarker_condition": rule.get("biomarker_condition", ""),
                "population": rule.get("population", ""), "source_page": rule.get("source_page", 0),
            })
            add_edge(dis_idx, rule_idx, "HAS_CONDITIONAL_RULE", extra={"condition": condition})

    for dm in treatment_knowledge.get("dose_modifications", []):
        if not isinstance(dm, dict): continue
        cond = dm.get("condition", ""); drug = dm.get("drug", "")
        if cond and drug:
            dm_idx = add_node("dose_modification", f"DM_{drug}_{cond[:40]}", "DM",
                              extra={"drug": drug, "condition": cond, "modification": dm.get("modification", "")})
            add_edge(dis_idx, dm_idx, "HAS_DOSE_MODIFICATION")

    _ev_registry: dict[str, str] = {}
    for ev in treatment_knowledge.get("key_evidence", []):
        trial = ev.get("trial", "") if isinstance(ev, dict) else ""
        if not trial: continue
        ev_idx = add_node("evidence", trial, "EV", extra={
            "finding": ev.get("finding", ""), "source_section": ev.get("source_section", ""),
            "source_page": ev.get("source_page", 0),
            "evidence_level": ev.get("recommendation_grade", "") or ev.get("evidence_level", ""),
        })
        add_edge(dis_idx, ev_idx, "SUPPORTED_BY")
        _ev_registry[trial] = ev_idx

    if relationship_store:
        for rel in relationship_store.all_relationships():
            src_idx = add_node("clinical_entity", rel.source, "ENT")
            tgt_idx = add_node("clinical_entity", rel.target, "ENT")
            add_edge(src_idx, tgt_idx, rel.relation.upper(), extra={
                "source_page": rel.source_page, "confidence": rel.confidence, "section": rel.section,
            })
        logger.info(f"[Graph] Added {len(relationship_store.all_relationships())} explicit relationship edges")

    fu = treatment_knowledge.get("follow_up", {})
    if isinstance(fu, dict):
        for test in (fu.get("monitoring_tests") or []):
            if test:
                fu_idx = add_node("follow_up", str(test), "FU")
                add_edge(dis_idx, fu_idx, "REQUIRES_FOLLOWUP")

    def _wire_stage_treatment(swt_item: dict, parent_idx: str, subtype_label: str):
        stage_label = swt_item.get("stage", "")
        if not stage_label: return
        intent = swt_item.get("intent", "")
        stage_key = f"{subtype_label}_{stage_label}" if subtype_label else stage_label
        stg_idx = add_node("stage", stage_key, "STG", extra={
            "stage_label": stage_label, "subtype": subtype_label, "intent": intent,
            "primary_treatment": swt_item.get("primary_treatment", ""),
            "source_pages": swt_item.get("source_pages", []), "conditions": swt_item.get("conditions", []),
        })
        add_edge(parent_idx, stg_idx, "HAS_STAGE", extra={"subtype": subtype_label, "intent": intent})
        for opt in (swt_item.get("options") or []):
            if not isinstance(opt, dict): continue
            reg_name = opt.get("regimen_name", "")
            if not reg_name: continue
            reg_key = f"{subtype_label}_{stage_label}_{reg_name}"
            reg_idx = add_node("regimen", reg_key, "REG", extra={
                "regimen_name": reg_name, "subtype": subtype_label, "stage": stage_label, "intent": intent,
                "modality": opt.get("modality", ""), "line": opt.get("line", ""), "notes": opt.get("notes", ""),
                "drugs": opt.get("drugs", []), "evidence_trial": opt.get("evidence_trial", ""),
                "condition": opt.get("condition", ""), "source_page": opt.get("source_page", 0),
            })
            add_edge(stg_idx, reg_idx, "TREATED_BY", extra={"line": opt.get("line", ""), "modality": opt.get("modality", "")})
            for drug in (opt.get("drugs") or []):
                if drug:
                    drg_idx = add_node("drug", drug, "DRG")
                    add_edge(reg_idx, drg_idx, "CONTAINS_DRUG", extra={"stage": stage_label, "subtype": subtype_label})
            trial = opt.get("evidence_trial", "")
            if trial:
                if trial not in _ev_registry:
                    _ev_registry[trial] = add_node("evidence", trial, "EV", extra={"finding": ""})
                add_edge(reg_idx, _ev_registry[trial], "SUPPORTED_BY")
        surg = swt_item.get("surgery", {})
        if isinstance(surg, dict) and surg.get("procedure"):
            pr_idx = add_node("procedure", f"SURG_{subtype_label}_{stage_label}_{surg['procedure'][:40]}", "PROC",
                              extra={"procedure_type": "surgery", "subtype": subtype_label, "stage": stage_label})
            add_edge(stg_idx, pr_idx, "HAS_PROCEDURE", extra={"procedure_type": "surgery"})
        rad = swt_item.get("radiation", {})
        if isinstance(rad, dict) and rad.get("protocol"):
            pr_idx = add_node("procedure", f"RAD_{subtype_label}_{stage_label}_{rad['protocol'][:40]}", "PROC",
                              extra={"procedure_type": "radiation", "subtype": subtype_label, "stage": stage_label})
            add_edge(stg_idx, pr_idx, "HAS_PROCEDURE", extra={"procedure_type": "radiation"})

    for subtype in understanding.get("subtypes", []):
        if isinstance(subtype, dict):
            subtype_name = subtype.get("name", "")
            parent_names = subtype.get("parent_diseases") or []
        else:
            subtype_name = str(subtype)
            parent_names = []
        if not subtype_name:
            continue
        sub_idx = add_node("subtype", subtype_name, "SUB")
        if parent_names:
            for parent_name in parent_names:
                parent_idx = add_node("disease", parent_name, "DIS")
                add_edge(parent_idx, sub_idx, "HAS_SUBTYPE")
        else:
            add_edge(dis_idx, sub_idx, "HAS_SUBTYPE")  # fallback only when no parent known
        sub_data  = subtype_knowledge.get(subtype_name, {})
        sub_diag  = sub_data.get("diagnosis") or {}
        sub_treat = sub_data.get("treatment") or {}
        for bm_item in sub_diag.get("biomarkers", []):
            bm_name = bm_item["name"] if isinstance(bm_item, dict) else str(bm_item)
            if bm_name:
                bm_idx = add_node("biomarker", f"{subtype_name}_{bm_name}", "BM",
                                  extra={"significance": bm_item.get("significance", "") if isinstance(bm_item, dict) else "", "subtype": subtype_name})
                add_edge(sub_idx, bm_idx, "HAS_BIOMARKER")
        for swt_item in (sub_treat.get("stage_wise_treatment") or []):
            if isinstance(swt_item, dict):
                _wire_stage_treatment(swt_item, sub_idx, subtype_name)
        if not sub_treat.get("stage_wise_treatment"):
            for reg in sub_treat.get("chemotherapy", {}).get("regimens", []):
                if not isinstance(reg, dict) or not reg.get("name"): continue
                reg_idx = add_node("regimen", f"{subtype_name}_{reg['name']}", "REG",
                                   extra={"regimen_name": reg["name"], "subtype": subtype_name, "drugs": reg.get("drugs", [])})
                add_edge(sub_idx, reg_idx, "TREATED_BY_REGIMEN")
                for drug in (reg.get("drugs") or []):
                    if drug:
                        drg_idx = add_node("drug", drug, "DRG")
                        add_edge(reg_idx, drg_idx, "CONTAINS_DRUG")
        for drug in sub_treat.get("targeted_therapy", {}).get("drugs", []):
            if drug:
                drg_idx = add_node("drug", drug, "DRG")
                add_edge(sub_idx, drg_idx, "TARGETED_DRUG")
        for drug in sub_treat.get("hormone_therapy", {}).get("drugs", []):
            if drug:
                drg_idx = add_node("drug", drug, "DRG")
                add_edge(sub_idx, drg_idx, "HORMONE_DRUG")

    for drug in understanding.get("drugs", []):
        drg_idx = add_node("drug", drug, "DRG")
        add_edge(dis_idx, drg_idx, "RECOMMENDED_DRUG")

    for swt_item in (treatment_knowledge.get("stage_wise_treatment") or []):
        if isinstance(swt_item, dict):
            _wire_stage_treatment(swt_item, dis_idx, "")

    for reg in treatment_knowledge.get("chemotherapy", {}).get("regimens", []):
        if not isinstance(reg, dict) or not reg.get("name"): continue
        reg_idx = add_node("regimen", reg["name"], "REG",
                           extra={"regimen_name": reg["name"], "indication": reg.get("indication", ""), "drugs": reg.get("drugs", [])})
        add_edge(dis_idx, reg_idx, "TREATED_BY_REGIMEN")
        for drug in (reg.get("drugs") or []):
            if drug:
                drg_idx = add_node("drug", drug, "DRG")
                add_edge(reg_idx, drg_idx, "CONTAINS_DRUG")

    node_type_counts = {}
    for n in nodes:
        node_type_counts[n["type"]] = node_type_counts.get(n["type"], 0) + 1
    logger.info(f"Knowledge graph: {len(nodes)} nodes | {len(edges)} edges | types={node_type_counts}")
    return {"nodes": nodes, "edges": edges, "index_map": index_map, "node_type_counts": node_type_counts}


# ═════════════════════════════════════════════════════════════════
# SECTION 19 — NEO4J GRAPH INDEXING + GRAPH SEARCH  (kept)
# ═════════════════════════════════════════════════════════════════

def push_graph_to_neo4j(graph: dict, doc_id: str) -> dict:
    if not _NEO4J_AVAILABLE:
        logger.warning("[Neo4j] neo4j driver not installed — skipping graph push")
        return {"pushed": False, "reason": "driver_not_installed"}
    if not (NEO4J_URI and NEO4J_PASSWORD):
        logger.warning("[Neo4j] NEO4J_URI / NEO4J_PASSWORD not set — skipping graph push")
        return {"pushed": False, "reason": "credentials_not_set"}

    driver = GraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        with driver.session() as session:
            session.run(
                "CREATE CONSTRAINT IF NOT EXISTS FOR (n:ClinicalNode) REQUIRE n.index IS UNIQUE"
            )
            for node in graph["nodes"]:
                props = {k: v for k, v in node.items() if isinstance(v, (str, int, float, bool)) or v is None}
                session.run(
                    """
                    MERGE (n:ClinicalNode {index: $index})
                    SET n += $props, n.node_type = $node_type, n.doc_id = $doc_id
                    """,
                    index=node["index"], props=props, node_type=node["type"], doc_id=doc_id,
                )
            for edge in graph["edges"]:
                props = {k: v for k, v in edge.items() if isinstance(v, (str, int, float, bool)) or v is None}
                session.run(
                    """
                    MATCH (a:ClinicalNode {index: $from_idx})
                    MATCH (b:ClinicalNode {index: $to_idx})
                    MERGE (a)-[r:RELATES {relationship: $rel, doc_id: $doc_id}]->(b)
                    SET r += $props
                    """,
                    from_idx=edge["from"], to_idx=edge["to"], rel=edge["relationship"],
                    doc_id=doc_id, props=props,
                )
        logger.info(f"[Neo4j] Pushed {len(graph['nodes'])} nodes / {len(graph['edges'])} edges | doc_id={doc_id}")
        return {"pushed": True, "nodes": len(graph["nodes"]), "edges": len(graph["edges"])}
    except Exception as exc:
        logger.error(f"[Neo4j] graph push failed: {exc}")
        return {"pushed": False, "reason": str(exc)}
    finally:
        driver.close()


def neo4j_graph_search(start_name: str, max_hops: int = 2, doc_id: Optional[str] = None) -> list[dict]:
    if not _NEO4J_AVAILABLE or not (NEO4J_URI and NEO4J_PASSWORD):
        logger.warning("[Neo4j] not configured — graph search skipped")
        return []

    driver = GraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        with driver.session() as session:
            query = (
                f"MATCH (start:ClinicalNode {{name: $name}})"
                f"-[r:RELATES*1..{max_hops}]-(connected:ClinicalNode) "
                "WHERE $doc_id IS NULL OR connected.doc_id = $doc_id "
                "RETURN DISTINCT connected.index AS index, connected.name AS name, "
                "       connected.node_type AS node_type "
                "LIMIT 100"
            )
            results = session.run(query, name=start_name, doc_id=doc_id)
            return [dict(r) for r in results]
    except Exception as exc:
        logger.error(f"[Neo4j] graph search failed: {exc}")
        return []
    finally:
        driver.close()


# ═════════════════════════════════════════════════════════════════
# SECTION 20 — BATCHING HELPERS  (kept — used by discovery + disease + relationships)
# ═════════════════════════════════════════════════════════════════

def _batch_sections(sections: list[Section], max_chars: int = MAX_AGENT_CHARS) -> list[list[Section]]:
    if not sections: return []
    batches: list[list[Section]] = []
    current: list[Section] = []
    current_len: int = 0
    for sec in sections:
        sec_len = len(sec.text)
        if current and (current_len + sec_len) > max_chars:
            batches.append(current)
            current = []; current_len = 0
        current.append(sec); current_len += sec_len
    if current: batches.append(current)
    return batches


def _build_agent_user_message(agent_key: str, sections: list[Section], batch_num: int = 1, total_batches: int = 1) -> str:
    parts = []
    for sec in sections:
        parts.append(
            f"=== SECTION: {sec.section_title} (pages {sec.page_start}–{sec.page_end}) ===\n"
            f"[SOURCE_PAGE_RANGE: {sec.page_start}-{sec.page_end}]\n"
            f"{sec.text}"
        )
    combined   = "\n\n".join(parts)
    batch_note = f"\n[Batch {batch_num}/{total_batches}]\n" if total_batches > 1 else ""
    return (
        f"Extract {agent_key.upper()} knowledge from these {len(sections)} retrieved sections:{batch_note}\n\n"
        f"IMPORTANT: For every extracted item, add a 'source_page' field with the page number "
        f"from the [SOURCE_PAGE_RANGE] where it appears.\n\n"
        f"{combined}"
    )


# ═════════════════════════════════════════════════════════════════
# SECTION 21 — DDX VALIDATION  (kept — dynamic LLM check, not hardcoded)
# ═════════════════════════════════════════════════════════════════


async def validate_discovered_entities(discovery: dict, semaphore: asyncio.Semaphore) -> dict:
    """
    Generic, non-hardcoded validation: ask the LLM to flag entities that are NOT
    valid instances of their assigned bucket type, using structural definitions
    only — no fixed reject list, works for any specialty.
    """
    payload = {
        "diseases": _extract_entity_names(discovery.get("diseases", [])),
        "subtypes": [
            s.get("name", "") if isinstance(s, dict) else str(s)
            for s in discovery.get("subtypes", [])
        ],
        "biomarkers": _extract_entity_names(discovery.get("biomarkers", [])),
    }
    if not any(payload.values()):
        return discovery

    try:
        async with semaphore:
            result = await _chat_async(
                system=(
                    "You are a clinical taxonomy reviewer. Given lists of entities labeled "
                    "as disease/subtype/biomarker, identify any entity that does NOT actually "
                    "fit its label by structural definition — e.g. a symptom, adverse event, "
                    "lab finding, or complication mislabeled as a disease; a severity/time-course "
                    "descriptor mislabeled as a subtype; a non-measurable entity mislabeled as a "
                    "biomarker. Use only general clinical taxonomy reasoning, not familiarity with "
                    "any specific condition.\n"
                    "Return ONLY valid JSON: "
                    '{"reject": {"diseases": [], "subtypes": [], "biomarkers": []}}'
                ),
                user=json.dumps(payload),
            )
    except Exception as e:
        logger.warning(f"[EntityValidator] failed: {e}")
        return discovery

    rejects = result.get("reject", {}) if isinstance(result, dict) else {}

    bad_diseases  = {n.lower() for n in rejects.get("diseases", [])}
    bad_subtypes  = {n.lower() for n in rejects.get("subtypes", [])}
    bad_biomarkers = {n.lower() for n in rejects.get("biomarkers", [])}

    if bad_diseases:
        discovery["diseases"] = [
            e for e in discovery.get("diseases", [])
            if (e.get("name", "") if isinstance(e, dict) else str(e)).lower() not in bad_diseases
        ]
    if bad_subtypes:
        discovery["subtypes"] = [
            s for s in discovery.get("subtypes", [])
            if (s.get("name", "") if isinstance(s, dict) else str(s)).lower() not in bad_subtypes
        ]
    if bad_biomarkers:
        discovery["biomarkers"] = [
            e for e in discovery.get("biomarkers", [])
            if (e.get("name", "") if isinstance(e, dict) else str(e)).lower() not in bad_biomarkers
        ]

    if bad_diseases or bad_subtypes or bad_biomarkers:
        logger.info(
            f"[EntityValidator] rejected diseases={bad_diseases} subtypes={bad_subtypes} biomarkers={bad_biomarkers}"
        )
    return discovery

async def validate_entity_relationships(discovery: dict, semaphore: asyncio.Semaphore) -> dict:
    """
    Specialty-agnostic relationship validation: checks that every subtype's
    parent_diseases actually exist in the discovered disease list, and asks
    the LLM to assign a parent for any subtype that's missing one — using
    only the entity names already discovered, no external knowledge injection.
    """
    disease_names = set(n.lower() for n in _extract_entity_names(discovery.get("diseases", [])))
    subtypes = discovery.get("subtypes", [])
    if not subtypes or not disease_names:
        return discovery

    orphans = []
    for s in subtypes:
        if not isinstance(s, dict):
            continue
        parents = s.get("parent_diseases") or []
        valid_parents = [p for p in parents if p.lower() in disease_names]
        s["parent_diseases"] = valid_parents  # drop hallucinated/invalid parent refs
        if not valid_parents:
            orphans.append(s.get("name", ""))

    orphans = [o for o in orphans if o]
    if not orphans:
        logger.info("[RelationshipValidator] all subtypes have valid parent_diseases")
        return discovery

    try:
        async with semaphore:
            result = await _chat_async(
                system=(
                    "You are a clinical taxonomy reviewer. Given a list of disease names "
                    "and a list of subtypes with no assigned parent, assign each subtype to "
                    "the single most plausible parent disease FROM THE GIVEN DISEASE LIST ONLY, "
                    "based on standard clinical classification reasoning. If no disease in the "
                    "list is a plausible parent for a subtype, omit that subtype from the output.\n"
                    "Return ONLY valid JSON: {\"assignments\": {\"<subtype>\": \"<disease>\"}}"
                ),
                user=json.dumps({
                    "diseases": _extract_entity_names(discovery.get("diseases", [])),
                    "orphan_subtypes": orphans,
                }),
            )
        assignments = result.get("assignments", {}) if isinstance(result, dict) else {}
    except Exception as e:
        logger.warning(f"[RelationshipValidator] LLM assignment failed: {e}")
        assignments = {}

    for s in subtypes:
        if not isinstance(s, dict):
            continue
        name = s.get("name", "")
        if name in orphans and name in assignments:
            assigned = assignments[name]
            if assigned.lower() in disease_names:
                s["parent_diseases"] = [assigned]

    still_orphaned = [s.get("name") for s in subtypes if isinstance(s, dict) and not s.get("parent_diseases")]
    if still_orphaned:
        logger.warning(f"[RelationshipValidator] subtypes with no resolvable parent: {still_orphaned}")

    discovery["subtypes"] = subtypes
    return discovery

def _validate_differential_diagnosis(ddx: list, disease_name: str, specialty: str) -> list:
    if not ddx or len(ddx) < 5: return ddx
    names = [d.get("condition", "") if isinstance(d, dict) else str(d) for d in ddx]
    try:
        result = _chat(
            system=(
                "You are a clinical knowledge quality reviewer. "
                "Given a disease and differential diagnoses, return only plausible differentials. "
                "Return ONLY valid JSON: {\"keep\": [\"condition1\"]}"
            ),
            user=f"Disease: {disease_name}\nSpecialty: {specialty}\nDifferentials:\n" + "\n".join(f"- {n}" for n in names if n),
        )
        keep_set = set(result.get("keep", names))
        filtered = [d for d in ddx if (d.get("condition", "") if isinstance(d, dict) else str(d)) in keep_set]
        if len(ddx) - len(filtered): logger.info(f"[DDx] Removed {len(ddx)-len(filtered)} irrelevant differentials")
        return filtered
    except Exception as exc:
        logger.warning(f"DDx validation failed ({exc})")
        return ddx


# ══════════════════════════════════════════════════════════════
# SECTION 22 — SKILL GENERATION  (kept — operates on diag/treat/subtype knowledge)
# ═════════════════════════════════════════════════════════════════

def _is_subtype_of(candidate: str, parent: str) -> bool:
    p_tokens = set(re.findall(r'\b\w+\b', parent.lower()))
    c_lower  = candidate.lower()
    if len(p_tokens) == 1:
        return list(p_tokens)[0] in c_lower
    return all(tok in c_lower for tok in p_tokens)


def _deduplicate_disease_list(diseases: list[str], primary: str) -> list[str]:
    seen_norm: set[str] = set()
    unique: list[str]   = []
    def _norm(s: str) -> str:
        return re.sub(r'\s+', ' ', s.strip().lower())
    for name in diseases:
        if not name or not isinstance(name, str): continue
        n = _norm(name)
        if not n or n in seen_norm: continue
        seen_norm.add(n)
        unique.append(name.strip())
    return [name for name in unique if _norm(name) != _norm(primary)]


def _build_disease_aliases(disease_name: str) -> list[str]:
    base   = disease_name.lower().strip()
    parts  = [base]
    spaced = base.replace("-", " ")
    if spaced != base: parts.append(spaced)
    if "/" in base: parts.append(base.replace("/", " "))
    words = base.split()
    if len(words) >= 2:
        initials = "".join(w[0] for w in words if len(w) > 1)
        if len(initials) >= 2: parts.append(initials)
    seen, unique = set(), []
    for p in parts:
        p = p.strip()
        if p and p not in seen:
            seen.add(p); unique.append(p)
    return unique


def _is_redundant_subtype(candidate: str, all_subtypes: list[str]) -> bool:
    c_lower = candidate.lower().strip()
    for other in all_subtypes:
        if other == candidate: continue
        o_lower = other.lower().strip()
        if c_lower == o_lower: return True
        if c_lower in o_lower and len(c_lower) < len(o_lower): return True
    return False


def _has_clinical_content(body: dict, skill_type: str, min_items: int = 2) -> bool:
    count = 0
    if skill_type == "diagnosis":
        if body.get("diagnostic_criteria"): count += 1
        count += len(body.get("biomarkers", []))
        count += len(body.get("staging", []))
        count += len(body.get("subtypes", []))
        cp = body.get("clinical_presentation", {})
        count += len(cp.get("symptoms", [])) + len(cp.get("signs", []))
        count += len(body.get("diagnostic_pathway", []))
        count += len(body.get("differential_diagnosis", []))
    elif skill_type == "treatment":
        count += len(body.get("stage_wise_treatment", []))
        count += len(body.get("chemotherapy", {}).get("regimens", []))
        count += len(body.get("targeted_therapy", {}).get("drugs", []))
        count += len(body.get("immunotherapy", {}).get("drugs", []))
        count += len(body.get("hormone_therapy", {}).get("drugs", []))
        count += len(body.get("contraindications", []))
        count += len(body.get("recommendations", []))
        count += len(body.get("if_then_rules", []))
        if body.get("treatment_principles"): count += 1
    return count >= min_items


def _normalize_keyword(v):
    if v is None: return None
    if isinstance(v, str): return v.strip() if v.strip() else None
    if isinstance(v, list):
        for item in v:
            result = _normalize_keyword(item)
            if result and isinstance(result, str): return result
        return None
    if isinstance(v, dict):
        for key in ["name", "drug", "trial", "regimen_name", "stage", "condition", "subtype", "disease_name", "finding"]:
            val = v.get(key)
            if val and isinstance(val, str) and val.strip(): return val.strip()
        return None
    try:
        s = str(v).strip()
        return s if s else None
    except: return None


def _resolve_per_disease_knowledge(disease_name: str, diagnosis_knowledge: dict, treatment_knowledge: dict) -> tuple[dict, dict]:
    aliases = _build_disease_aliases(disease_name)
    def _matches(obj) -> bool:
        try: return any(alias in json.dumps(obj).lower() for alias in aliases)
        except: return False
    diag = _empty_diagnosis()
    for f in ("staging", "biomarkers", "key_evidence", "differential_diagnosis", "exclusion_criteria"):
        items   = diagnosis_knowledge.get(f, [])
        matched = [i for i in items if _matches(i)]
        diag[f] = matched if matched else items
        if not matched: diag[f"_{f}_inherited"] = True
    items = diagnosis_knowledge.get("subtypes", [])
    matched = [s for s in items if (isinstance(s, dict) and _matches(s.get("name", ""))) or (isinstance(s, str) and _matches(s))]
    diag["subtypes"] = matched if matched else items
    if not matched: diag["_subtypes_inherited"] = True
    for f in ("disease_overview", "diagnostic_criteria", "investigations", "molecular_testing",
              "risk_stratification", "diagnostic_pathway", "special_populations", "clinical_presentation"):
        diag[f] = diagnosis_knowledge.get(f, diag.get(f, {})); diag[f"_{f}_inherited"] = True
    diag["source_pages"]   = diagnosis_knowledge.get("source_pages", [])
    diag["_disease_focus"] = disease_name
    own = sum(1 for k in ("staging", "biomarkers", "key_evidence", "subtypes") if diag.get(k) and not diag.get(f"_{k}_inherited"))
    diag["_specificity_ratio"] = round(own / 4, 3)

    treat = _empty_treatment()
    swt = treatment_knowledge.get("stage_wise_treatment", [])
    matched_swt = [s for s in swt if isinstance(s, dict) and _matches(s)]
    treat["stage_wise_treatment"] = matched_swt if matched_swt else swt
    if not matched_swt: treat["_stage_wise_treatment_inherited"] = True
    rules = treatment_knowledge.get("if_then_rules", [])
    matched_rules = [r for r in rules if isinstance(r, dict) and _matches(r)]
    treat["if_then_rules"] = matched_rules if matched_rules else rules
    if not matched_rules: treat["_if_then_rules_inherited"] = True
    for f in ("key_evidence", "recommendations"):
        items = treatment_knowledge.get(f, [])
        matched = [i for i in items if _matches(i)]
        treat[f] = matched if matched else items
        if not matched: treat[f"_{f}_inherited"] = True
    for mf in ("chemotherapy", "targeted_therapy", "immunotherapy", "hormone_therapy", "surgery", "radiation"):
        parent_val = treatment_knowledge.get(mf, treat.get(mf))
        treat[mf] = parent_val
        if not (parent_val and _matches(parent_val)):
            treat[f"_{mf}_inherited"] = True
    for uf in ("follow_up", "supportive_care", "monitoring", "dose_modifications", "contraindications",
               "risk_stratification", "special_populations"):
        treat[uf] = treatment_knowledge.get(uf, treat.get(uf, {})); treat[f"_{uf}_inherited"] = True
    treat["treatment_principles"] = treatment_knowledge.get("treatment_principles", "")
    treat["source_pages"]         = treatment_knowledge.get("source_pages", [])
    treat["_disease_focus"]       = disease_name
    own_t = sum(1 for k in ("stage_wise_treatment", "if_then_rules", "key_evidence") if treat.get(k) and not treat.get(f"_{k}_inherited"))
    treat["_specificity_ratio"] = round(own_t / 3, 3)
    return diag, treat


def _skill_index(doctor_id: str, disease_type: str, subtype: str, skill_type: str) -> str:
    base = re.sub(r"[^A-Z0-9]", "_", f"{disease_type}_{subtype}_{skill_type}".upper())
    return f"{doctor_id}_{base}"

def _relevant_relationships(
    entity_names: set[str], disease_label: str, subtype: str, relationship_store,
    skill_type: Optional[str] = None,   # NEW (Priority 2)
) -> list[dict]:
    if not relationship_store:
        return []
    names_lower = {n.lower() for n in entity_names if n}
    names_lower |= {disease_label.lower(), subtype.lower()}
    subtype_lower = subtype.lower().strip()
    out = []
    for rel in relationship_store.all_relationships():
        s, t = rel.source.lower(), rel.target.lower()

        rel_subtype = (rel.subtype or "").lower().strip()
        if rel_subtype and subtype_lower and subtype_lower != "general" and rel_subtype != subtype_lower:
            continue

        # NEW (Priority 2) — a diagnosis skill shouldn't carry "treats" /
        # "contraindicated_with" relationships, and a treatment skill
        # shouldn't carry "evaluates" / "differentiates" ones. Filter by
        # phase using the schema's own relation vocabulary (or embedding
        # fallback for novel relation types).
        if skill_type in ("diagnosis", "treatment"):
            if _relation_phase(rel.relation) != skill_type:
                continue

        if any(n and (n in s or s in n or n in t or t in n) for n in names_lower):
            out.append(rel.to_dict())
    return out

def validate_skill_relevance(
    skill_body: dict,
    subtype: str,
    disease_label: str,
    sibling_subtypes: list[str],
    skill_type: str,
) -> tuple[bool, str]:
    """
    Phase 3a — post-generation gate. Runs AFTER all the extraction/backfill/
    relevance-filter machinery, as a final sanity check before a skill is
    persisted. Rule-based, no LLM.

    Checks:
      1. Subtype purity — body must not contain a sibling subtype's name
         without also containing this subtype's own name (same rule as the
         relevance_score hard-exclusion, applied here as a final sweep over
         the fully-assembled body rather than per-entity during assembly).
      2. Non-empty content — an empty/near-empty skill is not useful and
         should not reach doctors for review.
      3. Relationship consistency — if the body carries a "relationships"
         list (Phase 2a/Patch C), every relationship explicitly tagged with
         a DIFFERENT subtype is a contamination signal.

    Returns (is_valid, reason). Rejection is logged, not silently dropped,
    so doctors/engineers can audit what got filtered and why.
    """
    if subtype and subtype != "General":
        body_text = json.dumps(skill_body).lower()
        subtype_lower = subtype.lower()
        for sib in sibling_subtypes:
            sib_lower = sib.lower().strip()
            if sib_lower and sib_lower != subtype_lower and sib_lower in body_text and subtype_lower not in body_text:
                return False, f"contains sibling subtype '{sib}' without self-reference"

    has_content = _has_clinical_content(skill_body, skill_type, min_items=1)
    if not has_content:
        return False, "no clinical content after filtering/backfill"

    rels = skill_body.get("relationships") or []
    if subtype and subtype != "General":
        for rel in rels:
            rel_subtype = (rel.get("subtype") or "").lower().strip()
            if rel_subtype and rel_subtype != subtype.lower().strip():
                return False, f"carries relationship tagged for a different subtype ('{rel.get('subtype')}')"

    return True, ""

def generate_skills(
    doctor_id:           str,
    doc_id:              str,
    guideline_name:      str,
    guideline_version:   str,
    understanding:       dict,
    diagnosis_knowledge: dict,
    treatment_knowledge: dict,
    subtype_knowledge:   dict[str, dict],
    per_disease_knowledge: dict[str, dict], 
    index_map:           dict,
    relationship_store:  Optional["RelationshipStore"] = None,   # NEW
) -> list[dict]:

    raw_subtype_entries: list[dict] = []
    for s in understanding.get("subtypes", []):
        if isinstance(s, str) and s.strip():
            raw_subtype_entries.append({"name": s.strip(), "parent_diseases": []})
        elif isinstance(s, dict):
            name = (s.get("name") or s.get("subtype") or "").strip()
            if name:
                raw_subtype_entries.append({
                    "name": name,
                    "parent_diseases": s.get("parent_diseases") or [],
                })

    seen_normalized: set[str] = set()
    case_deduped: list[dict]  = []
    for entry in raw_subtype_entries:
        norm = re.sub(r"\s+", " ", entry["name"].lower().strip())
        if norm not in seen_normalized:
            seen_normalized.add(norm)
            case_deduped.append(entry)

    subtype_names_only = [e["name"] for e in case_deduped]
    subtypes_with_parents = [
        e for e in case_deduped
        if not _is_redundant_subtype(e["name"], subtype_names_only)
    ]
    subtypes = [e["name"] for e in subtypes_with_parents]
    subtype_parent_lookup = {e["name"]: (e["parent_diseases"][0] if e["parent_diseases"] else None)
                              for e in subtypes_with_parents}
    understanding["subtypes"] = subtypes_with_parents  # keep dicts for graph builder

    disease_type = understanding.get("disease_type") or (
        (understanding.get("diseases") or ["Unknown"])[0]
    )
    logger.info(f"Subtypes after dedup: {len(subtypes)}")

    raw_diseases: list[str] = []
    for d in (understanding.get("diseases") or []):
        if isinstance(d, str) and d.strip(): raw_diseases.append(d.strip())
    dn = understanding.get("disease_name", "")
    if dn and dn.strip() and dn.strip() not in raw_diseases:
        raw_diseases.insert(0, dn.strip())
    subtype_set_lower   = {s.lower().strip() for s in subtypes}
    per_disease_targets = _deduplicate_disease_list(raw_diseases, disease_type)
    per_disease_targets = [
        d for d in per_disease_targets
        if not _is_subtype_of(d, disease_type)
        and d.lower().strip() not in subtype_set_lower
        and not _is_subtype_of(disease_type, d)
    ]
    per_disease_targets = per_disease_targets[:MAX_PER_DISEASE_SKILLS]

    skills: list[dict] = []
    now = datetime.now(timezone.utc).isoformat()

    def make_skill(subtype: str, skill_type: str, body: dict, disease_label: str = "") -> dict:
        effective_disease = disease_label or disease_type
        if effective_disease == "Unknown":
            real_diseases = understanding.get("diseases", [])
            non_unknown = [d for d in real_diseases if d and d != "Unknown"]
            if non_unknown:
                effective_disease = non_unknown[0]
        skill_id = str(uuid.uuid4())
        idx      = _skill_index(doctor_id, effective_disease, subtype, skill_type)
        keywords = [effective_disease, subtype, skill_type]

        # Dynamic keyword extraction — walks the body for ANY list-of-entities
        # field at any depth (biomarkers, staging, chemotherapy.regimens,
        # targeted_therapy.drugs, key_evidence, backfilled/expanded entities
        # from build_subtype_knowledge_package, future specialty fields, etc.)
        # instead of a hardcoded diagnosis/treatment field list. Keeps
        # trigger_keywords in sync with the richer relevance-filtered bodies.
        for path, pool in walk_entity_fields(body):
            for item in pool:
                nk = _normalize_keyword(item)
                if nk:
                    keywords.append(nk)
        for rule in (body.get("if_then_rules") or []):
            if isinstance(rule, dict):
                cond = rule.get("condition", "")
                if cond and isinstance(cond, str):
                    keywords.extend(re.findall(r'\b[A-Za-z0-9\-]+\b', cond)[:3])

        source_pages = _clean_pages(body.get("source_pages") or [])
        source_sections = list(dict.fromkeys(
            s.get("section_title") or s.get("title", "")
            for s in (body.get("_source_sections") or [])
            if isinstance(s, dict) and (s.get("section_title") or s.get("title"))
        )) if isinstance(body.get("_source_sections"), list) else []

        # NEW (Issue #6) — pull in relationship triples touching any entity
        # actually present in this skill's body, so diagnosis/treatment
        # agents get "HHV8 -> evaluates -> HHV8+ MCD" style reasoning
        # instead of only flat lists.
        entity_names_for_rels = {
            _entity_display_name(item) for _, pool in walk_entity_fields(body) for item in pool
        }
        rels = _relevant_relationships(entity_names_for_rels, effective_disease, subtype, relationship_store, skill_type=skill_type)
        if rels:
            body = {**body, "relationships": rels}

        cleaned_body = clean_empty_fields(body) or body

        skill_name = f"{subtype} {effective_disease} {skill_type.title()} Skill"

        # Build trigger_keywords FIRST — everything else depends on it.
        trigger_keywords_list = list(set(
            str(nk) for nk in [_normalize_keyword(k) for k in keywords if k]
            if nk and not isinstance(nk, dict)
        ))

        skill_description = (
            f"Structured {skill_type} knowledge for {subtype} {effective_disease} "
            f"extracted from {guideline_name} {guideline_version or ''}."
        )

        # Retrieval-time embedding text: name + description + disease/subtype +
        # trigger keywords ONLY — never skill_md or body. This is what
        # Phase 2's vector search matches against; the full body is loaded
        # lazily after a skill already wins on this lightweight text.
        _skill_search_text = (
            f"{skill_name}\n"
            f"{skill_description}\n"
            f"{effective_disease} {subtype} {skill_type}\n"
            f"{' '.join(trigger_keywords_list[:40])}"
        )
        skill_vector = _embed_text(_skill_search_text)

        # NEW (Issue #7) — QA signal: if trigger keywords contain entities
        # that never made it into the body, that's the exact "keywords
        # richer than skill content" symptom from the critique.
        tk_lower = {k.lower() for k in trigger_keywords_list}
        body_entity_names_lower = {n.lower() for n in entity_names_for_rels}
        completeness = len(tk_lower & body_entity_names_lower) / max(len(tk_lower), 1)
        if completeness < 0.5:
            logger.warning(
                f"[SkillCompleteness] '{skill_name}': only {completeness:.0%} of trigger "
                f"keywords are represented in the skill body — possible knowledge loss"
            )

        # NEW (Phase 3a) — final validation gate before this skill is built.
        sibling_subtypes_for_validation = [s for s in subtypes if s != subtype]
        is_valid, rejection_reason = validate_skill_relevance(
            cleaned_body, subtype, effective_disease, sibling_subtypes_for_validation, skill_type,
        )
        if not is_valid:
            logger.warning(f"[SkillValidation] REJECTED '{skill_name}': {rejection_reason}")
            return None   # signal to caller: do not append this skill

        return {
            "skill_id":          skill_id,
            "embedding":         skill_vector,      # <-- ADD THIS LINE
            "embedding_model":   EMBEDDING_MODEL,   # <-- ADD THIS LINE
            "skill_index":       idx,
            "doctor_id":         doctor_id,
            "doc_id":            doc_id,
            "skill_type":        skill_type,
            "skill_category":    "general",
            "disease_type":      effective_disease,
            "subtype":           subtype,
            "name":              skill_name,
            "description": skill_description,
            "trigger_keywords": trigger_keywords_list,
            "knowledge_completeness": round(completeness, 3),   # NEW — exposes the QA signal on the saved skill itself,
            "graph_path":        f"{effective_disease} > {subtype}",
            "guideline":         guideline_name,
            "guideline_version": guideline_version,
            "status":            "pending_review",
            "version":           1,
            "body":              cleaned_body,
            "source_pages":      source_pages,
            "source_sections":   source_sections,
            "references": {
                "guideline":    guideline_name,
                "version":      guideline_version,
                "doc_id":       doc_id,
                "graph_nodes":  [index_map.get(effective_disease), index_map.get(subtype)],
                "source_pages": source_pages,
            },
            "created_at": now,
            "updated_at": now,
        }
    # TIER 1: General skills
    for s in (make_skill("General", "diagnosis", diagnosis_knowledge),
              make_skill("General", "treatment", treatment_knowledge)):
        if s is not None:
            skills.append(s)
    logger.info(f"[TIER 1] General skills created for disease_type='{disease_type}'")

    # TIER 2: Per-disease skills
    TIER2_REJECT_THRESHOLD = 0.10
    for dis_name in per_disease_targets:
        dis_diag, dis_treat = per_disease_knowledge.get(dis_name, (None, None))
        if dis_diag is None:
            # fallback for diseases somehow missing from the batch (shouldn't normally happen)
            dis_diag, dis_treat = _resolve_per_disease_knowledge(dis_name, diagnosis_knowledge, treatment_knowledge)
            if dis_diag.get("_specificity_ratio", 0.0) < TIER2_REJECT_THRESHOLD:
                logger.warning(f"[TIER 2] '{dis_name}': specificity too low — skipping")
                continue
        if _has_clinical_content(dis_diag, "diagnosis"):
            s = make_skill("General", "diagnosis", dis_diag, disease_label=dis_name)
            if s is not None:
                skills.append(s)
        if _has_clinical_content(dis_treat, "treatment"):
            s = make_skill("General", "treatment", dis_treat, disease_label=dis_name)
            if s is not None:
                skills.append(s)

    # TIER 3: Per-subtype skills
    for subtype in subtypes:
        sub_data = subtype_knowledge.get(subtype, {})
        if sub_data.get("rejected"):
            logger.info(f"[TIER 3] Skipping rejected subtype '{subtype}': {sub_data.get('rejection_reason','')}")
            continue
        parent_for_subtype = subtype_parent_lookup.get(subtype) or disease_type

        # FIXED: no more unfiltered copy.deepcopy(diagnosis_knowledge) —
        # that dumped the ENTIRE parent disease's biomarkers/regimens into
        # every subtype skill regardless of relevance. Now we run the same
        # relevance-filtered package builder used in Section 16.
        if sub_data.get("diagnosis"):
            sub_diag = sub_data["diagnosis"]
        else:
            sub_diag, _ = build_subtype_knowledge_package(
                subtype_name=subtype, parent_disease=parent_for_subtype,
                subtype_diag={}, subtype_treat={},
                disease_diag=diagnosis_knowledge, disease_treat=treatment_knowledge,
                subtype_pages=None, relationship_store=None,   # no retrieval happened — token/parent tier only
            )
        if sub_data.get("treatment"):
            sub_treat = sub_data["treatment"]
        else:
            _, sub_treat = build_subtype_knowledge_package(
                subtype_name=subtype, parent_disease=parent_for_subtype,
                subtype_diag={}, subtype_treat={},
                disease_diag=diagnosis_knowledge, disease_treat=treatment_knowledge,
                subtype_pages=None, relationship_store=None,
            )
        sub_diag.setdefault("_subtype_focus", subtype)
        sub_treat.setdefault("_subtype_focus", subtype)

        # NEW (Phase 3b) — compute the inheritance delta: which entities in
        # this subtype's own body are marked inherited_from the parent vs.
        # genuinely subtype-specific. Cheap, deterministic, no LLM — reuses
        # walk_entity_fields, which already exists.
        def _inheritance_summary(body: dict) -> dict:
            own, inherited = 0, 0
            for _, pool in walk_entity_fields(body):
                for item in pool:
                    if isinstance(item, dict) and item.get("inherited_from"):
                        inherited += 1
                    else:
                        own += 1
            return {"own_entities": own, "inherited_entities": inherited, "inherits_from": parent_for_subtype}

        sub_diag["_inheritance"] = _inheritance_summary(sub_diag)
        sub_treat["_inheritance"] = _inheritance_summary(sub_treat)

        if _has_clinical_content(sub_diag, "diagnosis"):
            s = make_skill(subtype, "diagnosis", sub_diag, disease_label=parent_for_subtype)
            if s is not None:
                skills.append(s)
        if _has_clinical_content(sub_treat, "treatment"):
            s = make_skill(subtype, "treatment", sub_treat, disease_label=parent_for_subtype)
            if s is not None:
                skills.append(s)

    # Deduplicate by skill_index — keep the more specific one (non-Unknown wins)
    index_to_skill: dict[str, dict] = {}
    for skill in skills:
        idx = skill["skill_index"]
        existing = index_to_skill.get(idx)
        if existing is None:
            index_to_skill[idx] = skill
        else:
            if existing["disease_type"] == "Unknown" and skill["disease_type"] != "Unknown":
                index_to_skill[idx] = skill

    skills = list(index_to_skill.values())

    non_unknown_types = {s["skill_type"] for s in skills if s["disease_type"] != "Unknown"}
    if non_unknown_types:
        skills = [
            s for s in skills
            if s["disease_type"] != "Unknown"
            or s["skill_type"] not in non_unknown_types
        ]

    logger.info(
        f"Skills generated: {len(skills)} | tier1=2 | "
        f"tier2_targets={len(per_disease_targets)} | tier3_subtypes={len(subtypes)}"
    )


    return skills




# ═════════════════════════════════════════════════════════════════
# SECTION 23 — MONGODB STORAGE  (kept)
# ═════════════════════════════════════════════════════════════════

async def save_to_mongodb(
    mongo_uri: str, doctor_id: str, doc_id: str,
    understanding: dict, diagnosis_knowledge: dict, treatment_knowledge: dict,
    graph: dict, skills: list[dict],
    guideline_name: str, guideline_version: str,
    relationship_store: Optional[RelationshipStore] = None,
    neo4j_push_result:  Optional[dict]               = None,
    clinical_skills:    Optional[list[dict]]         = None,
) -> dict:
    motor_client = AsyncIOMotorClient(mongo_uri)
    db   = motor_client[MONGO_DB]
    now  = datetime.now(timezone.utc).isoformat()

    diag_coll      = db["phase1_diagnosis_skills"]
    treat_coll     = db["phase1_treatment_skills"]
    nodes_coll     = db["phase1_graph_nodes"]
    edges_coll     = db["phase1_graph_edges"]
    guideline_coll = db["phase1_guideline_versions"]
    job_coll       = db["phase1_processing_jobs"]
    rel_coll       = db["phase1_clinical_relationships"]
    clinical_skills_coll = db["phase1_clinical_skills"]

    saved: dict[str, Any] = {}

    if graph["nodes"]:
        await nodes_coll.insert_many(graph["nodes"])
        saved["graph_nodes"] = len(graph["nodes"])
    if graph["edges"]:
        await edges_coll.insert_many(graph["edges"])
        saved["graph_edges"] = len(graph["edges"])

    if relationship_store and len(relationship_store) > 0:
        rel_docs = [
            {**r.to_dict(), "doc_id": doc_id, "doctor_id": doctor_id, "created_at": now}
            for r in relationship_store.all_relationships()
        ]
        await rel_coll.insert_many(rel_docs)
        saved["clinical_relationships"] = len(rel_docs)

    if clinical_skills:
        skill_docs = [
            {
                **cs,
                "skill_id": str(uuid.uuid4()),
                "doc_id": doc_id,
                "doctor_id": doctor_id,
                "guideline_name": guideline_name,
                "guideline_version": guideline_version,
                "created_at": now,
            }
            for cs in clinical_skills
        ]
        await clinical_skills_coll.insert_many(skill_docs)
        saved["clinical_skills"] = len(skill_docs)

    for s in skills:
        if "source_pages" in s:
            s["source_pages"] = _clean_pages(s.get("source_pages") or [])
        if "body" in s and isinstance(s["body"], dict):
            body_pages = s["body"].get("source_pages")
            if body_pages is not None:
                s["body"]["source_pages"] = _clean_pages(body_pages or [])

    diag_skills  = [s for s in skills if s["skill_type"] == "diagnosis"]
    treat_skills = [s for s in skills if s["skill_type"] == "treatment"]
    if diag_skills:
        await diag_coll.insert_many(diag_skills)
        saved["diagnosis_skills"] = len(diag_skills)
    if treat_skills:
        await treat_coll.insert_many(treat_skills)
        saved["treatment_skills"] = len(treat_skills)

    # NEW — push the same embeddings already computed in generate_skills()
    # into the dedicated skill-vector Chroma collections, so Phase 2's
    # ChromaRetriever can find them. No re-embedding — reuses skill["embedding"].
    try:
        diag_client = chromadb.PersistentClient(path=CHROMA_PERSIST_PATH)
        diag_coll_v = diag_client.get_or_create_collection("phase1_diagnosis_skills_vectors")
        treat_coll_v = diag_client.get_or_create_collection("phase1_treatment_skills_vectors")

        def _push(coll, skills):
            ids, embeddings, metadatas = [], [], []
            for s in skills:
                if not s.get("embedding"):
                    continue
                ids.append(s["skill_id"])
                embeddings.append(s["embedding"])
                metadatas.append({
                    "doctor_id":    s.get("doctor_id", ""),
                    "doc_id":       s.get("doc_id", ""),
                    "disease_type": s.get("disease_type", ""),
                    "subtype":      s.get("subtype", ""),
                    "skill_type":   s.get("skill_type", ""),
                })
            if ids:
                coll.upsert(ids=ids, embeddings=embeddings, metadatas=metadatas)

        _push(diag_coll_v, diag_skills)
        _push(treat_coll_v, treat_skills)
        logger.info(f"[Chroma] Pushed {len(diag_skills)} diag + {len(treat_skills)} treat skill vectors")
    except Exception as e:
        logger.warning(f"[Chroma] Skill vector push failed: {e}")

    disease_type = understanding.get("disease_type") or (
        (understanding.get("diseases") or ["Unknown"])[0]
    )

    guideline_doc = {
        "doc_id": doc_id, "doctor_id": doctor_id,
        "guideline_name": guideline_name, "guideline_version": guideline_version,
        "disease_type": disease_type, "specialty": understanding.get("specialty", ""),
        "disease_name": understanding.get("disease_name", ""),
        "subtypes": understanding.get("subtypes", []),
        "diseases": understanding.get("diseases", []),
        "index_map": graph["index_map"],
        "skill_ids": [s["skill_id"] for s in skills],
        "source_pages_coverage": _clean_pages([p for s in skills for p in s.get("source_pages", [])]),
        "relationship_count": len(relationship_store) if relationship_store else 0,
        "neo4j_push": neo4j_push_result or {},
        "created_at": now, "updated_at": now,
        "pipeline_version": "v14",
        "embedding_model": EMBEDDING_MODEL,
        "vector_store": "chromadb",
        "chroma_collection": CHROMA_COLLECTION_NAME,
    }
    await guideline_coll.insert_one(guideline_doc)
    saved["guideline_version_record"] = 1

    await job_coll.update_one(
        {"doc_id": doc_id},
        {"$set": {"status": "completed", "completed_at": now, "summary": saved}},
        upsert=True,
    )
    motor_client.close()
    logger.info(f"MongoDB save complete | doc_id={doc_id} | summary={saved}")
    return saved


# ═════════════════════════════════════════════════════════════════
# SECTION 24 — DOCTOR PREVIEW  (kept)
# ═════════════════════════════════════════════════════════════════

def build_doctor_preview(
    understanding: dict, diagnosis_knowledge: dict, treatment_knowledge: dict,
    skills: list[dict], graph: dict,
    guideline_name: str, guideline_version: str,
    relationship_store: Optional[RelationshipStore] = None,
    neo4j_push_result:  Optional[dict]               = None,
) -> dict:
    disease_type     = understanding.get("disease_type", "Unknown")
    all_source_pages = _clean_pages([p for s in skills for p in s.get("source_pages", [])])

    return {
        "guideline": {"name": guideline_name, "version": guideline_version},
        "summary": {
            "disease_name":   understanding.get("disease_name", ""),
            "disease_type":   disease_type,
            "specialty":      understanding.get("specialty", ""),
            "diseases":       understanding.get("diseases", []),
            "subtypes":       understanding.get("subtypes", []),
            "stages":         understanding.get("stages", []),
            "biomarkers":     understanding.get("biomarkers", []),
            "drugs":          understanding.get("drugs", []),
            "regimens":       understanding.get("regimens", []),
            "investigations": understanding.get("investigations", []),
        },
        "graph": {
            "total_nodes":      len(graph["nodes"]),
            "total_edges":      len(graph["edges"]),
            "node_types":       sorted({n["type"] for n in graph["nodes"]}),
            "edge_types":       sorted({e["relationship"] for e in graph["edges"]}),
            "node_type_counts": graph.get("node_type_counts", {}),
            "neo4j":            neo4j_push_result or {},
        },
        "coverage": {
            "source_pages_covered": all_source_pages,
            "total_pages_covered":  len(all_source_pages),
        },
        "relationship_summary": {
            "total_relationships": len(relationship_store) if relationship_store else 0,
            "relation_types": list({r.relation for r in relationship_store.all_relationships()}) if relationship_store else [],
        },
        "embedding_model": EMBEDDING_MODEL,
        "vector_store":    "chromadb",
        "skills_preview": [
            {
                "skill_id":         s["skill_id"],
                "skill_index":      s["skill_index"],
                "name":             s["name"],
                "skill_type":       s["skill_type"],
                "skill_category":   s.get("skill_category", "general"),
                "subtype":          s["subtype"],
                "trigger_keywords": s["trigger_keywords"],
                "graph_path":       s["graph_path"],
                "body_sections":    list(s["body"].keys()),
                "source_pages":     s.get("source_pages", []),
                "is_generic_subtype": s.get("is_generic_subtype", False),
                "specificity_score":  s.get("specificity_score"),
            }
            for s in skills
        ],
        "diagnosis_knowledge": diagnosis_knowledge,
        "treatment_knowledge": treatment_knowledge,
    }


# ═════════════════════════════════════════════════════════════════
# SECTION 25 — DOCTOR APPROVAL  (kept)
# ═════════════════════════════════════════════════════════════════

async def approve_and_save(
    pipeline_result: dict,
    doctor_id: str,
    approved_skill_ids: list[str],
    skills: list[dict],
    edited_skills: dict[str, dict] = {},
) -> dict:
    now = datetime.now(timezone.utc).isoformat()
    for skill in skills:
        sid = skill["skill_id"]
        if sid in approved_skill_ids:
            skill["status"] = "approved"; skill["updated_at"] = now
            skill["approved_by"] = doctor_id; skill["approved_at"] = now
        if sid in edited_skills:
            skill["body"]       = edited_skills[sid]
            skill["updated_at"] = now
            skill["edited"]     = True
    return await save_to_mongodb(
        mongo_uri=MONGO_URI, doctor_id=doctor_id, doc_id=pipeline_result["doc_id"],
        understanding=pipeline_result["understanding"],
        diagnosis_knowledge=pipeline_result["diagnosis_knowledge"],
        treatment_knowledge=pipeline_result["treatment_knowledge"],
        graph=pipeline_result["graph"], skills=skills,
        guideline_name=pipeline_result["preview"]["guideline"]["name"],
        guideline_version=pipeline_result["preview"]["guideline"]["version"],
    )


# ═════════════════════════════════════════════════════════════════
# SECTION 26 — MAIN PIPELINE v14  (Retrieval-First Agentic Architecture
#                                  + Structural Capability Extraction)
# ═════════════════════════════════════════════════════════════════

async def run_phase1_pipeline(
    file_bytes: bytes,
    filename:   str,
    doctor_id:  str,
    doc_id:     Optional[str] = None,
    save_to_db: bool          = False,
) -> dict:
    """
    v14 Pipeline — Retrieval-First Agentic Architecture + Structural Capability
    Extraction (the only functional change vs v13 is inside collect_skills(),
    called in step 10b below — everything else is identical to v13).

        Upload PDF -> Build Sections -> OpenRouter Embedding -> Store in ChromaDB
        -> Discovery Agent (all sections, full read)
        -> Query Builder (no LLM) -> Parallel Retrieval per disease (no LLM)
        -> Coverage Validator (no LLM)
        -> Disease Skill Generation (1 diag + 1 treat LLM call per disease)
        -> Subtype Skill Generation (batched LLM calls, ~20 subtypes per call)
        -> Skill Aggregation (LLM-declared + structurally-derived, no LLM)
        -> Relationship Extraction -> Knowledge Graph -> Neo4j -> MongoDB
    """
    doc_id = doc_id or str(uuid.uuid4())
    _LLM_CACHE.clear()

    relationship_store = RelationshipStore()

    logger.info(
        f"Phase 1 pipeline v14 started | doc_id={doc_id} | file={filename} | "
        f"embedding={EMBEDDING_MODEL} | vector_store=chromadb"
    )

    # 1) PDF / text extraction
    raw_text, toc_entries = extract_text(file_bytes, filename)
    logger.info(f"Text extracted | {len(raw_text):,} chars | {len(toc_entries)} ToC entries")

    page_markers = re.findall(r"---\s*PAGE\s+(\d+)\s*---", raw_text)
    total_pages  = int(page_markers[-1]) if page_markers else max(len(raw_text) // 3000, 1)

    # 2) Section building
    sections = build_sections(raw_text, toc_entries)
    logger.info(f"Sections ready: {len(sections)} | total_pages={total_pages}")

    # 3) OpenRouter embedding + ChromaDB storage
    stored_count = store_section_embeddings(sections, doc_id, doctor_id, mongo_uri=MONGO_URI)
    logger.info(f"Sections embedded & stored in ChromaDB: {stored_count}/{len(sections)}")

    # 4) Discovery agent — reads ALL sections, identifies diseases/subtypes/biomarkers/stages
    semaphore = asyncio.Semaphore(MAX_CONCURRENT_GROQ_CALLS)
    discovery_result = await run_discovery_agent(doc_id=doc_id, semaphore=semaphore, all_sections=sections)
    discovery_result = await validate_discovered_entities(discovery_result, semaphore)   # ADD THIS LINE
    discovery_result = await validate_entity_relationships(discovery_result, semaphore)   # ADD THIS LINE
    discovery_result = await _classify_primary_vs_differential(discovery_result, filename, semaphore)

    all_diseases = _unique_list(
        discovery_result.get("primary_diseases", []) + discovery_result.get("differential_diseases", [])
    ) or discovery_result.get("diseases", [])

    if not all_diseases:
        all_diseases = ["Unknown"]
        logger.warning("[Pipeline] Discovery found no diseases — falling back to 'Unknown'")

    # 5) Query Builder (no LLM) + Parallel Retrieval per disease (no LLM)
    retrieved_by_disease: dict[str, dict[str, list[Section]]] = {}
    for d in all_diseases:
        retrieved_by_disease[d] = await retrieve_context_for_disease(d, doc_id)

    # 6) Coverage Validator (no LLM) — recover zero-retrieval diseases
    gaps = validate_disease_coverage(retrieved_by_disease)
    if gaps:
        retrieved_by_disease = await revalidate_and_retry(gaps, doc_id, retrieved_by_disease)

    coverage_report = compute_disease_page_coverage(retrieved_by_disease, total_pages)

    # 7) Disease Skill Generation — batched LLM call(s) covering EVERY disease
    #    (primary + differential), reusing the contexts already retrieved in step 5.
    primary_disease = (discovery_result.get("primary_diseases") or all_diseases)[0]

    per_disease_knowledge = await generate_disease_knowledge_batched(
        all_diseases, retrieved_by_disease, semaphore,
        discovery_biomarkers=_extract_entity_names(discovery_result.get("biomarkers", [])),
        discovery_investigations=_extract_entity_names(discovery_result.get("procedures", [])),
        discovery_drugs=_extract_entity_names(discovery_result.get("drugs", [])),
    )

    diagnosis_knowledge, treatment_knowledge = per_disease_knowledge.get(
        primary_disease, (_empty_diagnosis(), _empty_treatment())
    )

    merged_understanding = {
        "disease_name":      primary_disease,
        "disease_type":      primary_disease,
        "diseases":          all_diseases,
        "subtypes":          discovery_result.get("subtypes", []),
        "biomarkers":        discovery_result.get("biomarkers", []),
        "stages":            discovery_result.get("stages", []),
        "drugs":             discovery_result.get("drugs", []),
        "investigations":    discovery_result.get("procedures", []),
        "procedures":        discovery_result.get("procedures", []),
        "regimens":          [],
        "specialty":         "",
        "specialties":       [],
        "guideline_name":    filename,
        "guideline_version": "",
    }

    # ── Specialty fallback: infer from diseases + filename ──
    if all_diseases and all_diseases != ["Unknown"]:
        try:
            spec_result = _chat(
                system=(
                    "You are a medical classification expert. "
                    "Given a list of diseases from a clinical guideline and its filename, "
                    "return the primary medical specialty. "
                    "Return ONLY valid JSON: {\"specialty\": \"<specialty>\"}"
                ),
                user=f"Guideline filename: {filename}\nDiseases found: {all_diseases[:5]}",
            )
            if spec_result.get("specialty"):
                merged_understanding["specialty"] = spec_result["specialty"]
                logger.info(f"[SpecialtyFallback] inferred specialty='{spec_result['specialty']}'")
        except Exception as _spec_exc:
            logger.warning(f"[SpecialtyFallback] failed: {_spec_exc}")

    disease_name      = primary_disease
    disease_type      = primary_disease
    specialty         = merged_understanding.get("specialty", "")
    guideline_name    = filename
    guideline_version = ""

    # 8) Clinical relationship extraction (feeds Knowledge Graph) — still reads all sections once
    semaphore_rel = asyncio.Semaphore(MAX_CONCURRENT_GROQ_CALLS)
    relationships = await _extract_relationships(sections, semaphore_rel)
    for rel in relationships:
        relationship_store.add(rel)
    logger.info(f"Relationships extracted: {len(relationship_store)}")

    # 9) Synthesis (consolidation only, no validation gate)
    diagnosis_knowledge = _consolidate(diagnosis_knowledge, _SYNTH_DIAG_SYSTEM, f"diagnosis/{disease_name}")
    treatment_knowledge = _consolidate(treatment_knowledge, _SYNTH_TREAT_SYSTEM, f"treatment/{disease_name}")

    ddx = diagnosis_knowledge.get("differential_diagnosis", [])
    if ddx:
        diagnosis_knowledge["differential_diagnosis"] = _validate_differential_diagnosis(ddx, disease_name, specialty)

    # 10) Subtype Skill Generation — batched, no per-subtype sleep
    subtypes = _extract_entity_names(
        merged_understanding.get("subtypes", [])
    )
    logger.info(
        f"[Subtype Names] count={len(subtypes)} sample={subtypes[:10]}"
    )

    # build the parent map from discovery output (requires fix #4's subtype shape)
    subtype_parent_map = {}
    for s in discovery_result.get("subtypes", []):
        if isinstance(s, dict):
            name = s.get("name", "").strip()
            parents = s.get("parent_diseases") or []
            if name and parents:
                subtype_parent_map[name] = parents[0]

    subtype_knowledge = await derive_subtype_knowledge_batched(
        subtypes=subtypes, doc_id=doc_id, semaphore=semaphore,
        subtype_parent_map=subtype_parent_map,
        biomarkers=merged_understanding.get("biomarkers", []),
        stages=merged_understanding.get("stages", []),
        risk_groups=discovery_result.get("risk_groups", []),
        disease_diag=diagnosis_knowledge,       # NEW — enables backfill
        disease_treat=treatment_knowledge,      # NEW — enables backfill
        relationship_store=relationship_store,  # NEW — enables Tier 1 scoring
    )

    logger.info(f"[Coverage] Subtype discovered={len(subtypes)}")
    logger.info(f"[Coverage] Subtype knowledge generated={len(subtype_knowledge)}")
    non_rejected = sum(1 for v in subtype_knowledge.values() if not v.get("rejected"))
    logger.info(f"[Coverage] Subtype knowledge non-rejected={non_rejected}")

    # 10b) Aggregate skills — LLM-declared (unchanged) + structurally-derived
    #      (NEW in v14, via derive_skills_from_knowledge inside collect_skills).
    #      Zero extra LLM calls either way.
    clinical_skills = collect_skills(per_disease_knowledge, subtype_knowledge)
    entity_count = len(all_diseases) + len(subtypes)
    llm_derived_count        = sum(1 for s in clinical_skills if s.get("derivation") == "llm")
    structural_derived_count = sum(1 for s in clinical_skills if s.get("derivation") == "structural")
    logger.info(
        f"[SkillCoverage] diseases={len(all_diseases)} subtypes={len(subtypes)} "
        f"skills={len(clinical_skills)} (llm={llm_derived_count} structural={structural_derived_count})"
    )
    if entity_count > 0 and len(clinical_skills) < entity_count:
        logger.warning(
            f"[SkillCoverage] skill_count({len(clinical_skills)}) < entity_count({entity_count}) "
            "— possible under-generation"
        )

    # 11) Knowledge Graph build
    graph = build_knowledge_graph(
        doctor_id=doctor_id, doc_id=doc_id,
        understanding=merged_understanding,
        diagnosis_knowledge=diagnosis_knowledge,
        treatment_knowledge=treatment_knowledge,
        subtype_knowledge=subtype_knowledge,
        relationship_store=relationship_store,
    )

    # 12) Push Knowledge Graph to Neo4j
    neo4j_push_result = push_graph_to_neo4j(graph, doc_id)

    # 13) Skill generation
    skills = generate_skills(
        doctor_id=doctor_id, doc_id=doc_id,
        guideline_name=guideline_name, guideline_version=guideline_version,
        understanding=merged_understanding,
        diagnosis_knowledge=diagnosis_knowledge,
        treatment_knowledge=treatment_knowledge,
        subtype_knowledge=subtype_knowledge,
        per_disease_knowledge=per_disease_knowledge,
        index_map=graph["index_map"],
        relationship_store=relationship_store,   # NEW
    )

    for skill in skills:
        subtype = skill.get("subtype", "")
        if subtype and subtype != "General":
            sub_data = subtype_knowledge.get(subtype, {})
            if sub_data.get("is_generic"):
                skill["is_generic_subtype"]     = True
                skill["specificity_score"]      = sub_data.get("specificity_score", 0.0)
                skill["generic_subtype_reason"] = (
                    f"Guideline lacks subtype-specific sections for '{subtype}'. "
                    f"Skill is largely inherited from parent disease."
                )

    preview = build_doctor_preview(
        understanding=merged_understanding,
        diagnosis_knowledge=diagnosis_knowledge,
        treatment_knowledge=treatment_knowledge,
        skills=skills, graph=graph,
        guideline_name=guideline_name,
        guideline_version=guideline_version,
        relationship_store=relationship_store,
        neo4j_push_result=neo4j_push_result,
    )

    all_skill_pages    = _clean_pages([p for s in skills for p in s.get("source_pages", [])])
    final_coverage_pct = round(len(set(all_skill_pages)) / max(total_pages, 1) * 100, 1)

    await save_registries(doctor_id=doctor_id, doc_id=doc_id, discovery=discovery_result, mongo_uri=MONGO_URI)

    await save_subtype_taxonomy(
        doctor_id=doctor_id, doc_id=doc_id,
        primary_disease=primary_disease,
        subtype_parent_map=subtype_parent_map,
        mongo_uri=MONGO_URI,
    )

    result = {
        "doc_id":              doc_id,
        "understanding":       merged_understanding,
        "diagnosis_knowledge": diagnosis_knowledge,
        "treatment_knowledge": treatment_knowledge,
        "subtype_knowledge":   subtype_knowledge,
        "graph":               graph,
        "neo4j_push":          neo4j_push_result,
        "skills":              skills,
        "preview":             preview,
        "saved":               {},
        "discovery_result":    discovery_result,
        "pipeline_version":    "v14",
        "total_pages":         total_pages,
        "sections_processed":  len(sections),
        "sections_indexed_chromadb": stored_count,
        "llm_cache_size":      len(_LLM_CACHE),
        "specialty":           specialty,
        "disease_name":        disease_name,
        "disease_type":        disease_type,
        "embedding_model":     EMBEDDING_MODEL,
        "vector_store":        "chromadb",
        "relationship_summary": {
            "total_relationships": len(relationship_store),
            "relation_types": list({r.relation for r in relationship_store.all_relationships()}),
        },
        "coverage_report": {
            "initial_coverage_pct": coverage_report["coverage_pct"],
            "final_coverage_pct":   final_coverage_pct,
            "covered_pages":        coverage_report["covered_pages"],
            "zero_retrieval_diseases_recovered": gaps,
        },
        "retrieval_summary": {
            d: {qt: len(secs) for qt, secs in ctx.items()}
            for d, ctx in retrieved_by_disease.items()
        },
        "retrieval_config": {
            "embedding_model":         EMBEDDING_MODEL,
            "vector_store":            "chromadb",
            "similarity_threshold":    SIMILARITY_THRESHOLD,
            "max_retrieve_limit":      MAX_RETRIEVE_LIMIT,
            "max_agent_chars":         MAX_AGENT_CHARS,
            "subtype_batch_size":      SUBTYPE_BATCH_SIZE,
            "subtype_min_specificity": SUBTYPE_MIN_SPECIFICITY,
        },
       "quality_warnings": {
            "generic_subtype_skills": [s["skill_index"] for s in skills if s.get("is_generic_subtype")],
            "low_completeness_skills": [
                s["skill_index"] for s in skills
                if len({k.lower() for k in s.get("trigger_keywords", [])}
                       & {k.lower() for k in re.findall(r"[A-Za-z0-9\-]{3,}", json.dumps(s.get("body", {})))}
                   ) / max(len(s.get("trigger_keywords", [])), 1) < 0.5
            ],
        },
        "clinical_skills": clinical_skills,
        "skill_coverage": {
            "diseases": len(all_diseases),
            "subtypes": len(subtypes),
            "skills": len(clinical_skills),
            "llm_derived_skills":        llm_derived_count,
            "structurally_derived_skills": structural_derived_count,
        },
    }

    if save_to_db:
        saved = await save_to_mongodb(
            mongo_uri=MONGO_URI, doctor_id=doctor_id, doc_id=doc_id,
            understanding=merged_understanding,
            diagnosis_knowledge=diagnosis_knowledge,
            treatment_knowledge=treatment_knowledge,
            graph=graph, skills=skills,
            guideline_name=guideline_name,
            guideline_version=guideline_version,
            relationship_store=relationship_store,
            neo4j_push_result=neo4j_push_result,
            clinical_skills=clinical_skills,
        )
        result["saved"] = saved

    return result